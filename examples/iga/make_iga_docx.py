#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_iga_docx.py
KooRemapper IGA (Isogeometric Analysis) 가이드 DOCX 생성 스크립트

실행 방법:
    python make_iga_docx.py
    (examples/iga/ 디렉토리에서 실행)

출력: iga_guide.docx

요구사항:
    pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# 색상 상수
# ─────────────────────────────────────────────
COLOR_H1      = RGBColor(0x1F, 0x38, 0x64)   # Navy Blue
COLOR_H2      = RGBColor(0x17, 0x37, 0x5E)   # Dark Teal
COLOR_H3      = RGBColor(0x2E, 0x2E, 0x2E)   # Dark Gray
COLOR_TABLE_HDR = RGBColor(0xD9, 0xD9, 0xD9) # Light Gray header
COLOR_CODE_BG   = RGBColor(0xF0, 0xF0, 0xF0) # Code background

FONT_BODY = "Malgun Gothic"
FONT_CODE = "Courier New"
FONT_SIZE_BODY = Pt(10)
FONT_SIZE_CODE = Pt(9)

# ─────────────────────────────────────────────
# 헬퍼 함수들
# ─────────────────────────────────────────────

def rgb_to_hex(rgb_color):
    """RGBColor (tuple subclass) to hex string"""
    r, g, b = rgb_color[0], rgb_color[1], rgb_color[2]
    return f"{r:02X}{g:02X}{b:02X}"


def set_cell_background(cell, rgb_color):
    """테이블 셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = rgb_to_hex(rgb_color)
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_border(cell):
    """테이블 셀 테두리 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_paragraph_shading(paragraph, rgb_color):
    """단락 배경색 설정 (코드 블록용)"""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_color = rgb_to_hex(rgb_color)
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), hex_color)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def add_heading1(doc, text):
    """Heading 1: Navy Blue, Bold"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_BODY
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_H1
    # 하단 테두리 (구분선 효과)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_heading2(doc, text):
    """Heading 2: Dark Teal, Bold"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_BODY
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_H2
    return p


def add_heading3(doc, text):
    """Heading 3: Dark Gray, Bold"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_BODY
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_H3
    return p


def add_body(doc, text):
    """본문 단락: Malgun Gothic 10pt"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = FONT_SIZE_BODY
    return p


def add_bullet(doc, text, level=0):
    """글머리 기호 단락"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = FONT_SIZE_BODY
    return p


def add_code_block(doc, code_text):
    """코드 블록: Courier New 9pt, 회색 배경, 각 줄 별도 단락"""
    lines = code_text.split('\n')
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        set_paragraph_shading(p, COLOR_CODE_BG)
        run = p.add_run(line if line else ' ')
        run.font.name = FONT_CODE
        run.font.size = FONT_SIZE_CODE
    # 코드 블록 후 여백
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows, col_widths=None):
    """테이블 생성: 헤더 회색, 테두리 포함"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 헤더 행
    hdr_row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, COLOR_TABLE_HDR)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(9)

    # 데이터 행
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            # 코드 스타일 컬럼 (첫 번째가 파라미터명일 경우)
            if c_idx == 0 and (cell_text.startswith('`') or cell_text.startswith('R') or
                               cell_text.startswith('I') or cell_text.startswith('element_') or
                               cell_text.startswith('bbox_') or cell_text.startswith('target_') or
                               cell_text.startswith('offset') or cell_text.startswith('ir') or
                               cell_text.startswith('styp') or cell_text.startswith('tollg') or
                               cell_text.startswith('nisr') or cell_text.startswith('niss') or
                               cell_text.startswith('nist') or cell_text.startswith('pr') or
                               cell_text.startswith('ps') or cell_text.startswith('pt')):
                run = p.add_run(cell_text.strip('`'))
                run.font.name = FONT_CODE
                run.font.size = Pt(9)
            else:
                run = p.add_run(cell_text)
                run.font.name = FONT_BODY
                run.font.size = Pt(9)

    # 열 너비 설정
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    # 테이블 후 여백
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


# ─────────────────────────────────────────────
# 문서 생성 메인
# ─────────────────────────────────────────────

def create_iga_guide():
    doc = Document()

    # 페이지 여백 설정
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    # ──────────────────────────────────────────
    # 제목
    # ──────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("IGA (Isogeometric Analysis) 가이드")
    title_run.bold = True
    title_run.font.name = FONT_BODY
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = COLOR_H1

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(16)
    sub_run = sub_p.add_run("KooRemapper v1.3.0 기준  |  LS-DYNA R12+ 전용")
    sub_run.font.name = FONT_BODY
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = COLOR_H2

    # ══════════════════════════════════════════
    # 1. 개요
    # ══════════════════════════════════════════
    add_heading1(doc, "1. 개요")

    add_heading2(doc, "1.1 IGA란 무엇인가")
    add_body(doc,
        "IGA(Isogeometric Analysis, 등기하해석)는 CAD와 CAE를 통합하는 수치해석 방법론이다. "
        "기존 유한요소법(FEM)이 Lagrange 다항식 기반의 형상 함수를 사용하는 반면, "
        "IGA는 NURBS(Non-Uniform Rational B-Spline)를 형상 함수로 직접 사용한다."
    )

    add_body(doc, "FEM과 IGA의 주요 차이점:")
    add_table(doc,
        ["항목", "FEM (HEX8/TET4)", "IGA (NURBS)"],
        [
            ["형상 함수", "Lagrange 다항식", "NURBS 기저함수"],
            ["형상 표현", "근사 (절점 보간)", "정확 (CAD와 동일)"],
            ["연속성", "C0 (요소 경계)", "Cp-1 (p차 기준)"],
            ["제어점", "해석 자유도와 일치", "CAD 제어점 역할"],
            ["메시 세분화", "h-refinement (분할)", "k-refinement (차수+세분화 동시)"],
            ["곡면 품질", "메시에 의존", "NURBS로 항상 보장"],
        ],
        col_widths=[3.5, 5.0, 5.5]
    )
    add_body(doc,
        "IGA는 특히 얇은 쉘, 유체-구조 연성(FSI), 접촉 해석에서 FEM 대비 높은 정밀도를 보인다."
    )

    add_heading2(doc, "1.2 LS-DYNA에서의 IGA")
    add_body(doc,
        "LS-DYNA는 R12 버전부터 IGA solid 해석을 지원한다. KooRemapper가 활용하는 핵심 방식은 "
        "Trimmed NURBS Volume으로, 다음과 같은 구조를 가진다:"
    )
    add_bullet(doc, "*IGA_DEV_VOLUME_XYZ + TETMSH=-1 옵션")
    add_bullet(doc, "기존 FE tet/hex mesh를 내부 형상으로 사용 (trimming 경계)")
    add_bullet(doc, "NURBS 직육면체 박스(trivariate B-spline patch)가 FE mesh를 완전히 감싸는 구조")
    add_bullet(doc, "LS-DYNA가 FE mesh의 면을 trim 경계로 인식하여 NURBS 적분점을 내부에만 배치")
    add_body(doc,
        "이 방식의 장점은 기존 FE mesh를 그대로 활용하면서 NURBS의 고차 연속성 및 수렴성을 얻을 수 있다는 점이다."
    )

    add_heading2(doc, "1.3 KooRemapper가 자동화하는 것")
    add_body(doc, "KooRemapper의 iga 명령은 다음 과정을 완전 자동화한다:")
    steps = [
        "Bounding Box 계산: 대상 FE 파트의 모든 절점 좌표에서 최소/최대값 탐색",
        "Offset 확장: YAML 설정에 따른 bbox 확장 (고정값 / 비율 스케일 / 자동)",
        "NURBS 제어점 생성: 2×2×2 직육면체 제어점 배치 (8개 코너)",
        "파라미터 출력: *PARAMETER_LOCAL로 모든 값 기록 (LS-DYNA 참조 가능)",
        "재질 복사: 원본 FE 파트의 MAT 카드를 새 MID로 복사",
        "IGA 키워드 생성: *IGA_DEV_STABILIZATION, *PART, *SECTION_IGA_SOLID, "
        "*IGA_DEV_VOLUME_XYZ, *IGA_SOLID, *IGA_3D_NURBS_XYZ, *IGA_REFINE_SOLID",
        "파일 분리 출력: 파트별 _iga_pN.k 파일 생성",
        "메인 파일에 *INCLUDE 삽입: FE mesh와 IGA 정의를 연결",
    ]
    for i, s in enumerate(steps, 1):
        add_bullet(doc, f"{i}. {s}")

    # ══════════════════════════════════════════
    # 2. 작동 원리
    # ══════════════════════════════════════════
    add_heading1(doc, "2. 작동 원리")

    add_heading2(doc, "2.1 Trimmed NURBS Volume 개념")
    add_body(doc,
        "FE mesh를 trim 경계로 사용하는 방식으로, NURBS 박스가 FE mesh를 감싸는 구조이다:"
    )
    add_code_block(doc,
        "                   NURBS 박스 (IGA_3D_NURBS_XYZ)\n"
        "        ┌─────────────────────────────────┐\n"
        "        │  offset                         │\n"
        "        │  ┌─────────────────────────┐    │\n"
        "        │  │  FE Solid Part          │    │\n"
        "        │  │  (TET4/HEX8 mesh)       │    │\n"
        "        │  │                         │    │\n"
        "        │  └─────────────────────────┘    │\n"
        "        │                  offset         │\n"
        "        └─────────────────────────────────┘"
    )
    add_bullet(doc, "FE mesh: 원본 LS-DYNA FE solid 파트 (trim 경계 역할)")
    add_bullet(doc, "NURBS 박스: FE mesh bbox를 offset만큼 확장한 직육면체")
    add_bullet(doc, "TETMSH=-1: LS-DYNA에 FE mesh의 외면을 trim 경계로 사용하라는 지시")
    add_bullet(doc, "적분점: NURBS 박스 내부 + FE mesh 내부 영역에만 배치됨")

    add_heading2(doc, "2.2 생성 흐름")
    add_code_block(doc,
        "FE K-파일 로드\n"
        "    ↓\n"
        "대상 PID 절점 좌표 수집 (added elements 포함)\n"
        "    ↓\n"
        "BBox 계산: xmin/xmax/ymin/ymax/zmin/zmax\n"
        "    ↓\n"
        "Offset 계산 (bbox_scale or fixed offset or auto)\n"
        "    ↓\n"
        "새 ID 할당: newId = ++maxPartId_\n"
        "새 MID 할당: newMid = ++maxMaterialId_\n"
        "    ↓\n"
        "원본 MAT 블록 추출 + MID 교체\n"
        "    ↓\n"
        "generateIGAContent() → *PARAMETER_LOCAL + 카드 직렬화\n"
        "    ↓\n"
        "igaFiles_ 목록에 추가\n"
        "    ↓\n"
        "writeOutput() 호출 시:\n"
        "  - 메인 K파일: FE mesh 전체 + *INCLUDE _iga_pN.k\n"
        "  - _iga_pN.k: IGA 정의 + *END"
    )

    add_heading2(doc, "2.3 ID 할당 규칙")
    add_body(doc,
        "IGA 파트는 FE 파트와 다른 PID, SECID, MID를 사용해야 한다 (LS-DYNA 요건)."
    )
    add_bullet(doc, "PID = SECID = newId: 하나의 IGA 파트에 PID와 SECID를 동일하게 설정 (단일 NURBS 패치)")
    add_bullet(doc, "newId: 현재 모델의 maxPartId_에 +1")
    add_bullet(doc, "newMid: 현재 모델의 maxMaterialId_에 +1 (FE 파트 MID와 반드시 달라야 함)")

    # ══════════════════════════════════════════
    # 3. YAML 문법
    # ══════════════════════════════════════════
    add_heading1(doc, "3. YAML 문법")

    add_heading2(doc, "3.1 독립 실행 모드 (standalone iga command)")
    add_code_block(doc, "KooRemapper iga config.yaml")
    add_body(doc, "YAML 최상위 필드 구조:")
    add_code_block(doc,
        "base_model: block_2x2x1.k    # 입력 K-파일 (필수)\n"
        "output: result                # 출력 prefix (필수, 확장자 없이)\n"
        "\n"
        "targets:\n"
        "  - target_pid: 1\n"
        "    element_size: 4.0\n"
        "    # ... 추가 파라미터"
    )

    add_heading2(doc, "3.2 assemble 모드 (assemble 명령 내 operation)")
    add_code_block(doc,
        "base_model: model.k\n"
        "output: result\n"
        "\n"
        "operations:\n"
        "  - type: iga\n"
        "    targets:\n"
        "      - target_pid: 1\n"
        "        element_size: 4.0\n"
        "      - target_pid: 2\n"
        "        element_size: 3.0\n"
        "        element_size_r: 2.0"
    )

    add_heading2(doc, "3.3 target 필드 전체 목록")
    add_table(doc,
        ["필드", "타입", "기본값", "설명"],
        [
            ["target_pid",     "int",   "필수",    "대상 FE 파트 PID"],
            ["element_size",   "float", "1.0",     "NURBS 박스의 균일 복셀 크기 (rr=rs=rt)"],
            ["element_size_r", "float", "0.0",     "r(x)방향 복셀 크기 (0이면 element_size 사용)"],
            ["element_size_s", "float", "0.0",     "s(y)방향 복셀 크기 (0이면 element_size 사용)"],
            ["element_size_t", "float", "0.0",     "t(z)방향 복셀 크기 (0이면 element_size 사용)"],
            ["offset",         "float", "-1.0",    "bbox 고정 확장량 (-1이면 auto = element_size 사용)"],
            ["bbox_scale",     "float", "0.0",     "균일 bbox 스케일 (0=비활성; 1.5 = 각 측면 +25% 확장)"],
            ["bbox_scale_r",   "float", "0.0",     "r(x)방향 bbox 스케일 (0이면 bbox_scale 사용)"],
            ["bbox_scale_s",   "float", "0.0",     "s(y)방향 bbox 스케일"],
            ["bbox_scale_t",   "float", "0.0",     "t(z)방향 bbox 스케일"],
            ["ir",             "int",   "0",       "적분 방식 (0=reduced Gauss, 1=full Gauss)"],
            ["styp",           "int",   "4",       "IGA_DEV_STABILIZATION의 LCP 안정화 타입"],
            ["tollg",          "float", "1.0e-3",  "LCP 안정화 임계값"],
            ["pr",             "int",   "1",       "r(x)방향 NURBS 다항식 차수 (1=선형, 2=2차, 3=3차)"],
            ["ps",             "int",   "1",       "s(y)방향 NURBS 다항식 차수"],
            ["pt",             "int",   "1",       "t(z)방향 NURBS 다항식 차수"],
            ["nisr",           "int",   "1",       "r방향 적분점 수"],
            ["niss",           "int",   "1",       "s방향 적분점 수"],
            ["nist",           "int",   "1",       "t방향 적분점 수"],
        ],
        col_widths=[3.5, 1.5, 2.0, 7.0]
    )

    # ══════════════════════════════════════════
    # 4. bbox 확장 계산
    # ══════════════════════════════════════════
    add_heading1(doc, "4. bbox 확장(offset) 계산 방법")
    add_body(doc,
        "NURBS 박스는 FE 파트의 bounding box보다 반드시 커야 한다 "
        "(FE mesh가 박스 안에 완전히 포함되어야 함). "
        "Offset은 각 축 방향으로 bbox를 얼마나 확장할지를 결정한다."
    )

    add_heading2(doc, "4.1 우선순위 (높은 순서)")
    add_code_block(doc,
        "1순위: bbox_scale_r/s/t (축별 스케일)\n"
        "2순위: bbox_scale (균일 스케일)\n"
        "3순위: offset >= 0 (고정 확장량)\n"
        "4순위: auto = element_size (기본, offset=-1일 때)"
    )
    add_body(doc, "각 축에 대해 독립적으로 계산된다.")

    add_heading2(doc, "4.2 수식")
    add_heading3(doc, "bbox_scale 또는 bbox_scale_r/s/t 사용 시:")
    add_code_block(doc,
        "off_axis = (scale - 1.0) / 2.0 × len_axis\n"
        "\n"
        "예) bbox_scale=1.5, lenR=20.0:\n"
        "  offR = (1.5 - 1.0) / 2.0 × 20.0 = 5.0\n"
        "결과: rxminn = xmin - 5.0, rxmaxx = xmax + 5.0"
    )
    add_heading3(doc, "고정 offset 사용 시 (offset >= 0):")
    add_code_block(doc, "offR = offS = offT = offset (지정값 그대로)")
    add_heading3(doc, "자동(auto) 사용 시 (offset = -1, 기본값):")
    add_code_block(doc,
        "offR = rr (element_size_r 또는 element_size)\n"
        "offS = rs (element_size_s 또는 element_size)\n"
        "offT = rt (element_size_t 또는 element_size)"
    )

    add_heading2(doc, "4.3 확장된 bbox 파라미터 (사전 계산)")
    add_body(doc,
        "계산된 확장 bbox는 *PARAMETER_LOCAL에 사전 계산된 값으로 저장된다. "
        "이는 LS-DYNA의 *PARAMETER_EXPRESSION_LOCAL 호환성 문제를 피하기 위한 설계이다."
    )
    add_table(doc,
        ["파라미터", "값 (수식)"],
        [
            ["Rrxminn", "xmin - offR    (NURBS 박스 X 최솟값)"],
            ["Rrxmaxx", "xmax + offR    (NURBS 박스 X 최댓값)"],
            ["Rryminn", "ymin - offS    (NURBS 박스 Y 최솟값)"],
            ["Rrymaxx", "ymax + offS    (NURBS 박스 Y 최댓값)"],
            ["Rrzminn", "zmin - offT    (NURBS 박스 Z 최솟값)"],
            ["Rrzmaxx", "zmax + offT    (NURBS 박스 Z 최댓값)"],
        ],
        col_widths=[3.5, 10.5]
    )

    # ══════════════════════════════════════════
    # 5. NURBS 파라미터
    # ══════════════════════════════════════════
    add_heading1(doc, "5. NURBS 파라미터 의미")

    add_heading2(doc, "5.1 다항식 차수 (pr, ps, pt)")
    add_body(doc, "NURBS B-spline의 차수(degree)를 각 방향별로 지정한다.")
    add_table(doc,
        ["값", "의미", "연속성", "용도"],
        [
            ["1", "선형", "C0", "단순 형상, 빠른 계산"],
            ["2", "2차",  "C1", "응력 구배 표현, 얇은 부품"],
            ["3", "3차",  "C2", "고정밀 해석, 복잡한 응력분포"],
        ],
        col_widths=[1.5, 2.0, 2.5, 8.0]
    )
    add_body(doc, "실용 가이드:")
    add_bullet(doc, "얇은 평판 (두께 방향): pt=1 또는 pt=2로 시작")
    add_bullet(doc, "평면 방향: pr=2, ps=2로 수렴성 향상")
    add_bullet(doc, "계산 비용: 차수가 높을수록 비례하여 증가")

    add_heading2(doc, "5.2 k-refinement와 요소 크기 (rr, rs, rt)")
    add_body(doc,
        "rr/rs/rt는 *IGA_REFINE_SOLID에서 사용하는 목표 요소 크기이다. "
        "k-refinement는 차수를 높이면서 동시에 knot을 삽입하는 IGA 고유의 세분화 방식이다."
    )
    add_code_block(doc,
        "*IGA_REFINE_SOLID\n"
        "  rid    rtyp=2      ← 균일 refinement\n"
        "  hrtyp=2  rr  rs  rt  ← h-refinement 요소 크기 목표\n"
        "  itr=2  its=2  itt=2  ← 각 방향 minimum 분할 수"
    )
    add_body(doc, "rtyp=2: 균일(uniform) refinement / hrtyp=2: h-type (요소 크기 기반) refinement")

    add_heading2(doc, "5.3 적분점 수 (nisr, niss, nist)")
    add_table(doc,
        ["값", "의미"],
        [
            ["1", "기본 (reduced integration에서 충분)"],
            ["2", "정밀 (고차 NURBS에서 권장)"],
        ],
        col_widths=[2.0, 12.0]
    )

    add_heading2(doc, "5.4 적분 방식 (ir)")
    add_table(doc,
        ["값", "명칭", "특성"],
        [
            ["0", "Reduced Gauss", "빠름, 대부분의 경우 충분, 기본값"],
            ["1", "Full Gauss",    "정밀, volumetric locking 억제에 유리"],
        ],
        col_widths=[1.5, 4.0, 8.5]
    )
    add_body(doc,
        "ir=0이 기본이며, 대부분의 고체 해석에 적합하다. "
        "비압축성 재료(고무 등)나 높은 포아송비에서는 ir=1 권장."
    )

    add_heading2(doc, "5.5 LCP 안정화 (styp, tollg)")
    add_body(doc,
        "*IGA_DEV_STABILIZATION은 IGA trimmed volume 경계에서의 "
        "LCP(Linear Complementarity Problem) 안정화를 제어한다."
    )
    add_table(doc,
        ["파라미터", "기본값", "의미"],
        [
            ["styp=4",    "4",     "안정화 타입 (LS-DYNA 내부 정의)"],
            ["tollg=1e-3","1.0e-3","적분점 제거 임계값 (NURBS 박스 경계 근처 처리)"],
        ],
        col_widths=[3.0, 2.5, 8.5]
    )
    add_body(doc,
        "tollg를 너무 크게 하면 필요한 적분점이 제거되어 정밀도 저하, "
        "너무 작으면 경계 불안정 발생 가능."
    )

    # ══════════════════════════════════════════
    # 6. 생성 파일 구조
    # ══════════════════════════════════════════
    add_heading1(doc, "6. 생성 파일 구조")

    add_heading2(doc, "6.1 메인 파일 (<output>.k)")
    add_body(doc,
        "원본 FE mesh를 완전히 보존하며, IGA 파트는 파트 정의 자체는 그대로 둔다 "
        "(FE tetmesh로 사용). 파일 끝에 *INCLUDE로 IGA 파일을 참조한다."
    )
    add_code_block(doc,
        "*KEYWORD\n"
        "*TITLE\n"
        "  (원본 제목)\n"
        "*NODE\n"
        "  (원본 절점 전체)\n"
        "*ELEMENT_SOLID\n"
        "  (원본 요소 전체 - IGA의 fepid로 사용됨)\n"
        "*PART\n"
        "  (원본 FE 파트 정의들)\n"
        "*SECTION_SOLID\n"
        "  (원본 섹션)\n"
        "*MAT_ELASTIC (또는 원본 재질)\n"
        "  (원본 재질 - MID 1)\n"
        "*INCLUDE\n"
        " <output>_iga_p1.k\n"
        "*INCLUDE\n"
        " <output>_iga_p2.k    (다중 파트의 경우)\n"
        "*END"
    )

    add_heading2(doc, "6.2 IGA 파일 (<output>_iga_pN.k)")
    add_body(doc,
        "파트별로 별도 파일이 생성된다. "
        "파일명 형식: <output_basename>_iga_p<PID>.k"
    )
    add_code_block(doc,
        "*KEYWORD\n"
        "$ IGA solid wrapper for FE part N\n"
        "$ Generated by KooRemapper\n"
        "$---+----1----+----2----+----3----+----4----+----5----+----6----+----7----+----8\n"
        "*PARAMETER_LOCAL\n"
        "  (모든 파라미터)\n"
        "*MAT_ELASTIC (또는 원본 재질 타입)\n"
        "  (MID = newMid로 복사)\n"
        "*IGA_DEV_STABILIZATION\n"
        "*PART\n"
        "*SECTION_IGA_SOLID\n"
        "*IGA_DEV_VOLUME_XYZ\n"
        "*IGA_SOLID\n"
        "*IGA_3D_NURBS_XYZ\n"
        "*IGA_REFINE_SOLID\n"
        "*END"
    )

    # ══════════════════════════════════════════
    # 7. 생성 카드 상세 해설
    # ══════════════════════════════════════════
    add_heading1(doc, "7. 생성 카드 상세 해설")

    add_heading2(doc, "7.1 *PARAMETER_LOCAL")
    add_body(doc,
        "IGA 파일의 모든 수치값을 파라미터로 선언한다. "
        "파라미터명 앞의 'I'는 정수(Integer), 'R'은 실수(Real)를 의미한다."
    )
    add_body(doc, "전체 파라미터 목록:")
    add_table(doc,
        ["파라미터명", "타입", "의미"],
        [
            ["Iid",      "int",  "이 IGA 파트의 PID = SECID"],
            ["Imid",     "int",  "IGA 파트용 새 MID (FE와 다름)"],
            ["Ifepid",   "int",  "원본 FE 파트 PID (trim 경계)"],
            ["Rxmin",    "real", "FE bbox X 최솟값"],
            ["Rxmax",    "real", "FE bbox X 최댓값"],
            ["Rymin",    "real", "FE bbox Y 최솟값"],
            ["Rymax",    "real", "FE bbox Y 최댓값"],
            ["Rzmin",    "real", "FE bbox Z 최솟값"],
            ["Rzmax",    "real", "FE bbox Z 최댓값"],
            ["Rrr",      "real", "r방향 NURBS refinement 요소 크기"],
            ["Rrs",      "real", "s방향 NURBS refinement 요소 크기"],
            ["Rrt",      "real", "t방향 NURBS refinement 요소 크기"],
            ["Rofr",     "real", "r방향 bbox 확장량 (offset)"],
            ["Rofs",     "real", "s방향 bbox 확장량"],
            ["Roft",     "real", "t방향 bbox 확장량"],
            ["Iir",      "int",  "적분 방식 (0=reduced, 1=full)"],
            ["Istyp",    "int",  "LCP 안정화 타입"],
            ["Rtollg",   "real", "LCP 임계값"],
            ["Rrxminn",  "real", "NURBS 박스 X 최솟값 (= Rxmin - Rofr)"],
            ["Rrxmaxx",  "real", "NURBS 박스 X 최댓값 (= Rxmax + Rofr)"],
            ["Rryminn",  "real", "NURBS 박스 Y 최솟값 (= Rymin - Rofs)"],
            ["Rrymaxx",  "real", "NURBS 박스 Y 최댓값 (= Rymax + Rofs)"],
            ["Rrzminn",  "real", "NURBS 박스 Z 최솟값 (= Rzmin - Roft)"],
            ["Rrzmaxx",  "real", "NURBS 박스 Z 최댓값 (= Rzmax + Roft)"],
        ],
        col_widths=[3.0, 1.5, 9.5]
    )
    add_body(doc, "파라미터 참조 방식: LS-DYNA 카드에서 &id, &mid, &rxminn 등으로 참조.")
    add_body(doc, "형식 예시 (10-char 고정폭 필드):")
    add_code_block(doc,
        "Iid                3     ← 타입(I/R)+이름을 10자로 패딩, 값 10자 우측정렬\n"
        "Rxmin              0\n"
        "Rrxminn           -4"
    )

    add_heading2(doc, "7.2 재질 카드 (*MAT_*)")
    add_body(doc,
        "원본 FE 파트의 MAT 카드를 그대로 복사하되, MID 필드를 newMid로 교체한다. "
        "원본 MAT를 찾지 못한 경우 경고 주석 + 더미 *MAT_ELASTIC (모든 값 0.0) 삽입."
    )
    add_code_block(doc,
        "*MAT_ELASTIC\n"
        "$#     mid        ro         e        pr        da        db  not used\n"
        "         2  7.85E-09    210000       0.3"
    )

    add_heading2(doc, "7.3 *IGA_DEV_STABILIZATION")
    add_body(doc, "LCP 안정화 설정. Trimmed volume의 경계 처리를 위한 개발자 옵션.")
    add_code_block(doc,
        "*IGA_DEV_STABILIZATION\n"
        "$#      sid      styp                                   tollg\n"
        "       &id     &styp                                  &tollg"
    )
    add_table(doc,
        ["필드", "파라미터", "의미"],
        [
            ["SID",   "&id",    "Set ID (이 IGA 파트 ID와 동일)"],
            ["STYP",  "&styp",  "안정화 타입 (기본: 4)"],
            ["TOLLG", "&tollg", "적분점 제거 임계값 (기본: 1.0e-3)"],
        ],
        col_widths=[2.0, 3.0, 9.0]
    )

    add_heading2(doc, "7.4 *PART")
    add_code_block(doc,
        "*PART\n"
        "$#\n"
        "IGA_Part_1\n"
        "$#     pid     secid       mid     eosid      hgid      grav    adpopt      tmid\n"
        "      &id      &id     &mid"
    )
    add_bullet(doc, "제목 행: IGA_Part_<fepid> (원본 FE PID 번호 포함)")
    add_bullet(doc, "pid = secid = &id: 동일한 ID 공유 (NURBS patch 1개이므로)")
    add_bullet(doc, "mid = &mid: IGA 전용 재질 ID (FE와 반드시 다름)")

    add_heading2(doc, "7.5 *SECTION_IGA_SOLID")
    add_code_block(doc,
        "*SECTION_IGA_SOLID\n"
        "$#   secid    elform        ir\n"
        "      &id        0      &ir"
    )
    add_table(doc,
        ["필드", "값", "의미"],
        [
            ["SECID",  "&id", "섹션 ID (PID와 동일)"],
            ["ELFORM", "0",   "요소 정식화 자동 (LS-DYNA가 선택)"],
            ["IR",     "&ir", "적분 방식 (0=reduced, 1=full)"],
        ],
        col_widths=[2.5, 2.5, 9.0]
    )

    add_heading2(doc, "7.6 *IGA_DEV_VOLUME_XYZ")
    add_body(doc, "Trimmed NURBS volume 정의 (핵심 카드). TETMSH=-1이 FE mesh를 trim 경계로 사용하게 한다.")
    add_code_block(doc,
        "*IGA_DEV_VOLUME_XYZ\n"
        "$#     vid   patchid       pid      esid      fsid    TETMSH      MYTP\n"
        "      &id      &id                                  -1\n"
        "$#     PID of existing FEA solid with tetmesh\n"
        "   &fepid\n"
        "$#   brid1     brid2     brid3     brid4     brid5     brid6     brid7     brid8\n"
        "(빈 행)"
    )
    add_table(doc,
        ["필드", "값", "의미"],
        [
            ["VID",    "&id",     "Volume ID"],
            ["PATCHID","&id",     "NURBS patch ID (*IGA_3D_NURBS_XYZ의 patchid)"],
            ["PID",    "(빈칸)",  "이 볼륨이 속할 IGA 파트 PID"],
            ["TETMSH", "-1",      "FE mesh를 trim 경계로 사용 (핵심!)"],
            ["FE PID", "&fepid",  "trim 경계로 사용할 FE 파트 PID (Card 2)"],
        ],
        col_widths=[2.5, 3.0, 8.5]
    )

    add_heading2(doc, "7.7 *IGA_SOLID")
    add_code_block(doc,
        "*IGA_SOLID\n"
        "$#     sid       pid      nisr      niss      nist       rid\n"
        "      &id      &id        1        1        1      &id"
    )
    add_table(doc,
        ["필드", "값", "의미"],
        [
            ["SID",  "&id",  "Solid set ID"],
            ["PID",  "&id",  "연결할 IGA 파트 PID"],
            ["NISR", "nisr", "r방향 적분점 수"],
            ["NISS", "niss", "s방향 적분점 수"],
            ["NIST", "nist", "t방향 적분점 수"],
            ["RID",  "&id",  "*IGA_REFINE_SOLID의 RID 참조"],
        ],
        col_widths=[2.0, 2.5, 9.5]
    )

    add_heading2(doc, "7.8 *IGA_3D_NURBS_XYZ")
    add_body(doc, "NURBS 볼륨 패치 정의 (B-spline 제어점, knot 벡터).")
    add_code_block(doc,
        "*IGA_3D_NURBS_XYZ\n"
        "$# patchid        nr        ns        nt        pr        ps        pt\n"
        "      &id        2        2        2        1        1        1\n"
        "$#    unir      unis      unit\n"
        "        1        1        1\n"
        "$#            rfirst               rlast\n"
        "             &rxminn             &rxmaxx\n"
        "$#            sfirst               slast\n"
        "             &ryminn             &rymaxx\n"
        "$#            tfirst               tlast\n"
        "             &rzminn             &rzmaxx\n"
        "$#                 x                   y                   z                 wgt\n"
        "             &rxminn             &ryminn             &rzminn                 1.0\n"
        "             &rxmaxx             &ryminn             &rzminn                 1.0\n"
        "             &rxminn             &rymaxx             &rzminn                 1.0\n"
        "             &rxmaxx             &rymaxx             &rzminn                 1.0\n"
        "             &rxminn             &ryminn             &rzmaxx                 1.0\n"
        "             &rxmaxx             &ryminn             &rzmaxx                 1.0\n"
        "             &rxminn             &rymaxx             &rzmaxx                 1.0\n"
        "             &rxmaxx             &rymaxx             &rzmaxx                 1.0"
    )
    add_body(doc, "Card 1 필드 설명:")
    add_table(doc,
        ["필드", "값", "의미"],
        [
            ["PATCHID", "&id", "이 패치의 ID (IGA_DEV_VOLUME_XYZ와 일치)"],
            ["NR",      "2",   "r방향 제어점 수"],
            ["NS",      "2",   "s방향 제어점 수"],
            ["NT",      "2",   "t방향 제어점 수"],
            ["PR",      "pr",  "r방향 다항식 차수"],
            ["PS",      "ps",  "s방향 다항식 차수"],
            ["PT",      "pt",  "t방향 다항식 차수"],
        ],
        col_widths=[2.5, 2.0, 9.5]
    )
    add_body(doc,
        "unir=unis=unit=1: 균일(uniform) knot 분포 지시자. "
        "rfirst/rlast ~ tfirst/tlast: 각 방향 knot 범위 (= NURBS 박스 좌표 범위). "
        "제어점 8개: 2×2×2 코너점, 루프 순서 k(z) → j(y) → i(x), 가중치 wgt=1.0."
    )

    add_heading2(doc, "7.9 *IGA_REFINE_SOLID")
    add_code_block(doc,
        "*IGA_REFINE_SOLID\n"
        "$      rid      rtyp\n"
        "      &id        2\n"
        "$    hrtyp        rr        rs        rt\n"
        "        2      &rr      &rs      &rt\n"
        "$      itr       its       itt\n"
        "        2        2        2"
    )
    add_table(doc,
        ["필드", "값", "의미"],
        [
            ["RID",   "&id", "Refinement ID (*IGA_SOLID에서 참조)"],
            ["RTYP",  "2",   "Refinement 타입: 2=균일(uniform)"],
            ["HRTYP", "2",   "h-refinement 타입: 2=요소 크기 기반"],
            ["RR",    "&rr", "r방향 목표 요소 크기"],
            ["RS",    "&rs", "s방향 목표 요소 크기"],
            ["RT",    "&rt", "t방향 목표 요소 크기"],
            ["ITR",   "2",   "r방향 최소 refinement 반복 횟수"],
            ["ITS",   "2",   "s방향 최소 refinement 반복 횟수"],
            ["ITT",   "2",   "t방향 최소 refinement 반복 횟수"],
        ],
        col_widths=[2.0, 2.5, 9.5]
    )

    # ══════════════════════════════════════════
    # 8. 예제 파일
    # ══════════════════════════════════════════
    add_heading1(doc, "8. 예제 파일")

    add_heading2(doc, "8.1 block_2x2x1.k (기본 FE 베이스 모델)")
    add_body(doc, "20×10mm 평면 2×2 HEX8 메시, z방향 두 층(각 5mm 높이):")
    add_table(doc,
        ["항목", "내용"],
        [
            ["전체 크기",   "x=0~20, y=0~10, z=0~10 (mm)"],
            ["절점 수",     "18개 (6×3 그리드)"],
            ["요소 수",     "4개 HEX8 (2×2 패턴)"],
            ["PID 1 (하부)", "Lower layer, z=0~5, eid=1,2"],
            ["PID 2 (상부)", "Upper layer, z=5~10, eid=3,4"],
            ["재질",        "MAT_ELASTIC: ro=7.85e-9 (t/mm/s), E=210000 MPa, nu=0.3"],
            ["섹션",        "SECTION_SOLID: secid=1, elform=1 (constant stress)"],
        ],
        col_widths=[3.5, 10.5]
    )

    add_heading2(doc, "8.2 iga_single.yaml (최소 설정)")
    add_code_block(doc,
        "# 단일 파트 IGA - 최소 설정\n"
        "base_model: block_2x2x1.k\n"
        "output: iga_single_result\n"
        "\n"
        "targets:\n"
        "  - target_pid: 1\n"
        "    element_size: 4.0"
    )
    add_body(doc, "결과 해석:")
    add_table(doc,
        ["항목", "값"],
        [
            ["PID 1 bbox", "x[0,20], y[0,10], z[0,5]"],
            ["auto offset", "4.0 (= element_size)"],
            ["NURBS 박스", "x[-4,24], y[-4,14], z[-4,9]"],
            ["rr=rs=rt", "4.0 (균일 복셀)"],
            ["ir/styp/tollg", "0 / 4 / 1.0e-3 (기본값)"],
            ["pr/ps/pt", "1/1/1 (선형, 기본값)"],
            ["신규 IGA id", "3 (maxPartId+1)"],
            ["신규 IGA mid", "2 (maxMaterialId+1)"],
        ],
        col_widths=[4.0, 10.0]
    )

    add_heading2(doc, "8.3 iga_multipart.yaml (다중 파트 + 축별 복셀)")
    add_code_block(doc,
        "base_model: block_2x2x1.k\n"
        "output: iga_multipart_result\n"
        "\n"
        "targets:\n"
        "  - target_pid: 1\n"
        "    element_size: 4.0     # 균일 복셀 4mm\n"
        "    ir: 0                 # reduced Gauss\n"
        "    styp: 4\n"
        "    tollg: 1.0e-3\n"
        "  - target_pid: 2\n"
        "    element_size: 3.0     # t방향 기본값\n"
        "    element_size_r: 2.0   # r방향 더 촘촘\n"
        "    element_size_s: 2.0   # s방향 더 촘촘\n"
        "    ir: 1                 # full Gauss (정밀)\n"
        "    pr: 2                 # r방향 2차\n"
        "    ps: 2                 # s방향 2차\n"
        "    pt: 1                 # t방향 1차 유지"
    )
    add_body(doc, "PID 2 결과: rr=2.0, rs=2.0, rt=3.0 / NURBS 박스: x[-2,22], y[-2,12], z[2,13]")

    add_heading2(doc, "8.4 iga_scale.yaml (bbox_scale 비율 확장)")
    add_code_block(doc,
        "base_model: block_2x2x1.k\n"
        "output: iga_scale_result\n"
        "\n"
        "targets:\n"
        "  - target_pid: 1\n"
        "    element_size: 4.0\n"
        "    bbox_scale: 1.5         # 각 변 × 1.5 (양측 +25%)\n"
        "\n"
        "  - target_pid: 2\n"
        "    element_size: 3.0\n"
        "    bbox_scale_r: 2.0       # r: +50% 양측\n"
        "    bbox_scale_s: 1.3       # s: +15% 양측\n"
        "    bbox_scale_t: 3.0       # t: +100% 양측\n"
        "    ir: 1"
    )
    add_body(doc, "PID 1 계산: lenR=20→offR=5.0, lenS=10→offS=2.5, lenT=5→offT=1.25")
    add_body(doc, "PID 2 계산: offR=10.0, offS=1.5, offT=5.0 (축별 scale 독립 적용)")

    # ══════════════════════════════════════════
    # 9. 파라미터 선택 가이드
    # ══════════════════════════════════════════
    add_heading1(doc, "9. 주요 파라미터 선택 가이드")

    add_heading2(doc, "9.1 element_size 설정")
    add_body(doc, "권장: FE mesh 평균 요소 크기의 1~4배")
    add_code_block(doc,
        "얇은 판재 (두께 방향 HEX8 1층):\n"
        "  element_size_t: 두께/2 ~ 두께×2 (두께 방향 NURBS 간격)\n"
        "  element_size_r/s: 면내 HEX 크기의 2~4배\n"
        "\n"
        "3D 솔리드:\n"
        "  element_size: FE 요소 크기의 2~3배 (균일)"
    )
    add_body(doc,
        "너무 작은 element_size → 과도한 NURBS refinement → 계산 비용 급증. "
        "너무 큰 element_size → 해상도 부족 → 응력 분포 불량."
    )

    add_heading2(doc, "9.2 offset vs bbox_scale 선택")
    add_table(doc,
        ["방식", "YAML", "용도"],
        [
            ["고정 offset",   "offset: 2.0",        "파트 크기와 무관하게 일정한 여유 확보"],
            ["비율 scale",    "bbox_scale: 1.2",     "파트 크기 대비 일정 비율(20%) 여유"],
            ["auto (기본)",   "(offset 미지정)",     "element_size를 자동으로 offset 적용"],
        ],
        col_widths=[3.0, 4.0, 7.0]
    )

    add_heading2(doc, "9.3 다항식 차수 선택")
    add_code_block(doc,
        "평판/쉘 유사 파트:     pr: 2, ps: 2, pt: 1    (두께 방향은 선형)\n"
        "3D 볼륨 고정밀:        pr: 2, ps: 2, pt: 2    (전 방향 2차)\n"
        "빠른 초기 계산:        pr: 1, ps: 1, pt: 1    (선형, 기본값)"
    )

    add_heading2(doc, "9.4 적분 방식 선택")
    add_code_block(doc,
        "일반 구조 해석:     ir: 0  (reduced, 기본)\n"
        "비압축성 재료:      ir: 1  (full)\n"
        "고차 NURBS (p≥3): ir: 1  권장"
    )

    # ══════════════════════════════════════════
    # 10. LS-DYNA 실행
    # ══════════════════════════════════════════
    add_heading1(doc, "10. LS-DYNA 실행")

    add_heading2(doc, "10.1 실행 방법")
    add_body(doc, "IGA 파일은 단독 실행 불가 → 반드시 메인 K 파일로 실행:")
    add_code_block(doc,
        "# 단일 CPU\n"
        "ls-dyna i=iga_single_result.k\n"
        "\n"
        "# 병렬 (4 CPU)\n"
        "ls-dyna i=iga_single_result.k ncpu=4\n"
        "\n"
        "# MPP 버전\n"
        "mpirun -np 4 ls-dyna_mpp i=iga_single_result.k"
    )

    add_heading2(doc, "10.2 예상 Warning 메시지 (정상)")
    add_table(doc,
        ["Warning 번호", "내용", "원인"],
        [
            ["11541", "Structured deck disabled",  "IGA 개발자 키워드 사용 시 정상"],
            ["30128", "Massless node detected",     "NURBS 제어점은 질량 없음 (정상)"],
            ["30131", "Node not attached",          "NURBS 박스 내 미사용 절점"],
        ],
        col_widths=[3.0, 5.0, 6.0]
    )

    add_heading2(doc, "10.3 IGA 해석 결과 확인")
    add_bullet(doc, "FE mesh 부분은 기존과 동일하게 d3plot에서 확인 가능")
    add_bullet(doc, "IGA 파트 응력은 별도 IGA 결과 데이터베이스에 저장됨")
    add_bullet(doc, "LS-PrePost에서 IGA 결과 확인: Results → IGA → ...")

    # ══════════════════════════════════════════
    # 11. 알려진 제약사항
    # ══════════════════════════════════════════
    add_heading1(doc, "11. 알려진 제약사항")
    add_table(doc,
        ["제약", "내용"],
        [
            ["Shell 파트 불가",    "Solid 파트만 지원 (ELEMENT_SOLID)"],
            ["단일 NURBS 박스",    "L형/C형/복잡한 형상은 단일 직육면체로 표현 불가"],
            ["FE mesh 포함 필수",  "FE mesh가 NURBS 박스 안에 완전히 포함되어야 함"],
            ["LS-DYNA R12+ 필요", "*IGA_DEV_* 키워드는 R12 이상에서만 지원"],
            ["MID 공유 불가",      "IGA 파트와 FE 파트는 반드시 다른 MID 사용"],
            ["offset > 0 권장",   "NURBS 박스가 FE mesh와 정확히 일치하면 경계 문제 발생 가능"],
        ],
        col_widths=[4.0, 10.0]
    )

    # ══════════════════════════════════════════
    # 12. 변경 이력
    # ══════════════════════════════════════════
    add_heading1(doc, "12. 변경 이력")
    add_table(doc,
        ["버전", "내용"],
        [
            ["v1.2.x", "초기 IGA 기능 구현 (단일 파트, assemble op 내 지원)"],
            ["v1.3.0", "다중 파트 지원, bbox_scale (균일/축별), 축별 element_size, "
                       "standalone iga 명령 추가, *PARAMETER_EXPRESSION_LOCAL 호환성 문제 해결"],
        ],
        col_widths=[2.5, 11.5]
    )

    # ══════════════════════════════════════════
    # 부록 A. 전체 IGA 파일 예시
    # ══════════════════════════════════════════
    add_heading1(doc, "부록 A. 전체 IGA 파일 예시 (iga_single_result_iga_p1.k)")
    add_code_block(doc,
        "*KEYWORD\n"
        "$ IGA solid wrapper for FE part 1\n"
        "$ Generated by KooRemapper\n"
        "$---+----1----+----2----+----3----+----4----+----5----+----6----+----7----+----8\n"
        "*PARAMETER_LOCAL\n"
        "$    PRMR1      VAL1\n"
        "Iid                3\n"
        "Imid               2\n"
        "Ifepid             1\n"
        "Rxmin              0\n"
        "Rxmax             20\n"
        "Rymin              0\n"
        "Rymax             10\n"
        "Rzmin              0\n"
        "Rzmax              5\n"
        "Rrr                4\n"
        "Rrs                4\n"
        "Rrt                4\n"
        "Rofr               4\n"
        "Rofs               4\n"
        "Roft               4\n"
        "Iir                0\n"
        "Istyp              4\n"
        "Rtollg         0.001\n"
        "Rrxminn           -4\n"
        "Rrxmaxx           24\n"
        "Rryminn           -4\n"
        "Rrymaxx           14\n"
        "Rrzminn           -4\n"
        "Rrzmaxx            9\n"
        "*MAT_ELASTIC\n"
        "$#     mid        ro         e        pr        da        db  not used\n"
        "         2  7.85E-09    210000       0.3\n"
        "*IGA_DEV_STABILIZATION\n"
        "$#      sid      styp                                   tollg\n"
        "       &id     &styp                                  &tollg\n"
        "*PART\n"
        "$#\n"
        "IGA_Part_1\n"
        "$#     pid     secid       mid     eosid      hgid      grav    adpopt      tmid\n"
        "      &id      &id     &mid\n"
        "*SECTION_IGA_SOLID\n"
        "$#   secid    elform        ir\n"
        "      &id        0      &ir\n"
        "*IGA_DEV_VOLUME_XYZ\n"
        "$#     vid   patchid       pid      esid      fsid    TETMSH      MYTP\n"
        "      &id      &id                                  -1\n"
        "$#     PID of existing FEA solid with tetmesh\n"
        "   &fepid\n"
        "$#   brid1     brid2     brid3     brid4     brid5     brid6     brid7     brid8\n"
        "\n"
        "*IGA_SOLID\n"
        "$#     sid       pid      nisr      niss      nist       rid\n"
        "      &id      &id        1        1        1      &id\n"
        "*IGA_3D_NURBS_XYZ\n"
        "$# patchid        nr        ns        nt        pr        ps        pt\n"
        "      &id        2        2        2        1        1        1\n"
        "$#    unir      unis      unit\n"
        "        1        1        1\n"
        "$#            rfirst               rlast\n"
        "             &rxminn             &rxmaxx\n"
        "$#            sfirst               slast\n"
        "             &ryminn             &rymaxx\n"
        "$#            tfirst               tlast\n"
        "             &rzminn             &rzmaxx\n"
        "$#                 x                   y                   z                 wgt\n"
        "             &rxminn             &ryminn             &rzminn                 1.0\n"
        "             &rxmaxx             &ryminn             &rzminn                 1.0\n"
        "             &rxminn             &rymaxx             &rzminn                 1.0\n"
        "             &rxmaxx             &rymaxx             &rzminn                 1.0\n"
        "             &rxminn             &ryminn             &rzmaxx                 1.0\n"
        "             &rxmaxx             &ryminn             &rzmaxx                 1.0\n"
        "             &rxminn             &rymaxx             &rzmaxx                 1.0\n"
        "             &rxmaxx             &rymaxx             &rzmaxx                 1.0\n"
        "*IGA_REFINE_SOLID\n"
        "$      rid      rtyp\n"
        "      &id        2\n"
        "$    hrtyp        rr        rs        rt\n"
        "        2      &rr      &rs      &rt\n"
        "$      itr       its       itt\n"
        "        2        2        2\n"
        "*END"
    )

    # ══════════════════════════════════════════
    # 부록 B. 자주 묻는 질문
    # ══════════════════════════════════════════
    add_heading1(doc, "부록 B. 자주 묻는 질문")

    faq_items = [
        (
            "Q: IGA 파트와 FE 파트가 같은 공간을 공유해도 되나요?",
            "A: 네. IGA 파트(NURBS 박스)는 FE 파트와 물리적으로 겹칩니다. "
            "LS-DYNA가 TETMSH=-1과 fepid를 통해 FE mesh를 trim 경계로 사용하기 때문에, "
            "IGA와 FE 해석 영역은 동일한 형상을 공유합니다."
        ),
        (
            "Q: FE와 IGA 파트 간 접촉/구속은 어떻게 처리하나요?",
            "A: TETMSH=-1 방식에서 LS-DYNA가 내부적으로 coupling을 처리합니다. "
            "별도 *CONSTRAINED_* 카드가 필요하지 않습니다."
        ),
        (
            "Q: 여러 번 assemble op를 거친 후 IGA를 적용해도 되나요?",
            "A: 네. applyIGA()는 addedElements_와 modifiedNodePositions_도 포함하여 "
            "bbox를 계산하므로, replace/restack 이후에 적용해도 정상 동작합니다."
        ),
        (
            "Q: offset을 0으로 설정하면 어떻게 되나요?",
            "A: offset: 0.0으로 설정하면 NURBS 박스 = FE bbox (확장 없음). "
            "FE mesh 표면이 NURBS 박스와 정확히 일치하게 되어 경계 처리 문제가 발생할 수 있습니다. "
            "최소 요소 크기의 10% 이상 offset을 권장합니다."
        ),
        (
            "Q: *MAT_024 같은 탄소성 재질도 지원하나요?",
            "A: 네. extractMaterialBlock()이 *MAT_* 키워드 종류와 무관하게 원본 카드를 복사합니다. "
            "MAT_ELASTIC, MAT_024, MAT_RIGID 등 모두 지원됩니다."
        ),
        (
            "Q: 생성된 IGA 파일을 수동으로 편집해도 되나요?",
            "A: *PARAMETER_LOCAL의 값을 수정하면 연결된 모든 카드에 자동 반영됩니다. "
            "단, 파라미터명 형식(Iname, Rname, 10자 필드)을 유지해야 합니다."
        ),
    ]
    for q, a in faq_items:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(6)
        p_q.paragraph_format.space_after = Pt(2)
        run_q = p_q.add_run(q)
        run_q.bold = True
        run_q.font.name = FONT_BODY
        run_q.font.size = FONT_SIZE_BODY
        add_body(doc, a)

    # 저장
    output_path = "iga_guide.docx"
    doc.save(output_path)
    print(f"[OK] 생성 완료: {output_path}")
    print(f"     페이지 여백: 상2.5cm 하2.0cm 좌2.5cm 우2.0cm")
    print(f"     포함 섹션: 12개 + 부록 2개")


if __name__ == "__main__":
    create_iga_guide()
