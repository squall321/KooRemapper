"""
KooRemapper Explicit Stabilization Guide generator
python-docx 기반 Word 문서 생성
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 기본 스타일 설정 ───────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = 'Malgun Gothic'
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x4A, 0x86, 0xC8)
    return p

def add_para(text, bold=False, italic=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent * 0.5)
    run = p.add_run(text)
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(10)
    run.bold = bold
    run.italic = italic
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '2E74B5')
        run = cell.paragraphs[0].add_run(h)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_note(text, color='FFF3CD'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run('📌 ' + text)
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x7B, 0x52, 0x00)
    return p

def add_warning(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run('⚠ ' + text)
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return p

# ════════════════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('KooRemapper')
run.font.name = 'Malgun Gothic'
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Explicit 해석 안정화 가이드')
run.font.name = 'Malgun Gothic'
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('stabilize explicit — 12단계 레벨 시스템')
run.font.name = 'Malgun Gothic'
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

doc.add_paragraph()
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 1. 개요
# ════════════════════════════════════════════════════════════════
add_heading('1. 개요', 1)
add_para(
    'LS-DYNA explicit 해석에서 발산(divergence) 또는 불안정 현상이 발생할 때 '
    '단계적으로 적용할 수 있는 12단계 안정화 레벨 시스템입니다.'
)
doc.add_paragraph()
add_para('기본 사용법:', bold=True)
add_code('KooRemapper stabilize config.yaml')
doc.add_paragraph()

add_para('YAML 최소 구성:', bold=True)
add_code('model:  your_model.k')
add_code('output: your_model_stable.k')
add_code('stabilize: explicit')
add_code('level: 5          # 1~12')
doc.add_paragraph()

add_para('레벨 없이 수동 완전 제어:', bold=True)
add_code('model:  your_model.k')
add_code('output: your_model_stable.k')
add_code('stabilize: explicit')
add_code('# level 키 생략 → 명시된 옵션만 적용')
add_code('tssfac: 0.67')
add_code('ihq: 6')
add_code('soft: 2')
doc.add_paragraph()

add_para('레벨 기반 + 일부 override:', bold=True)
add_code('stabilize: explicit')
add_code('level: 7           # 레벨 기본값 적용')
add_code('tssfac: 0.60       # 이 항목만 override')
add_code('soft: 1            # 이 항목만 override')

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 2. Implicit vs Explicit 발산 차이
# ════════════════════════════════════════════════════════════════
add_heading('2. Implicit vs Explicit 발산 구조 차이', 1)

add_table(
    ['구분', 'Implicit', 'Explicit'],
    [
        ['발산 원인', '비선형 뉴턴법 수렴 실패 (단일)', '시간스텝/호글래스/접촉/요소변형 (복합)'],
        ['해결 방향', '솔버 파라미터 강화 (단일 축)', '원인별 옵션 조합 (복수 축)'],
        ['레벨 의미', '더 많은 반복, 더 작은 허용치', '원인에 따른 옵션 선택'],
        ['최후 수단', 'Arc-length (하중경로 추적)', '요소 침식 (ERODE)'],
        ['물리 변경', '없음 (솔버 파라미터만)', 'ERODE 시 질량/체적 비보존'],
    ],
    col_widths=[2.5, 6.5, 6.5]
)
doc.add_paragraph()

add_para('Explicit 발산의 4개 독립 축:', bold=True)
axes = [
    ('Axis 1: 시간 스텝', 'TSSFAC, DT2MS(경고), ERODE'),
    ('Axis 2: 호글래스 모드', 'IHQ, QH'),
    ('Axis 3: 접촉 침투/채터링', 'SOFT, SBOPT, DEPTH, IGNORE'),
    ('Axis 4: 재료/요소 발산', 'OSU, INN, ESORT, 쉘 옵션'),
]
for name, opts in axes:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name + ':  ')
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(10)
    run.font.bold = True
    run2 = p.add_run(opts)
    run2.font.name = 'Malgun Gothic'
    run2.font.size = Pt(10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 3. 12단계 레벨 스펙트럼
# ════════════════════════════════════════════════════════════════
add_heading('3. 12단계 레벨 스펙트럼', 1)
add_para('각 레벨은 이전 레벨의 모든 설정을 포함합니다 (누적 적용).')
doc.add_paragraph()

level_overview = [
    ('1', '에너지 진단', '없음', 'CONTROL_ENERGY 전체, ISLCHK=2, WRPANG=10'),
    ('2', '정확도 기반', '없음', 'OSU=1, INN=2, ESORT(solid+shell)=1'),
    ('3', '시간 스텝 1', '11% 감소', 'TSSFAC=0.80, INN=4'),
    ('4', '호글래스 1', '소', 'IHQ=4 (FB stiffness), QH=0.10'),
    ('5', '쉘 안정화', '소~중 ★쉘만', 'BWC=1, MITER=2, IRNXX=-2 / 쉘 없으면 skip'),
    ('6', '접촉 1단계', '중', 'SOFT=1, SBOPT=2, DEPTH=3, ORIEN=2, SHLTHK=1, XPENE=2.0'),
    ('7', '시간 2 + 체적점성', '25% 감소', 'TSSFAC=0.67, Q1=1.5/Q2=0.06 강제 설정'),
    ('8', '접촉 2단계', '중~대', 'SOFT=2(pinball), SBOPT=3, DEPTH=5, IGNORE=1, MAXPAR=1.15, NSBCS=5'),
    ('9', '호글래스 최강', '대', 'IHQ=6 (BB), QH=1.0'),
    ('10', '시간 3', '35% 감소', 'TSSFAC=0.60, NSBCS=2'),
    ('11', '침식 활성화', '대 ★물리 변경', 'ERODE=11, ENMASS=2, NSBCS=1, TSSFAC=0.55'),
    ('12', '최대 보수', '최대', 'TSSFAC=0.50, Q1=2.0/Q2=0.10'),
]

add_table(
    ['Lv', '이름', '연산비용 영향', '핵심 추가 설정'],
    level_overview,
    col_widths=[0.8, 2.8, 2.5, 9.4]
)
doc.add_paragraph()

add_note('레벨 5는 *ELEMENT_SHELL 키워드가 없는 순수 솔리드 모델에서 자동으로 skip됩니다.')
add_warning('레벨 11은 요소 침식(ERODE)으로 물리가 변경됩니다. confirm_erosion: true 또는 interactive 확인 필요.')
doc.add_paragraph()

# ── 누적 상태 스냅샷 ─────────────────────────────────────────
add_heading('3.1 누적 상태 스냅샷', 2)
add_table(
    ['Lv', 'TSSFAC', 'IHQ', 'QH', 'SOFT', 'SBOPT', 'DEPTH', 'INN', 'Q1/Q2', 'ERODE', 'NSBCS'],
    [
        ['1',  '0.90', '1(기본)', '0.10', '0(기본)', '0', '2', '1', '기본',     '-', '-'],
        ['2',  '0.90', '1',       '0.10', '0',       '0', '2', '2', '기본',     '-', '-'],
        ['3',  '0.80', '1',       '0.10', '0',       '0', '2', '4', '기본',     '-', '-'],
        ['4',  '0.80', '4',       '0.10', '0',       '0', '2', '4', '기본',     '-', '-'],
        ['5',  '0.80', '4',       '0.10', '0',       '0', '2', '4', '기본',     '-', '-'],
        ['6',  '0.80', '4',       '0.10', '1',       '2', '3', '4', '기본',     '-', '-'],
        ['7',  '0.67', '4',       '0.10', '1',       '2', '3', '4', '1.5/0.06', '-', '-'],
        ['8',  '0.67', '4',       '0.10', '2',       '3', '5', '4', '1.5/0.06', '-', '5'],
        ['9',  '0.67', '6',       '1.00', '2',       '3', '5', '4', '1.5/0.06', '-', '5'],
        ['10', '0.60', '6',       '1.00', '2',       '3', '5', '4', '1.5/0.06', '-', '2'],
        ['11', '0.55', '6',       '1.00', '2',       '3', '5', '4', '1.5/0.06', '11','1'],
        ['12', '0.50', '6',       '1.00', '2',       '3', '5', '4', '2.0/0.10', '11','1'],
    ],
    col_widths=[0.7, 1.3, 1.5, 0.9, 1.1, 1.2, 1.2, 0.9, 1.8, 1.3, 1.2]
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 4. 옵션 상세 레퍼런스
# ════════════════════════════════════════════════════════════════
add_heading('4. 옵션 상세 레퍼런스', 1)

# ── 4.1 시간 스텝 ────────────────────────────────────────────
add_heading('4.1 시간 스텝  (*CONTROL_TIMESTEP)', 2)
add_table(
    ['옵션', '기본값', '레벨 적용값', '설명'],
    [
        ['tssfac', '0.90', 'Lv3:0.80 → Lv7:0.67 → Lv10:0.60 → Lv11:0.55 → Lv12:0.50',
         'Courant 안전계수. dt_actual = tssfac × dt_critical. 낮을수록 안정, 연산비용 증가.'],
        ['erode', '0', 'Lv11: 11 (solid+shell)',
         '요소 침식. 1=솔리드, 10=쉘, 11=솔리드+쉘, 111=전체. ⚠물리 변경.'],
        ['DT2MS', '0', 'WARNING only',
         'Mass scaling. 관성력 왜곡 → 동적 정확도 손실. 경고 출력만, 강제 변경 없음.'],
    ],
    col_widths=[2.0, 2.2, 5.0, 6.3]
)
doc.add_paragraph()

add_para('TSSFAC 값 선택 기준:', bold=True)
add_table(
    ['TSSFAC', '적용 레벨', '권장 케이스'],
    [
        ['0.90', '기본', 'LS-DYNA 기본. 대부분의 금속 충격 해석.'],
        ['0.80', 'Lv 3', '첫 보수화. 가벼운 진동/불안정.'],
        ['0.67', 'Lv 7', '고무, 폼 표준 권장값. LS-DYNA 공식 권고.'],
        ['0.60', 'Lv10', '심각한 요소 변형 진행 중.'],
        ['0.55', 'Lv11', '침식 모드 병행.'],
        ['0.50', 'Lv12', '최대 안전. 충격파/폭발/HVI 해석.'],
    ],
    col_widths=[2.0, 2.5, 11.0]
)
doc.add_paragraph()

# ── 4.2 호글래스 ────────────────────────────────────────────
add_heading('4.2 호글래스  (*CONTROL_HOURGLASS)', 2)
add_para('글로벌 기본값 설정. 파트별 *HOURGLASS 카드가 있으면 그쪽이 우선됨.', italic=True)
doc.add_paragraph()

add_table(
    ['IHQ', '이름', '특성', '주요 용도', '레벨'],
    [
        ['1', '표준 점성', '속도 비례 감쇠. 기본값.', '충격, 폭발', '기본'],
        ['2', 'FB 점성', 'Flanagan-Belytschko 점성.', '일반', '-'],
        ['3', 'FB 점성+체적', 'FB 점성 + 정적 체적 보정.', '비압축 개선', '-'],
        ['4', 'FB 강성', '변위 비례. 준정적/벤딩에 강함.', '금속 준정적, 벤딩 지배', 'Lv 4'],
        ['5', 'FB 강성 (rubber)', 'IHQ=4와 동일, 계수 다름.', '고무/폼', '-'],
        ['6', 'BB 물리 안정화', 'Assumed strain. 가장 정확.', '비압축 대변형, 고무', 'Lv 9'],
    ],
    col_widths=[1.0, 2.5, 4.5, 3.5, 1.5]
)
doc.add_paragraph()

add_table(
    ['옵션', '기본값', 'IHQ 1~5 권장', 'IHQ 6+ 권장', '설명'],
    [
        ['QH', '0.10', '0.05~0.15 (0.15 초과 시 불안정)', '1.0 표준',
         'IHQ=6에서 QH는 assumed strain 스케일. QH=1.0이 완전 안정화.'],
    ],
    col_widths=[1.5, 2.0, 4.0, 2.5, 5.5]
)
doc.add_paragraph()

# ── 4.3 정확도 ────────────────────────────────────────────
add_heading('4.3 정확도  (*CONTROL_ACCURACY)', 2)
add_table(
    ['옵션', '기본', '적용값', '레벨', '설명'],
    [
        ['OSU', '0', '1', 'Lv 2', '2차 Jaumann 객관 응력 업데이트. 대변형/대회전 허위 응력 방지.'],
        ['INN', '1', '2→4', 'Lv 2→3', '불변량 노드 번호. 2=쉘 한정. 4=솔리드 포함. 대변형 수치 일관성.'],
    ],
    col_widths=[1.5, 1.2, 1.8, 1.5, 9.5]
)
doc.add_paragraph()

# ── 4.4 쉘 안정화 ─────────────────────────────────────────
add_heading('4.4 쉘 안정화  (*CONTROL_SHELL)  ── ELEMENT_SHELL 없으면 자동 skip', 2)
add_table(
    ['옵션', '기본', '적용값', '레벨', '설명'],
    [
        ['BWC', '2', '1', 'Lv 5', '뒤틀린 Belytschko-Tsay 쉘에 워핑 강성 추가. 국부 불안정 억제.'],
        ['MITER', '1', '2', 'Lv 5', '쉘 소성 반복. 1=3 secant(기본). 2=Full iterative (대변형 정확도↑).'],
        ['IRNXX', '-1', '-2', 'Lv 5', '노달 파이버 업데이트. -1=매 cycle 재계산. -2=증분 업데이트 (대회전 일관성).'],
        ['ESORT', '0', '1', 'Lv 2', '퇴화 QUAD → TRI 자동 전환. d3hsp 경고 감소.'],
        ['WRPANG', '20°', '10°', 'Lv 1', '워핑 경고 각도. 낮을수록 엄격한 진단.'],
    ],
    col_widths=[1.8, 1.2, 1.8, 1.5, 9.2]
)
doc.add_paragraph()

# ── 4.5 접촉 전역 ─────────────────────────────────────────
add_heading('4.5 접촉 전역  (*CONTROL_CONTACT)', 2)
add_table(
    ['옵션', '기본', '적용값', '레벨', '설명'],
    [
        ['ORIEN', '1', '2', 'Lv 6', '세그먼트 법선 자동 재정렬. 1=part만. 2=수동 정의 포함.'],
        ['SHLTHK', '0', '1→2', 'Lv6→8', '쉘 두께를 접촉 오프셋에 포함. 0=미포함. 1=비강체만. 2=전체.'],
        ['XPENE', '4.0', '2.0', 'Lv 6', '최대 침투 허용 배수. 낮을수록 엄격한 침투 검사.'],
        ['ENMASS', '0', '1→2', 'Lv8→11', '침식 노드 접촉 유지. 0=제거. 1=솔리드. 2=솔리드+쉘.'],
        ['NSBCS', '자동', '5→2→1', 'Lv8→10→11', '버킷 검색 주기(cycle). 낮을수록 자주 검색. NSBCS=1: 매 cycle.'],
        ['ISLCHK', '1', '2', 'Lv 1', '초기 침투 전수 검사. 결과를 d3hsp에 보고 (진단 전용).'],
    ],
    col_widths=[1.8, 1.3, 1.8, 1.8, 8.8]
)
doc.add_paragraph()

# ── 4.6 접촉 per-contact ──────────────────────────────────
add_heading('4.6 접촉 per-contact  (모든 *CONTACT_* 카드 일괄 적용)', 2)
add_table(
    ['옵션', '기본', 'Lv6 값', 'Lv8+ 값', '설명'],
    [
        ['SOFT', '0', '1', '2',
         '0=표준(bulk modulus 기반). 1=soft constraint(이종 재료). 2=pinball(대변형 최강).'],
        ['SBOPT', '0', '2', '3',
         '0=표준. 2=뒤틀린 세그먼트 법선 보정. 3=edge-to-edge + 뒤틀림 보정.'],
        ['DEPTH', '2', '3', '5',
         '접촉 검색 깊이(요소 레이어 수). 높을수록 두꺼운 부품 통과 접촉도 감지.'],
        ['IGNORE', '0', '-', '1',
         '초기 침투 처리. 0=무시(발산 원인). 1=offset 처리. 조립체에서 거의 필수.'],
        ['MAXPAR', '1.025', '-', '1.15',
         'SOFT=2 전용. 세그먼트 크기 대비 침투 허용 배수. 대변형 시 1.15~1.25 권장.'],
    ],
    col_widths=[1.8, 1.3, 1.5, 1.8, 9.1]
)
doc.add_paragraph()

add_heading('SOFT 선택 상세', 3)
add_table(
    ['SOFT', '알고리즘', '권장 케이스', '비용'],
    [
        ['0 (기본)', '표준 penalty (bulk modulus 기반)', '동종 재료 금속 접촉', '낮음'],
        ['1 (Lv 6)', 'Soft constraint (노드 질량 기반)', '이종 재료: 고무-금속, 폼-강체', '중간'],
        ['2 (Lv 8)', 'Pinball segment-based', '대변형, 대슬라이딩, 복잡한 접촉', '높음'],
    ],
    col_widths=[2.0, 5.5, 5.5, 2.5]
)
doc.add_paragraph()

# ── 4.7 체적 점성 ─────────────────────────────────────────
add_heading('4.7 체적 점성  (*CONTROL_BULK_VISCOSITY)', 2)
add_table(
    ['옵션', '기본', 'Lv7 설정', 'Lv12 설정', '설명'],
    [
        ['Q1', '1.5', '1.5 (강제)', '2.0',
         '2차 체적 점성. 충격파 압력 스파이크 분산. 고무=1.5, 충격파/폭발=2.0.'],
        ['Q2', '0.06', '0.06 (강제)', '0.10',
         '1차 체적 점성. 팽창파 저속 진동 감쇠. 범용=0.06, 충격파=0.10.'],
    ],
    col_widths=[1.5, 1.5, 2.5, 2.5, 7.5]
)
add_note('Lv7에서 Q1/Q2를 강제 override 합니다. 이미 다른 값이 있어도 1.5/0.06으로 변경되며 콘솔에 notice가 출력됩니다.')
doc.add_paragraph()

# ── 4.8 에너지 추적 ───────────────────────────────────────
add_heading('4.8 에너지 추적  (*CONTROL_ENERGY)  ── Lv 1에서 모두 활성화', 2)
add_table(
    ['옵션', '기본', '설정값', '확인 지표', '해석'],
    [
        ['HGEN', '0', '2', 'glstat: hourglass energy', 'HG energy > 전체의 10% → IHQ 업그레이드(Lv4/9) 필요'],
        ['SLNTEN', '0', '2', 'glstat: sliding energy', '접촉 에너지 급증 → 접촉 설정 문제(Lv6/8)'],
        ['RWEN', '0', '2', 'glstat: external work', '에너지 보존 확인용'],
        ['RYLEN', '0', '2', 'glstat: damping energy', '감쇠 에너지 과다 → 불필요한 CONTROL_DAMPING 확인'],
    ],
    col_widths=[1.8, 1.3, 1.8, 3.8, 6.8]
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 5. 진단 워크플로우
# ════════════════════════════════════════════════════════════════
add_heading('5. 발산 유형별 진단 워크플로우', 1)

add_table(
    ['증상 (glstat/d3plot)', '원인 추정', '권장 레벨'],
    [
        ['시뮬레이션 초반부터 에너지 폭주', 'TSSFAC 너무 공격적 또는 초기 침투', 'Lv3 + Lv6(IGNORE=1)'],
        ['체커보드 변형 패턴 (d3plot)', '호글래스 모드 발생', 'Lv4 (IHQ=4), Lv9 (IHQ=6)'],
        ['HG energy > 전체 에너지 10%', '호글래스 억제 불충분', 'Lv4 → Lv9로 upgrade'],
        ['접촉면에서 진동/튀김(chattering)', '이종 재료 penalty 불균형', 'Lv6 (SOFT=1) → Lv8 (SOFT=2)'],
        ['초기 몇 step에서 에너지 피크', '초기 침투 (initial penetration)', 'Lv6 ISLCHK=2 확인 후 Lv8 IGNORE=1'],
        ['dt가 주기적으로 0에 가까워짐', '요소 심각 변형 (negative Jacobian 임박)', 'Lv10 TSSFAC=0.60'],
        ['Negative Jacobian으로 종료', '요소가 반전됨', 'Lv11 ERODE=11 (물리 변경 주의)'],
        ['쉘 요소만 불안정', '쉘 워핑 또는 소성 수렴 불량', 'Lv5 BWC=1 + MITER=2'],
        ['충격파 압력 진동', '체적 점성 부족', 'Lv12 Q1=2.0/Q2=0.10'],
        ['대슬라이딩 접촉 손실', 'SOFT=0/1의 검색 범위 부족', 'Lv8 SOFT=2 + DEPTH=5'],
    ],
    col_widths=[5.5, 4.5, 5.5]
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 6. Full Manual 예시
# ════════════════════════════════════════════════════════════════
add_heading('6. Full Manual Mode 전체 옵션 예시', 1)
add_para('level 키를 생략하고 모든 옵션을 직접 지정하는 방식. full_manual.yaml 참조.', italic=True)
doc.add_paragraph()

add_table(
    ['옵션 키', '타입', '기본값', '설명'],
    [
        ['model', 'string', '-', '입력 K-file 경로 (필수)'],
        ['output', 'string', '-', '출력 K-file 경로 (필수)'],
        ['stabilize', 'string', '-', '"explicit" 고정'],
        ['level', 'int', '없음', '1~12. 생략 시 수동 모드.'],
        ['confirm_erosion', 'bool', 'false', 'true: ERODE 적용 시 interactive 확인 생략'],
        ['tssfac', 'float', '0.90', 'Courant 안전계수'],
        ['erode', 'int', '0', '요소 침식 플래그 (1/10/11/111)'],
        ['ihq', 'int', '1', '호글래스 타입 (1~6)'],
        ['qh', 'float', '0.10', '호글래스 계수'],
        ['osu', 'int', '0', '객관 응력 업데이트 (0/1)'],
        ['inn', 'int', '1', '불변량 노드 번호 (1/2/4)'],
        ['esort_solid', 'int', '0', 'CONTROL_SOLID ESORT (0/1)'],
        ['esort_shell', 'int', '0', 'CONTROL_SHELL ESORT (0/1)'],
        ['bwc', 'int', '2', '쉘 워핑 강성 (1/2)'],
        ['miter', 'int', '1', '쉘 소성 반복 (1/2/3)'],
        ['irnxx', 'int', '-1', '쉘 노달 파이버 (-1/-2)'],
        ['wrpang', 'float', '20', '쉘 워핑 경고 각도 (도)'],
        ['orien', 'int', '1', '접촉 법선 재정렬 (1/2)'],
        ['shlthk', 'int', '0', '접촉 쉘 두께 포함 (0/1/2)'],
        ['xpene', 'float', '4.0', '최대 침투 배수'],
        ['enmass', 'int', '0', '침식 노드 접촉 유지 (0/1/2)'],
        ['nsbcs', 'int', '자동', '버킷 검색 주기 (cycle)'],
        ['islchk', 'int', '1', '초기 침투 검사 (1/2)'],
        ['soft', 'int', '0', 'per-contact penalty 타입 (0/1/2)'],
        ['sbopt', 'int', '0', 'per-contact 세그먼트 검색 (0/2/3)'],
        ['depth', 'int', '2', 'per-contact 검색 깊이'],
        ['ignore', 'int', '0', 'per-contact 초기 침투 처리 (0/1/2)'],
        ['maxpar', 'float', '1.025', 'per-contact SOFT=2 침투 배수'],
        ['bulk_q1', 'float', '1.5', '2차 체적 점성 계수'],
        ['bulk_q2', 'float', '0.06', '1차 체적 점성 계수'],
        ['hgen', 'int', '0', '호글래스 에너지 추적 (0/2)'],
        ['rwen', 'int', '0', '리블릿 에너지 추적 (0/2)'],
        ['slnten', 'int', '0', '슬라이딩 에너지 추적 (0/2)'],
        ['rylen', 'int', '0', '레일리 에너지 추적 (0/2)'],
    ],
    col_widths=[2.5, 1.5, 2.0, 9.5]
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 7. optimize rubber와의 관계
# ════════════════════════════════════════════════════════════════
add_heading('7. optimize rubber 명령과의 관계', 1)
add_para(
    'stabilize explicit과 optimize rubber는 독립적이며 함께 사용 가능합니다. '
    '일반적으로 matswap → optimize → stabilize 순서로 적용합니다.'
)
doc.add_paragraph()

add_table(
    ['구분', 'optimize rubber', 'stabilize explicit'],
    [
        ['목적', '고무 재료 특화 최적화', '발산/불안정 해석 일반 안정화'],
        ['CONTROL_ACCURACY INN', 'INN=4 강제', 'INN=2(Lv2)→4(Lv3)'],
        ['CONTROL_ENERGY', '전체 활성화', '전체 활성화 (Lv1)'],
        ['CONTROL_TIMESTEP TSSFAC', '0.67 강제', 'Lv7에서 0.67'],
        ['CONTACT SOFT/SBOPT', '대상 PID 한정', '모든 contact 전체 적용'],
        ['IHQ 호글래스', '없음', 'Lv4: IHQ=4, Lv9: IHQ=6'],
        ['ERODE 침식', '없음', 'Lv11 (interactive 확인)'],
        ['Q1/Q2 체적 점성', '경고만', 'Lv7: 1.5/0.06, Lv12: 2.0/0.10'],
    ],
    col_widths=[4.5, 5.5, 5.5]
)
doc.add_paragraph()

add_code('# 권장 적용 순서 예시')
add_code('# Step 1: matswap (재료 교환)')
add_code('KooRemapper matswap matswap.yaml')
add_code('')
add_code('# Step 2: optimize rubber (고무 특화 최적화)')
add_code('KooRemapper optimize optimize.yaml')
add_code('')
add_code('# Step 3: stabilize explicit (발산 방지 안정화)')
add_code('KooRemapper stabilize stabilize.yaml')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 8. 예제 파일 목록
# ════════════════════════════════════════════════════════════════
add_heading('8. 예제 파일 목록 (examples/explicit/)', 1)
add_table(
    ['파일', '설명'],
    [
        ['level01.yaml', 'Lv 1: 에너지 진단 (비용 없음, 첫 번째 단계)'],
        ['level02.yaml', 'Lv 2: 정확도 기반 (OSU=1, INN=2, ESORT)'],
        ['level03.yaml', 'Lv 3: 시간 스텝 1 (TSSFAC=0.80, INN=4)'],
        ['level04.yaml', 'Lv 4: 호글래스 1 (IHQ=4 FB stiffness)'],
        ['level05.yaml', 'Lv 5: 쉘 안정화 (BWC=1, MITER=2, IRNXX=-2) / 쉘 없으면 skip'],
        ['level06.yaml', 'Lv 6: 접촉 1단계 (SOFT=1, SBOPT=2, DEPTH=3)'],
        ['level07.yaml', 'Lv 7: 시간 2 + 체적점성 (TSSFAC=0.67, Q1/Q2 설정)'],
        ['level08.yaml', 'Lv 8: 접촉 2단계 (SOFT=2 pinball, DEPTH=5, NSBCS=5)'],
        ['level09.yaml', 'Lv 9: 호글래스 최강 (IHQ=6 BB, QH=1.0)'],
        ['level10.yaml', 'Lv10: 시간 3 (TSSFAC=0.60, NSBCS=2)'],
        ['level11.yaml', 'Lv11: 침식 활성화 (ERODE=11, ENMASS=2) ⚠물리 변경'],
        ['level12.yaml', 'Lv12: 최대 보수 (TSSFAC=0.50, Q1=2.0/Q2=0.10)'],
        ['full_manual.yaml', '전체 옵션 직접 지정 모드. 모든 키 설명 주석 포함.'],
        ['explicit_stabilization_guide.docx', '이 문서'],
    ],
    col_widths=[4.5, 11.0]
)
doc.add_paragraph()

# 저장
out = 'd:/KooRemapper/examples/explicit/explicit_stabilization_guide.docx'
doc.save(out)
print(f'Generated: {out}')
