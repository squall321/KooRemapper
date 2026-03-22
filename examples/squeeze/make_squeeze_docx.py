#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_squeeze_docx.py
KooRemapper Squeeze 커맨드 가이드 DOCX 생성 스크립트

실행 방법:
    python make_squeeze_docx.py
    (examples/squeeze/ 디렉토리에서 실행)

출력: squeeze_guide.docx

요구사항:
    pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# 색상 상수
# ─────────────────────────────────────────────
COLOR_H1        = RGBColor(0x1F, 0x38, 0x64)   # Navy Blue
COLOR_H2        = RGBColor(0x17, 0x37, 0x5E)   # Dark Teal
COLOR_H3        = RGBColor(0x2E, 0x2E, 0x2E)   # Dark Gray
COLOR_TABLE_HDR = RGBColor(0xD9, 0xD9, 0xD9)   # Light Gray header
COLOR_CODE_BG   = RGBColor(0xF0, 0xF0, 0xF0)   # Code background
COLOR_NOTE_BG   = RGBColor(0xFF, 0xF3, 0xCD)   # Note background (yellow tint)
COLOR_WARN_BG   = RGBColor(0xFF, 0xE0, 0xE0)   # Warning background (red tint)

FONT_BODY = "Malgun Gothic"
FONT_CODE = "Courier New"
FONT_SIZE_BODY = Pt(10)
FONT_SIZE_CODE = Pt(9)


# ─────────────────────────────────────────────
# 헬퍼 함수
# ─────────────────────────────────────────────

def rgb_to_hex(rgb_color):
    r, g, b = rgb_color[0], rgb_color[1], rgb_color[2]
    return f"{r:02X}{g:02X}{b:02X}"


def set_cell_background(cell, rgb_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = rgb_to_hex(rgb_color)
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_paragraph_shading(paragraph, rgb_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_color = rgb_to_hex(rgb_color)
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), hex_color)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_BODY
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_H1
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_BODY
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_H2
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_BODY
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_H3
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = FONT_SIZE_BODY
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = FONT_SIZE_BODY
    return p


def add_code_block(doc, code_text):
    lines = code_text.split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        set_paragraph_shading(p, COLOR_CODE_BG)
        run = p.add_run(line if line else ' ')
        run.font.name = FONT_CODE
        run.font.size = FONT_SIZE_CODE
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(6)


def add_note(doc, text, bg=None):
    """노트/경고 박스"""
    if bg is None:
        bg = COLOR_NOTE_BG
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    set_paragraph_shading(p, bg)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = FONT_SIZE_BODY


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, COLOR_TABLE_HDR)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(9)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = FONT_BODY
            run.font.size = Pt(9)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def add_spacer(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)


# ─────────────────────────────────────────────
# 문서 생성 메인
# ─────────────────────────────────────────────

def create_squeeze_guide():
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    # ──────────────────────────────────────────
    # 표지
    # ──────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("Squeeze 커맨드 가이드")
    title_run.bold = True
    title_run.font.name = FONT_BODY
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = COLOR_H1

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(4)
    sub_run = sub_p.add_run("KooRemapper v1.3.2 기준")
    sub_run.font.name = FONT_BODY
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = COLOR_H3

    sub2_p = doc.add_paragraph()
    sub2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2_p.paragraph_format.space_after = Pt(20)
    sub2_run = sub2_p.add_run("간섭끼워맞춤 · 초기 변형 · 팽윤 · Dynamic Relaxation 통합 가이드")
    sub2_run.font.name = FONT_BODY
    sub2_run.font.size = Pt(10)
    sub2_run.font.color.rgb = COLOR_H3

    doc.add_page_break()

    # ──────────────────────────────────────────
    # 1. 개요
    # ──────────────────────────────────────────
    add_heading1(doc, "1. 개요")

    add_body(doc,
        "squeeze 커맨드는 간섭끼워맞춤(interference fit), 팽윤(swelling), 초기 변형 등으로 인해 "
        "파트가 변형된 상태를 LS-DYNA에 초기 조건으로 전달하는 도구다.")

    add_spacer(doc)
    add_heading2(doc, "1.1 핵심 아이디어")

    add_body(doc, "squeeze는 두 가지 정보를 동시에 생성하여 LS-DYNA에 전달한다.")
    add_bullet(doc, "압착·팽창된 기하 형상 — 노드 좌표를 스케일하여 이미 변형된 형상으로 표현")
    add_bullet(doc, "초기 조건 — 그 변형에 대응하는 응력 또는 변형률을 dynain 파일로 기록")
    add_bullet(doc, "LS-DYNA가 해석 시 두 정보를 합산하여 올바른 평형 상태에서 시작")

    add_spacer(doc)
    add_heading2(doc, "1.2 실행 방법")

    add_code_block(doc, "KooRemapper squeeze <mesh.k> <config.yaml> <output_prefix>")

    add_table(doc,
        ["인자", "설명"],
        [
            ["mesh.k",        "원본 기하 K-파일"],
            ["config.yaml",   "설정 파일 (모드·파트·재료·DR 옵션)"],
            ["output_prefix", "출력 파일명 접두어"],
        ],
        col_widths=[4.0, 12.0]
    )

    add_heading2(doc, "1.3 출력 파일")

    add_table(doc,
        ["파일", "내용"],
        [
            ["<prefix>.k",      "압착된 메시 + *INCLUDE 연결 (dynain, 열팽창 카드, DR 카드 포함)"],
            ["<prefix>.dynain", "초기 응력(*INITIAL_STRESS_SOLID) 또는 변형률(*INITIAL_STRAIN_SOLID)"],
        ],
        col_widths=[4.5, 11.5]
    )

    # ──────────────────────────────────────────
    # 2. 초기 조건 방식 3가지
    # ──────────────────────────────────────────
    add_heading1(doc, "2. 초기 조건 방식 3가지")

    add_table(doc,
        ["방식", "YAML 키", "재료 필요 여부", "노드 이동", "LS-DYNA 카드"],
        [
            ["응력 (기본)",       "eps_x/y/z",                  "필수",         "O", "*INITIAL_STRESS_SOLID"],
            ["변형률",            "eps_x/y/z + strain_mode: true", "불필요",    "O", "*INITIAL_STRAIN_SOLID"],
            ["열팽창 (swelling)", "swelling",                    "K-파일 *MAT 필수", "X",
             "*MAT_ADD_THERMAL_EXPANSION\n+ *INITIAL_TEMPERATURE_NODE"],
        ],
        col_widths=[2.8, 4.5, 3.2, 1.5, 4.0]
    )

    # ──────────────────────────────────────────
    # 3. 응력 모드 (기본)
    # ──────────────────────────────────────────
    add_heading1(doc, "3. 응력 모드 (기본)")

    add_body(doc,
        "eps_x/y/z를 Hooke's law로 응력으로 변환하여 *INITIAL_STRESS_SOLID로 기록한다. "
        "재료(E, nu)가 반드시 필요하다.")

    add_heading2(doc, "3.1 동작 흐름")

    add_code_block(doc,
        "eps_x/y/z 입력\n"
        "      │\n"
        "      ▼  노드 좌표 스케일  (center + (pos - center) × (1 + eps))\n"
        "   압착된 기하 형상\n"
        "      │\n"
        "      ▼  역변형률 (-eps) → Hooke's law\n"
        "   *INITIAL_STRESS_SOLID\n"
        "      │\n"
        "      ▼  LS-DYNA 해석\n"
        "   응력과 기하가 평형 상태로 수렴")

    add_heading2(doc, "3.2 재료 지정 방법")

    add_heading3(doc, "A. YAML에 재료 직접 지정 (ex01)")

    add_body(doc, "K-파일에 *MAT 카드가 없을 때. 모든 squeeze 파트에 동일 재료 적용.")

    add_code_block(doc,
        "parts:\n"
        "  - pid: 1\n"
        "    eps_x: -0.01    # x방향 1% 압축\n"
        "    eps_y: -0.01    # y방향 1% 압축\n"
        "    eps_z:  0.0\n"
        "\n"
        "material:\n"
        "  E: 210000.0    # MPa\n"
        "  nu: 0.3")

    add_heading3(doc, "B. K-파일 재료 사용 (ex02)")

    add_body(doc,
        "K-파일에 *MAT_ELASTIC 등이 있을 때. material: 섹션 생략 가능. "
        "파트별로 각자의 재료 자동 사용.")

    add_code_block(doc,
        "parts:\n"
        "  - pid: 1\n"
        "    eps_x: -0.015\n"
        "    eps_y: -0.015\n"
        "    eps_z:  0.0\n"
        "  - pid: 2\n"
        "    eps_x: -0.010\n"
        "    eps_y:  0.0\n"
        "    eps_z:  0.0\n"
        "\n"
        "# material: 섹션 없음 → K-파일 *MAT 자동 참조 (파트별 개별 적용)")

    add_note(doc,
        "우선순위: YAML material:  >  K-파일 *MAT. "
        "YAML에 지정하면 K-파일 재료를 완전히 덮어씀.")

    add_heading2(doc, "3.3 생성되는 *INITIAL_STRESS_SOLID 구조")

    add_code_block(doc,
        "*INITIAL_STRESS_SOLID\n"
        "$#    eid    nint   nhisv   large     ics   ncomp\n"
        "       101       1       0       0       0       6\n"
        "$#      sx        sy        sz       sxy       syz       sxz\n"
        " -2.308e+03  -2.308e+03   0.000e+00  0.000e+00  0.000e+00  0.000e+00")

    # ──────────────────────────────────────────
    # 4. 변형률 모드
    # ──────────────────────────────────────────
    add_heading1(doc, "4. 변형률 모드 (strain_mode: true)")

    add_body(doc,
        "eps 값을 응력으로 변환하지 않고 *INITIAL_STRAIN_SOLID에 직접 기록한다. "
        "재료(E, nu)가 없어도 동작한다.")

    add_heading2(doc, "4.1 사용하는 경우")

    add_bullet(doc, "재료의 E, nu를 모를 때")
    add_bullet(doc, "재료가 비선형(소성, 초탄성)이어서 Hooke's law를 쓸 수 없을 때")
    add_bullet(doc, "LS-DYNA 재료 모델이 직접 응력을 계산하게 하고 싶을 때")

    add_heading2(doc, "4.2 동작 흐름")

    add_code_block(doc,
        "eps_x/y/z 입력\n"
        "      │\n"
        "      ▼  노드 좌표 스케일 (응력 모드와 동일)\n"
        "   압착된 기하 형상\n"
        "      │\n"
        "      ▼  역변형률 (-eps) 그대로 기록 (Hooke's law 없음)\n"
        "   *INITIAL_STRAIN_SOLID\n"
        "      │\n"
        "      ▼  LS-DYNA 해석 (재료 모델이 응력 자동 계산)\n"
        "   응력과 기하가 평형 상태로 수렴")

    add_heading2(doc, "4.3 YAML 예제 (ex03)")

    add_code_block(doc,
        "strain_mode: true   # 최상위 키\n"
        "\n"
        "parts:\n"
        "  - pid: 1\n"
        "    eps_x: -0.01\n"
        "    eps_y: -0.01\n"
        "    eps_z:  0.0\n"
        "\n"
        "# material: 섹션 불필요")

    add_heading2(doc, "4.4 재료 지정과의 관계")

    add_table(doc,
        ["재료 상태", "strain_mode 동작"],
        [
            ["YAML material: 있음",  "dynain 생성에는 사용 안 함 (무시)"],
            ["K-파일 *MAT 있음",     "dynain 생성에는 사용 안 함 (LS-DYNA 해석에는 필요)"],
            ["재료 완전히 없음",     "정상 동작 — *INITIAL_STRAIN_SOLID만 생성"],
        ],
        col_widths=[5.0, 11.0]
    )

    add_note(doc,
        "주의: LS-DYNA 해석 자체에는 재료 모델이 필요하다. "
        "strain_mode는 KooRemapper가 응력 계산을 위해 재료를 요구하지 않는다는 의미다. "
        "K-파일에는 *MAT 카드가 있어야 LS-DYNA 해석이 가능하다.",
        bg=COLOR_WARN_BG)

    add_heading2(doc, "4.5 생성되는 *INITIAL_STRAIN_SOLID 구조")

    add_code_block(doc,
        "*INITIAL_STRAIN_SOLID\n"
        "$#    eid    nint   nhisv   large\n"
        "       101       1       0       0\n"
        "$#       eps11        eps22        eps33        eps12        eps23        eps13\n"
        "  -1.0000e-02  -1.0000e-02   0.0000e+00   0.0000e+00   0.0000e+00   0.0000e+00")

    # ──────────────────────────────────────────
    # 5. 열팽창 모드
    # ──────────────────────────────────────────
    add_heading1(doc, "5. 열팽창 모드 (swelling)")

    add_body(doc,
        "재료가 균일하게 등방 팽창하는 경우에 사용한다. "
        "노드를 이동시키지 않고 *MAT_ADD_THERMAL_EXPANSION으로 팽창을 표현한다.")

    add_heading2(doc, "5.1 사용하는 경우")

    add_bullet(doc, "배터리 전극, 고분자, 젤, 흡수재 등 등방 팽창 재료")
    add_bullet(doc, "팽창 후 형상이 아닌 팽창 전 형상을 기준으로 해석할 때")

    add_heading2(doc, "5.2 동작 흐름")

    add_code_block(doc,
        "swelling: 0.03 (3% 팽윤)\n"
        "      │\n"
        "      ▼  노드 이동 없음 (기하 형상 유지)\n"
        "   *MAT_ADD_THERMAL_EXPANSION  (CTE = swelling값, MID = *PART의 MID)\n"
        "   *INITIAL_TEMPERATURE_NODE  (T = 1.0, 모든 파트 노드)\n"
        "      │\n"
        "      ▼  LS-DYNA 해석\n"
        "   ΔT=1 × CTE=0.03 → 3% 열팽창 자동 적용")

    add_heading2(doc, "5.3 YAML 예제 (ex04)")

    add_code_block(doc,
        "parts:\n"
        "  - pid: 1\n"
        "    swelling: 0.03    # 3% 등방 팽윤\n"
        "\n"
        "  - pid: 2\n"
        "    swelling: 0.05    # 5% 팽윤\n"
        "\n"
        "# K-파일에 *MAT_* 카드 필수")

    add_heading2(doc, "5.4 생성되는 카드 예시 (PID 1, MID 1)")

    add_code_block(doc,
        "*MAT_ADD_THERMAL_EXPANSION\n"
        "$#     mid      lcid     mult\n"
        "         1         0   0.030\n"
        "*INITIAL_TEMPERATURE_NODE\n"
        "$#     nid      temp       loc\n"
        "      1001       1.0         0\n"
        "      1002       1.0         0\n"
        "      ...")

    add_heading2(doc, "5.5 제약사항")

    add_bullet(doc, "한 파트에 swelling과 eps_x/y/z 동시 사용 불가")
    add_bullet(doc, "K-파일에 해당 파트의 *MAT 카드 필수 (MID를 *PART에서 읽어 연결)")
    add_bullet(doc, "*MAT 없으면 PID를 MID로 대체 (경고 출력)")
    add_bullet(doc, "dynain 파일에 포함되지 않음 — 카드가 직접 .k 파일에 삽입")

    # ──────────────────────────────────────────
    # 6. 재료 지정 분기 요약
    # ──────────────────────────────────────────
    add_heading1(doc, "6. 재료 지정 분기 요약")

    add_code_block(doc,
        "squeeze 실행\n"
        "      │\n"
        "      ├─ swelling 파트 ──────────────────────── K-파일 *MAT 연결\n"
        "      │                                         (없으면 PID → MID 대체)\n"
        "      │\n"
        "      └─ eps 파트\n"
        "            │\n"
        "            ├─ strain_mode: true ────────────── 재료 불필요\n"
        "            │                                   *INITIAL_STRAIN_SOLID\n"
        "            │\n"
        "            └─ strain_mode: false (기본)\n"
        "                  │\n"
        "                  ├─ YAML material: 있음 ──────  전 파트 동일 재료\n"
        "                  │                              *INITIAL_STRESS_SOLID\n"
        "                  │\n"
        "                  ├─ K-파일 *MAT 있음 ─────────  파트별 재료 자동\n"
        "                  │                              *INITIAL_STRESS_SOLID\n"
        "                  │\n"
        "                  └─ 재료 없음 ───────────────── ERROR\n"
        "                                                  → strain_mode: true 사용 권장")

    # ──────────────────────────────────────────
    # 7. Dynamic Relaxation 연동
    # ──────────────────────────────────────────
    add_heading1(doc, "7. Dynamic Relaxation 연동 (relax: 섹션)")

    add_body(doc,
        "squeeze 출력 파일에 DR 키워드를 자동 삽입한다. "
        "별도로 KooRemapper relax 커맨드를 실행하지 않아도 된다. "
        "relax: 섹션 진입만으로 DR이 활성화된다.")

    add_heading2(doc, "7.1 YAML 옵션")

    add_code_block(doc,
        "relax:\n"
        "  level: 2          # 1(빠름) ~ 5(최대 보수적)\n"
        "  mode: explicit    # explicit(IDRFLG=1) | implicit(IDRFLG=5)\n"
        "  drterm: 0.0       # DR 종료 시간 (0 = 수렴 판정까지)\n"
        "  endtime: 1.0      # *CONTROL_TERMINATION 삽입 (생략 가능)\n"
        "  d3drlf: true      # *DATABASE_BINARY_D3DRLF 출력\n"
        "  # 개별 오버라이드 (선택)\n"
        "  nrcyck: 250       # 수렴 체크 간격\n"
        "  drtol: 0.001      # 수렴 tolerance\n"
        "  drfctr: 0.995     # 속도 감쇠 계수\n"
        "  tssfdr: 0.90      # DR 중 timestep 스케일\n"
        "  irelal: 0         # 자동 제어 (0=off, 1=on)\n"
        "  edttl: 0.04       # 자동 제어 tolerance")

    add_heading2(doc, "7.2 레벨 프리셋")

    add_table(doc,
        ["레벨", "이름", "NRCYCK", "DRTOL", "DRFCTR", "TSSFDR", "IRELAL"],
        [
            ["1", "빠름",   "500", "0.010",  "0.990", "0.95", "0"],
            ["2", "표준",   "250", "0.001",  "0.995", "0.90", "0"],
            ["3", "안정",   "100", "0.001",  "0.998", "0.80", "0"],
            ["4", "보수",    "50", "1e-4",   "0.999", "0.67", "1"],
            ["5", "최대",    "25", "1e-5",   "0.999", "0.50", "1"],
        ],
        col_widths=[1.5, 2.0, 2.2, 2.2, 2.2, 2.2, 2.0]
    )

    add_heading2(doc, "7.3 삽입되는 DR 카드 (explicit, level=2)")

    add_code_block(doc,
        "*CONTROL_DYNAMIC_RELAXATION\n"
        "$#  nrcyck    drtol   drfctr   drterm   tssfdr   irelal    edttl    idrflg\n"
        "       250  0.001000  0.995000       0.0  0.900000        0  0.040000        1\n"
        "*DATABASE_BINARY_D3DRLF\n"
        "$#      dt      lcid\n"
        " 1.000e-03         0")

    # ──────────────────────────────────────────
    # 8. 전체 YAML 옵션 레퍼런스
    # ──────────────────────────────────────────
    add_heading1(doc, "8. 전체 YAML 옵션 레퍼런스")

    add_table(doc,
        ["키", "기본값", "설명"],
        [
            ["strain_mode",   "false",     "true: *INITIAL_STRAIN_SOLID (재료 불필요)\nfalse: *INITIAL_STRESS_SOLID (재료 필수)"],
            ["parts[].pid",   "(필수)",    "파트 ID"],
            ["parts[].eps_x", "0.0",       "x방향 공학 변형률 (음수=압축, 양수=인장)"],
            ["parts[].eps_y", "0.0",       "y방향 공학 변형률"],
            ["parts[].eps_z", "0.0",       "z방향 공학 변형률"],
            ["parts[].swelling", "0.0",    "등방 팽윤율 (eps와 동시 사용 불가)"],
            ["material.E",    "(없음)",    "Young's modulus [MPa] (생략 시 K-파일 *MAT 사용)"],
            ["material.nu",   "(없음)",    "Poisson's ratio"],
            ["relax.level",   "2",         "DR 강도 레벨 (1~5)"],
            ["relax.mode",    "explicit",  "explicit(IDRFLG=1) | implicit(IDRFLG=5)"],
            ["relax.drterm",  "0.0",       "DR 종료 시간 (0=수렴 판정)"],
            ["relax.endtime", "(없음)",    "*CONTROL_TERMINATION 삽입 (생략 시 미삽입)"],
            ["relax.d3drlf",  "true",      "*DATABASE_BINARY_D3DRLF 출력 여부"],
            ["relax.nrcyck",  "(레벨값)",  "수렴 체크 간격 오버라이드"],
            ["relax.drtol",   "(레벨값)",  "수렴 tolerance 오버라이드"],
            ["relax.drfctr",  "(레벨값)",  "속도 감쇠 계수 오버라이드"],
            ["relax.tssfdr",  "(레벨값)",  "DR 중 timestep 스케일 오버라이드"],
            ["relax.irelal",  "(레벨값)",  "자동 제어 오버라이드 (0=off, 1=on)"],
            ["relax.edttl",   "(레벨값)",  "자동 제어 tolerance 오버라이드"],
        ],
        col_widths=[3.5, 2.5, 10.0]
    )

    # ──────────────────────────────────────────
    # 9. 예제 파일 목록
    # ──────────────────────────────────────────
    add_heading1(doc, "9. 예제 파일 목록")

    add_table(doc,
        ["파일", "모드", "재료", "DR"],
        [
            ["ex01_stress_yaml_material.yaml",  "응력",          "YAML 직접 지정",   "없음"],
            ["ex02_stress_kfile_material.yaml",  "응력",          "K-파일 자동",      "없음"],
            ["ex03_strain_no_material.yaml",     "변형률",        "불필요",           "없음"],
            ["ex04_swelling.yaml",               "열팽창",        "K-파일 필수",      "없음"],
            ["ex05_mixed_with_dr.yaml",          "응력+열팽창",   "YAML+K-파일",      "포함"],
        ],
        col_widths=[5.5, 2.5, 3.5, 2.0]
    )

    # ──────────────────────────────────────────
    # 10. 워크플로우별 시나리오
    # ──────────────────────────────────────────
    add_heading1(doc, "10. 워크플로우별 시나리오")

    add_heading2(doc, "시나리오 A: 볼베어링 간섭끼워맞춤")

    add_body(doc, "재료를 알고 있음 → 응력 모드 + DR")

    add_code_block(doc,
        "KooRemapper squeeze bearing.k bearing_squeeze.yaml bearing_out\n"
        "# → bearing_out.k (INITIAL_STRESS_SOLID + DR 카드)\n"
        "# → bearing_out.dynain\n"
        "# LS-DYNA로 bearing_out.k 실행 → DR 평형 후 접촉 해석")

    add_heading2(doc, "시나리오 B: 배터리 셀 팽윤")

    add_body(doc, "재료는 K-파일에 있고 등방 팽창 → 열팽창 모드")

    add_code_block(doc,
        "KooRemapper squeeze battery.k battery_swell.yaml battery_out\n"
        "# → battery_out.k (MAT_ADD_THERMAL_EXPANSION + INITIAL_TEMPERATURE_NODE)\n"
        "# LS-DYNA 해석 시 팽윤 자동 적용")

    add_heading2(doc, "시나리오 C: 비선형 재료 초기 변형")

    add_body(doc, "재료 모델이 복잡하거나 없음 → 변형률 모드")

    add_code_block(doc,
        "KooRemapper squeeze part.k squeeze_strain.yaml part_out\n"
        "# → part_out.k (INITIAL_STRAIN_SOLID)\n"
        "# LS-DYNA 재료 모델이 직접 응력 결정")

    # ──────────────────────────────────────────
    # 11. 자주 하는 실수
    # ──────────────────────────────────────────
    add_heading1(doc, "11. 자주 하는 실수")

    add_table(doc,
        ["증상", "원인", "해결"],
        [
            ["No material specified 에러",
             "응력 모드인데 재료 없음",
             "YAML material: 추가 또는 strain_mode: true"],
            ["dynain 요소 수 = 0",
             "K-파일에 해당 PID 없음",
             "K-파일과 YAML의 PID 일치 확인"],
            ["swelling 카드 생성 안 됨",
             "*PART 카드 없어 MID 탐색 실패",
             "K-파일에 *PART 카드 추가"],
            ["LS-DYNA가 초기 응력 무시",
             "dynain이 *INCLUDE 되지 않음",
             ".k 파일에 *INCLUDE <prefix>.dynain 확인"],
            ["DR이 수렴 안 됨",
             "레벨 낮거나 변형이 큼",
             "relax.level 값 높이기 (3~4)"],
        ],
        col_widths=[4.5, 5.0, 6.5]
    )

    # ──────────────────────────────────────────
    # 저장
    # ──────────────────────────────────────────
    out_path = "squeeze_guide.docx"
    doc.save(out_path)
    print(f"저장 완료: {out_path}")


if __name__ == "__main__":
    create_squeeze_guide()
