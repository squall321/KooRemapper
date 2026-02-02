#!/usr/bin/env python3
"""
KooRemapper 이론 문서 DOCX 생성기
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """테이블 셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_heading_with_number(doc, text, level):
    """번호가 있는 제목 추가"""
    doc.add_heading(text, level=level)

def add_equation(doc, equation_text):
    """수식 추가 (일반 텍스트로)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(equation_text)
    run.italic = True
    run.font.name = 'Cambria Math'
    return p

def create_document():
    doc = Document()

    # 문서 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)

    # 제목 페이지
    title = doc.add_heading('KooRemapper 이론 문서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('메시 매핑 및 응력/변형률 계산 알고리즘')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Version 1.0\n').bold = True
    info.add_run('Date: 2026-01-20')

    doc.add_page_break()

    # 목차
    doc.add_heading('목차', level=1)
    toc_items = [
        '1. 개요',
        '2. 메시 매핑 알고리즘',
        '   2.1 트랜스파이닛 보간법',
        '   2.2 호장 길이 기반 매개변수화',
        '   2.3 Coons 패치',
        '   2.4 구조화 그리드 인덱싱',
        '3. 변형 해석 알고리즘',
        '   3.1 변형 구배 텐서',
        '   3.2 변형률 텐서',
        '   3.3 응력 텐서',
        '4. 수치 해석 기법',
        '   4.1 가우스 적분',
        '   4.2 형상 함수',
        '   4.3 야코비안 행렬',
        '5. 재료 모델',
        '6. 참고 문헌',
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ============== 1. 개요 ==============
    doc.add_heading('1. 개요', level=1)

    doc.add_paragraph(
        'KooRemapper는 LS-DYNA 유한요소 해석을 위한 메시 매핑 및 프리스트레스 계산 도구입니다. '
        '본 문서는 프로그램에서 사용된 핵심 수학적 알고리즘과 이론적 배경을 상세히 설명합니다.'
    )

    doc.add_heading('1.1 주요 기능', level=2)
    features = [
        '메시 매핑: 평면(flat) 메시를 곡면(bent) 참조 형상으로 변환',
        '변형률 계산: 기준 형상과 변형 형상 간의 변형률 계산',
        '응력 계산: 재료 물성을 이용한 Cauchy 응력 계산',
        'DYNAIN 출력: LS-DYNA 프리스트레스 입력 파일 생성',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('1.2 좌표계 정의', level=2)
    doc.add_paragraph('프로그램은 다음 좌표계를 사용합니다:')

    coords = [
        ('물리 좌표계 (Physical Coordinates)', '(x, y, z) - 실제 공간 좌표'),
        ('매개변수 좌표계 (Parametric Coordinates)', '(u, v, w) ∈ [0,1]³ - 정규화된 좌표'),
        ('자연 좌표계 (Natural Coordinates)', '(ξ, η, ζ) ∈ [-1,1]³ - 요소 내부 좌표'),
    ]
    for name, desc in coords:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(name + ': ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ============== 2. 메시 매핑 알고리즘 ==============
    doc.add_heading('2. 메시 매핑 알고리즘', level=1)

    # 2.1 트랜스파이닛 보간법
    doc.add_heading('2.1 트랜스파이닛 보간법 (Transfinite Interpolation)', level=2)

    doc.add_heading('2.1.1 이론적 배경', level=3)
    doc.add_paragraph(
        '트랜스파이닛 보간법은 Gordon & Hall (1973)에 의해 개발된 방법으로, '
        '경계 조건을 정확히 만족하면서 내부 영역을 보간합니다.'
    )

    doc.add_heading('2.1.2 3차원 트랜스파이닛 보간', level=3)
    doc.add_paragraph('매개변수 좌표 (u, v, w) ∈ [0,1]³에서 물리 좌표 P(u, v, w)로의 변환:')

    doc.add_paragraph('기본 공식:', style='Intense Quote')
    add_equation(doc, 'P(u,v,w) = Lc + Ld + Le - Lcd - Lce - Lde + B')

    doc.add_paragraph('여기서:')
    terms = [
        'Lc: u-방향 선형 보간',
        'Ld: v-방향 선형 보간',
        'Le: w-방향 선형 보간',
        'Lcd, Lce, Lde: 이중 보정항',
        'B: 삼중선형 코너 보정항',
    ]
    for t in terms:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_heading('2.1.3 Edge-Based 보간 (구현된 방식)', level=3)
    doc.add_paragraph('KooRemapper에서는 단순화된 Edge-Based 보간법을 사용합니다:')

    doc.add_paragraph('Step 1: i-축 모서리에서 점 선택', style='Intense Quote')
    code1 = '''p₀₀ = edge[0].interpolate(u)  // j=0, k=0 모서리
p₁₀ = edge[1].interpolate(u)  // j=N, k=0 모서리
p₀₁ = edge[2].interpolate(u)  // j=0, k=P 모서리
p₁₁ = edge[3].interpolate(u)  // j=N, k=P 모서리'''
    p = doc.add_paragraph(code1)
    p.runs[0].font.name = 'Consolas'
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph('Step 2: (v, w) 평면에서 쌍선형 보간', style='Intense Quote')
    add_equation(doc, 'P(u,v,w) = (1-w)·[(1-v)·p₀₀ + v·p₁₀] + w·[(1-v)·p₀₁ + v·p₁₁]')

    doc.add_heading('2.1.4 삼선형 보간 (Trilinear Interpolation)', level=3)
    doc.add_paragraph('가장 단순한 형태로, 8개 코너 노드만 사용:')
    add_equation(doc, 'P(u,v,w) = Σᵢ Nᵢ(u,v,w) · Cᵢ')

    doc.add_paragraph('형상함수:')
    add_equation(doc, 'Nᵢ(u,v,w) = (1/8)(1 ± u)(1 ± v)(1 ± w)')

    # 노드 테이블
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'
    headers = ['노드', 'u', 'v', 'w']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')

    node_data = [
        ('0', '-', '-', '-'),
        ('1', '+', '-', '-'),
        ('2', '+', '+', '-'),
        ('3', '-', '+', '-'),
        ('4', '-', '-', '+'),
        ('5', '+', '-', '+'),
        ('6', '+', '+', '+'),
        ('7', '-', '+', '+'),
    ]
    for i, (n, u, v, w) in enumerate(node_data):
        row = table.rows[i+1]
        row.cells[0].text = n
        row.cells[1].text = u
        row.cells[2].text = v
        row.cells[3].text = w

    doc.add_paragraph()

    # 2.2 호장 길이 기반 매개변수화
    doc.add_heading('2.2 호장 길이 기반 매개변수화 (Arc-length Parameterization)', level=2)

    doc.add_heading('2.2.1 이론', level=3)
    doc.add_paragraph('곡선을 따라 균일한 매개변수 분포를 보장하기 위해 호장 길이를 사용합니다.')

    doc.add_heading('2.2.2 알고리즘', level=3)

    doc.add_paragraph('Step 1: 누적 호장 계산', style='Intense Quote')
    add_equation(doc, 's₀ = 0')
    add_equation(doc, 'sᵢ = sᵢ₋₁ + ‖Pᵢ - Pᵢ₋₁‖')
    add_equation(doc, 'L_total = sₙ')

    doc.add_paragraph('Step 2: 매개변수 t에 대한 위치 계산', style='Intense Quote')
    doc.add_paragraph('목표 호장:')
    add_equation(doc, 's_target = t · L_total')

    doc.add_paragraph('선분 검색: sᵢ ≤ s_target < sᵢ₊₁ 인 i 찾기')

    doc.add_paragraph('국소 매개변수:')
    add_equation(doc, 't_local = (s_target - sᵢ) / (sᵢ₊₁ - sᵢ)')

    doc.add_paragraph('보간 결과:')
    add_equation(doc, 'P(t) = (1 - t_local) · Pᵢ + t_local · Pᵢ₊₁')

    doc.add_heading('2.2.3 장점', level=3)
    advantages = [
        '곡선 길이에 비례한 균일한 점 분포',
        '물리적 대응성 보장',
        '곡률 변화에 자연스럽게 적응',
    ]
    for a in advantages:
        doc.add_paragraph(a, style='List Bullet')

    # 2.3 Coons 패치
    doc.add_heading('2.3 Coons 패치 (Coons Patch)', level=2)

    doc.add_heading('2.3.1 이론적 배경', level=3)
    doc.add_paragraph(
        'Coons (1967)가 제안한 곡면 보간법으로, 4개의 경계 곡선으로 정의된 곡면을 생성합니다.'
    )

    doc.add_heading('2.3.2 수학적 정의', level=3)
    doc.add_paragraph('4개 경계 곡선:')
    curves = [
        'C₀(t): s=0 경계 (아래)',
        'C₁(t): s=1 경계 (위)',
        'D₀(s): t=0 경계 (왼쪽)',
        'D₁(s): t=1 경계 (오른쪽)',
    ]
    for c in curves:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_paragraph('Coons 패치 공식:', style='Intense Quote')
    add_equation(doc, 'P(s,t) = Lc(s,t) + Ld(s,t) - B(s,t)')

    doc.add_paragraph('여기서:')
    add_equation(doc, 'Lc(s,t) = (1-s)·D₀(t) + s·D₁(t)')
    add_equation(doc, 'Ld(s,t) = (1-t)·C₀(s) + t·C₁(s)')
    add_equation(doc, 'B(s,t) = (1-s)(1-t)·P₀₀ + s(1-t)·P₁₀ + (1-s)t·P₀₁ + st·P₁₁')

    # 2.4 구조화 그리드 인덱싱
    doc.add_heading('2.4 구조화 그리드 인덱싱', level=2)

    doc.add_heading('2.4.1 목적', level=3)
    doc.add_paragraph('비정렬(unstructured) 메시 요소에 (i, j, k) 구조화 인덱스를 할당합니다.')

    doc.add_heading('2.4.2 BFS 기반 알고리즘', level=3)

    doc.add_paragraph('Step 1: 코너 요소 탐색', style='Intense Quote')
    doc.add_paragraph('코너 요소 조건: 정확히 3개의 면 이웃을 가진 요소')

    doc.add_paragraph('Step 2: 기하 기반 축 방향 결정', style='Intense Quote')
    doc.add_paragraph('요소 중심(centroid)의 상대 위치에서 축 방향 계산')

    doc.add_paragraph('Step 3: BFS 전파', style='Intense Quote')
    doc.add_paragraph('Queue를 이용한 너비 우선 탐색으로 인접 요소에 인덱스 할당')

    doc.add_paragraph('Step 4: 축 순서 정렬', style='Intense Quote')
    add_equation(doc, 'dim_k ≤ dim_j ≤ dim_i')

    doc.add_paragraph('Step 5: 노드 순서 정규화', style='Intense Quote')
    doc.add_paragraph('LS-DYNA HEX8 표준 노드 순서로 정렬')

    doc.add_page_break()

    # ============== 3. 변형 해석 알고리즘 ==============
    doc.add_heading('3. 변형 해석 알고리즘', level=1)

    # 3.1 변형 구배 텐서
    doc.add_heading('3.1 변형 구배 텐서 (Deformation Gradient)', level=2)

    doc.add_heading('3.1.1 정의', level=3)
    doc.add_paragraph(
        '변형 구배 텐서 F는 기준 형상에서 현재 형상으로의 국소 변형을 나타냅니다:'
    )
    add_equation(doc, 'F = ∂x / ∂X')

    doc.add_paragraph('여기서:')
    doc.add_paragraph('x: 현재(변형된) 좌표', style='List Bullet')
    doc.add_paragraph('X: 기준(미변형) 좌표', style='List Bullet')

    doc.add_heading('3.1.2 성분 표현', level=3)
    add_equation(doc, 'Fᵢⱼ = ∂xᵢ / ∂Xⱼ')

    doc.add_paragraph('3×3 행렬 형태:')
    matrix_text = '''        ┌ ∂x/∂X  ∂x/∂Y  ∂x/∂Z ┐
F =     │ ∂y/∂X  ∂y/∂Y  ∂y/∂Z │
        └ ∂z/∂X  ∂z/∂Y  ∂z/∂Z ┘'''
    p = doc.add_paragraph(matrix_text)
    p.runs[0].font.name = 'Consolas'

    doc.add_heading('3.1.3 유한요소 계산', level=3)
    doc.add_paragraph('자연 좌표계를 통한 계산:')
    add_equation(doc, 'F = J_def · J_ref⁻¹')

    doc.add_paragraph('여기서 야코비안 행렬:')
    add_equation(doc, 'Jᵢⱼ = Σₖ (∂Nₖ/∂ξⱼ) · xₖ,ᵢ')

    doc.add_heading('3.1.4 물리적 의미', level=3)
    meanings = [
        'det(F) > 0: 요소 유효 (양의 체적)',
        'det(F) < 0: 요소 뒤집힘 (음의 야코비안)',
        'det(F) = 1: 비압축성 변형',
        'F = I: 변형 없음',
    ]
    for m in meanings:
        doc.add_paragraph(m, style='List Bullet')

    # 3.2 변형률 텐서
    doc.add_heading('3.2 변형률 텐서 (Strain Tensor)', level=2)

    doc.add_heading('3.2.1 공학 변형률 (Engineering Strain)', level=3)
    doc.add_paragraph('작은 변형 가정 하에서:')
    add_equation(doc, 'ε = (1/2)(F + Fᵀ) - I')

    doc.add_paragraph('성분 표현:')
    add_equation(doc, 'εᵢⱼ = (1/2)(∂uᵢ/∂Xⱼ + ∂uⱼ/∂Xᵢ)')
    doc.add_paragraph('여기서 변위 u = x - X')

    doc.add_heading('3.2.2 Green-Lagrange 변형률', level=3)
    doc.add_paragraph('큰 변형에 적합한 비선형 변형률:')
    add_equation(doc, 'E = (1/2)(Fᵀ·F - I) = (1/2)(C - I)')
    doc.add_paragraph('여기서 C = Fᵀ·F 는 우 Cauchy-Green 변형 텐서')

    doc.add_heading('3.2.3 Voigt 표기법', level=3)
    doc.add_paragraph('대칭 텐서를 6개 독립 성분으로 표현:')
    voigt = '{ε} = [εxx, εyy, εzz, γxy, γyz, γxz]ᵀ'
    add_equation(doc, voigt)
    doc.add_paragraph('여기서 공학 전단변형률: γᵢⱼ = 2εᵢⱼ (i ≠ j)')

    doc.add_heading('3.2.4 주변형률 (Principal Strains)', level=3)
    doc.add_paragraph('변형률 텐서의 고유값 문제:')
    add_equation(doc, 'det(ε - λI) = 0')

    doc.add_paragraph('특성 방정식:')
    add_equation(doc, 'λ³ - I₁λ² + I₂λ - I₃ = 0')

    doc.add_paragraph('불변량:')
    doc.add_paragraph('I₁ = tr(ε) = εxx + εyy + εzz', style='List Bullet')
    doc.add_paragraph('I₂ = (1/2)[I₁² - tr(ε²)]', style='List Bullet')
    doc.add_paragraph('I₃ = det(ε)', style='List Bullet')

    doc.add_paragraph('Cardano 공식을 이용한 해석해:')
    add_equation(doc, 'p = I₂ - I₁²/3')
    add_equation(doc, 'q = 2I₁³/27 - I₁I₂/3 + I₃')
    add_equation(doc, 'λₖ = I₁/3 + 2√(-p/3)·cos((θ + 2πk)/3), k = 0,1,2')

    doc.add_heading('3.2.5 von Mises 등가 변형률', level=3)
    add_equation(doc, 'ε_eq = √(2/3 · ε\':ε\')')
    doc.add_paragraph('여기서 편차 변형률:')
    add_equation(doc, 'ε\'ᵢⱼ = εᵢⱼ - (1/3)εₖₖδᵢⱼ')

    doc.add_heading('3.2.6 체적 변형률', level=3)
    add_equation(doc, 'ε_vol = εxx + εyy + εzz = tr(ε)')

    # 3.3 응력 텐서
    doc.add_heading('3.3 응력 텐서 (Stress Tensor)', level=2)

    doc.add_heading('3.3.1 Hooke 법칙 (등방성 선형 탄성)', level=3)
    add_equation(doc, 'σ = λ·tr(ε)·I + 2μ·ε')

    doc.add_paragraph('Lamé 상수:')
    doc.add_paragraph('제1 Lamé 상수: λ = Eν / ((1+ν)(1-2ν))', style='List Bullet')
    doc.add_paragraph('전단 계수: μ = G = E / (2(1+ν))', style='List Bullet')

    doc.add_heading('3.3.2 탄성 강성 행렬', level=3)
    doc.add_paragraph('Voigt 표기법에서의 구성 행렬:')
    add_equation(doc, '{σ} = [C]{ε}')

    doc.add_paragraph('여기서:')
    doc.add_paragraph('C₁₁ = λ + 2μ = E(1-ν) / ((1+ν)(1-2ν))', style='List Bullet')
    doc.add_paragraph('C₁₂ = λ = Eν / ((1+ν)(1-2ν))', style='List Bullet')
    doc.add_paragraph('C₄₄ = μ = E / (2(1+ν))', style='List Bullet')

    doc.add_heading('3.3.3 von Mises 응력 (항복 기준)', level=3)
    add_equation(doc, 'σ_vm = √(3/2 · s:s)')

    doc.add_paragraph('여기서 편차 응력:')
    add_equation(doc, 'sᵢⱼ = σᵢⱼ - (1/3)σₖₖδᵢⱼ = σᵢⱼ - σₘδᵢⱼ')

    doc.add_paragraph('평균 응력 (정수압):')
    add_equation(doc, 'σₘ = (1/3)(σxx + σyy + σzz)')

    doc.add_paragraph('주응력 형태:')
    add_equation(doc, 'σ_vm = (1/√2)√[(σ₁-σ₂)² + (σ₂-σ₃)² + (σ₃-σ₁)²]')

    doc.add_heading('3.3.4 응력 삼축성 (Stress Triaxiality)', level=3)
    add_equation(doc, 'η = σₘ / σ_vm')

    doc.add_paragraph('물리적 의미:')
    doc.add_paragraph('η > 1/3: 인장 지배 (취성 파괴 위험)', style='List Bullet')
    doc.add_paragraph('η ≈ 0: 순수 전단', style='List Bullet')
    doc.add_paragraph('η < 0: 압축 지배', style='List Bullet')

    doc.add_heading('3.3.5 최대 전단응력 (Tresca)', level=3)
    add_equation(doc, 'τ_max = (σ₁ - σ₃) / 2')

    doc.add_page_break()

    # ============== 4. 수치 해석 기법 ==============
    doc.add_heading('4. 수치 해석 기법', level=1)

    # 4.1 가우스 적분
    doc.add_heading('4.1 가우스 적분 (Gauss Quadrature)', level=2)

    doc.add_heading('4.1.1 이론', level=3)
    doc.add_paragraph('적분을 가중 합으로 근사:')
    add_equation(doc, '∫₋₁¹ f(ξ)dξ ≈ Σᵢ wᵢ f(ξᵢ)')

    doc.add_heading('4.1.2 1D 가우스점', level=3)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['차수 n', '위치 ξᵢ', '가중치 wᵢ']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')

    gauss_data = [
        ('1', '0', '2'),
        ('2', '±1/√3', '1'),
        ('3', '0, ±√(3/5)', '8/9, 5/9'),
    ]
    for i, (n, pos, w) in enumerate(gauss_data):
        row = table.rows[i+1]
        row.cells[0].text = n
        row.cells[1].text = pos
        row.cells[2].text = w

    doc.add_paragraph()

    doc.add_heading('4.1.3 3D HEX8 요소', level=3)

    doc.add_paragraph('1-점 적분 (Reduced Integration):', style='Intense Quote')
    doc.add_paragraph('위치: (0, 0, 0), 가중치: 8', style='List Bullet')
    doc.add_paragraph('장점: 빠름, shear locking 방지', style='List Bullet')
    doc.add_paragraph('단점: hourglass 모드 가능', style='List Bullet')

    doc.add_paragraph('2×2×2 = 8-점 적분 (Full Integration):', style='Intense Quote')
    doc.add_paragraph('위치: (±1/√3, ±1/√3, ±1/√3), 가중치: 각 1', style='List Bullet')
    doc.add_paragraph('장점: 정확함', style='List Bullet')
    doc.add_paragraph('단점: shear locking 가능', style='List Bullet')

    # 4.2 형상 함수
    doc.add_heading('4.2 형상 함수 (Shape Functions)', level=2)

    doc.add_heading('4.2.1 HEX8 요소', level=3)
    doc.add_paragraph('8절점 육면체 요소의 형상함수:')
    add_equation(doc, 'Nᵢ(ξ,η,ζ) = (1/8)(1 + ξᵢξ)(1 + ηᵢη)(1 + ζᵢζ)')

    doc.add_heading('4.2.2 형상함수 미분', level=3)
    add_equation(doc, '∂Nᵢ/∂ξ = (1/8)ξᵢ(1 + ηᵢη)(1 + ζᵢζ)')
    add_equation(doc, '∂Nᵢ/∂η = (1/8)(1 + ξᵢξ)ηᵢ(1 + ζᵢζ)')
    add_equation(doc, '∂Nᵢ/∂ζ = (1/8)(1 + ξᵢξ)(1 + ηᵢη)ζᵢ')

    doc.add_heading('4.2.3 TET4 요소', level=3)
    doc.add_paragraph('4절점 사면체 요소:')
    add_equation(doc, 'N₁ = 1 - ξ - η - ζ')
    add_equation(doc, 'N₂ = ξ')
    add_equation(doc, 'N₃ = η')
    add_equation(doc, 'N₄ = ζ')
    doc.add_paragraph('좌표계: (ξ, η, ζ) ∈ [0,1], ξ + η + ζ ≤ 1')

    # 4.3 야코비안 행렬
    doc.add_heading('4.3 야코비안 행렬 (Jacobian Matrix)', level=2)

    doc.add_heading('4.3.1 정의', level=3)
    doc.add_paragraph('자연 좌표에서 물리 좌표로의 변환 야코비안:')
    matrix_j = '''        ┌ ∂x/∂ξ  ∂y/∂ξ  ∂z/∂ξ ┐
J =     │ ∂x/∂η  ∂y/∂η  ∂z/∂η │
        └ ∂x/∂ζ  ∂y/∂ζ  ∂z/∂ζ ┘'''
    p = doc.add_paragraph(matrix_j)
    p.runs[0].font.name = 'Consolas'

    doc.add_heading('4.3.2 계산', level=3)
    add_equation(doc, 'Jᵢⱼ = Σₖ (∂Nₖ/∂ξⱼ) · xₖᵢ')

    doc.add_heading('4.3.3 행렬식', level=3)
    add_equation(doc, '|J| = det(J)')
    doc.add_paragraph('물리적 의미: 자연 좌표계 단위 체적에 대한 물리 좌표계 체적비')
    doc.add_paragraph('|J| > 0: 유효한 요소', style='List Bullet')
    doc.add_paragraph('|J| ≤ 0: 뒤집힌 요소 (무효)', style='List Bullet')

    doc.add_page_break()

    # ============== 5. 재료 모델 ==============
    doc.add_heading('5. 재료 모델', level=1)

    doc.add_heading('5.1 등방성 선형 탄성', level=2)

    doc.add_heading('5.1.1 입력 물성', level=3)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['물성', '기호', '단위']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')

    prop_data = [
        ('탄성 계수', 'E', 'MPa'),
        ('포아송비', 'ν', '-'),
        ('밀도', 'ρ', 'ton/mm³'),
    ]
    for i, (name, sym, unit) in enumerate(prop_data):
        row = table.rows[i+1]
        row.cells[0].text = name
        row.cells[1].text = sym
        row.cells[2].text = unit

    doc.add_paragraph()

    doc.add_heading('5.1.2 유도 물성', level=3)
    doc.add_paragraph('전단 계수:')
    add_equation(doc, 'G = E / (2(1+ν))')

    doc.add_paragraph('체적 탄성 계수:')
    add_equation(doc, 'K = E / (3(1-2ν))')

    doc.add_paragraph('Lamé 상수:')
    add_equation(doc, 'λ = Eν / ((1+ν)(1-2ν)) = K - 2G/3')

    doc.add_heading('5.1.3 물성 제한', level=3)
    doc.add_paragraph('E > 0 (양의 탄성 계수)', style='List Bullet')
    doc.add_paragraph('-1 < ν < 0.5 (포아송비 범위)', style='List Bullet')
    doc.add_paragraph('ν = 0.5: 비압축성 (수치적 문제 발생)', style='List Bullet')

    doc.add_heading('5.2 점탄성 재료 (MAT_VISCOELASTIC)', level=2)

    doc.add_heading('5.2.1 입력 물성', level=3)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['물성', '기호', '의미']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')

    visco_data = [
        ('체적 탄성 계수', 'K (BULK)', '부피 변형 저항'),
        ('단시간 전단 계수', 'G₀', '즉시 전단 응답'),
        ('장시간 전단 계수', 'Gᵢ', '평형 전단 응답'),
        ('감쇠 계수', 'β', '이완 속도'),
    ]
    for i, (name, sym, meaning) in enumerate(visco_data):
        row = table.rows[i+1]
        row.cells[0].text = name
        row.cells[1].text = sym
        row.cells[2].text = meaning

    doc.add_paragraph()

    doc.add_heading('5.2.2 프리스트레스 계산', level=3)
    doc.add_paragraph('장시간 거동(평형 상태)을 사용:')
    add_equation(doc, 'G = Gᵢ')

    doc.add_paragraph('등가 탄성 계수 계산:')
    add_equation(doc, 'E = 9KG / (3K + G)')
    add_equation(doc, 'ν = (3K - 2G) / (2(3K + G))')

    doc.add_page_break()

    # ============== 6. 참고 문헌 ==============
    doc.add_heading('6. 참고 문헌', level=1)

    doc.add_heading('6.1 메시 매핑', level=2)
    refs_mapping = [
        'Gordon, W.J., Hall, C.A. (1973). "Transfinite element methods: Blending-function interpolation over arbitrary curved element domains." Numerische Mathematik, 21(2), 109-129.',
        'Coons, S.A. (1967). "Surfaces for computer-aided design of space forms." MIT Project MAC, TR-41.',
        'Farin, G. (2002). Curves and Surfaces for CAGD: A Practical Guide, 5th Edition. Morgan Kaufmann.',
    ]
    for i, ref in enumerate(refs_mapping):
        doc.add_paragraph(f'[{i+1}] {ref}')

    doc.add_heading('6.2 연속체 역학', level=2)
    refs_mech = [
        'Malvern, L.E. (1969). Introduction to the Mechanics of a Continuous Medium. Prentice-Hall.',
        'Holzapfel, G.A. (2000). Nonlinear Solid Mechanics: A Continuum Approach for Engineering. Wiley.',
        'Bonet, J., Wood, R.D. (2008). Nonlinear Continuum Mechanics for Finite Element Analysis, 2nd Edition. Cambridge University Press.',
    ]
    for i, ref in enumerate(refs_mech):
        doc.add_paragraph(f'[{i+4}] {ref}')

    doc.add_heading('6.3 유한요소법', level=2)
    refs_fem = [
        'Hughes, T.J.R. (2000). The Finite Element Method: Linear Static and Dynamic Finite Element Analysis. Dover Publications.',
        'Zienkiewicz, O.C., Taylor, R.L. (2005). The Finite Element Method, 6th Edition. Butterworth-Heinemann.',
        'Bathe, K.J. (2006). Finite Element Procedures. Prentice Hall.',
    ]
    for i, ref in enumerate(refs_fem):
        doc.add_paragraph(f'[{i+7}] {ref}')

    doc.add_heading('6.4 LS-DYNA', level=2)
    refs_dyna = [
        'Hallquist, J.O. (2006). LS-DYNA Theory Manual. Livermore Software Technology Corporation.',
        'LS-DYNA Keyword User\'s Manual. LSTC.',
    ]
    for i, ref in enumerate(refs_dyna):
        doc.add_paragraph(f'[{i+10}] {ref}')

    # 저장
    doc.save('d:/KooRemapper/docs/KooRemapper_Theory_Document.docx')
    print('Document saved: d:/KooRemapper/docs/KooRemapper_Theory_Document.docx')

if __name__ == '__main__':
    create_document()
