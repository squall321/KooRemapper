"""
KooRemapper Manual Builder
- Reads KooRemapper_Manual.md
- Adds meshfix section
- Adds table/figure captions throughout
- Generates KooRemapper_Manual.docx with proper Word styles
  - Heading 1/2/3 for section titles
  - Caption style ABOVE tables
  - Caption style BELOW figures/code-diagrams
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ─────────────────────────────────────────────────────────────────────────────
# meshfix section text (inserted before "## 40. 수학 이론")
# ─────────────────────────────────────────────────────────────────────────────

MESHFIX_SECTION = '''
## 40. meshfix — TET4 재메시 (Gmsh 기반)

### 용도
기존 TET4 파트를 Gmsh를 통해 **완전 재메시**하여 요소 품질을 개선하는 명령.
STL 경계 추출 → Gmsh 실행 → MSH2 파싱 → 원본 K파일에 스플라이스하는 파이프라인으로 동작하며,
Gmsh 실행 파일(`gmsh.exe`)이 `dist/gmsh/` 또는 `dist/gmsh-<ver>/` 디렉터리에 있어야 한다.

### 사용법

```bash
KooRemapper.exe meshfix <config.yaml>
```

### YAML 설정 전체

```yaml
model:   input.k      # 입력 K파일
output:  output.k     # 출력 K파일
pid:     1            # 재메시할 파트 ID (TET4)

# ─── 요소 크기 제어 ────────────────────────────────────────────────────────
lc_target:    5.0     # 목표 평균 요소 크기 (기본: 1.0, 단위: 모델 단위)
lc_min:      -1.0     # 최소 요소 크기 (-1 = 자동: edge_min×0.8 또는 dt 기반)
lc_max:      -1.0     # 최대 요소 크기 (-1 = lc_target × 2)

# ─── dt 기반 lc_min (lc_min 대신 사용 가능) ───────────────────────────────
min_dt:       1.0e-6  # LS-DYNA explicit 시간 증분 하한 (초)
density:      2.7e-9  # 밀도 (t/mm³)
E:            70000.0 # 탄성계수 (MPa)
nu:           0.33    # 포아송 비

# ─── 얇은 형상 처리 ────────────────────────────────────────────────────────
min_layers_thin: 2    # 얇은 방향 최소 요소 레이어 수 (기본: 2)

# ─── 적응형 사이즈 필드 ────────────────────────────────────────────────────
adaptive:     true    # bbox 코너 거리 기반 MathEval 필드 활성화 (기본: true)
decay_factor: 8.0     # 코너 세밀 영역 크기 = lc_min × decay_factor

# ─── Gmsh 메셔 설정 ────────────────────────────────────────────────────────
algorithm:       hxt  # 3D 메셔: hxt (병렬, 기본) | frontal3d | del3d
optimize_netgen: true # Gmsh 내장 Netgen 최적화 활성화
optimize_passes: 3    # Mesh 3 이후 추가 OptimizeMesh "Netgen" 호출 횟수

# ─── 표면 STL 전처리 ───────────────────────────────────────────────────────
refine_surface:  auto # auto | 0(off) | 1~3 (conforming feature-edge 세분화)
smooth_surface:  0    # feature-edge Laplacian 스무딩 스텝 수 (0=off)

# ─── 경계 노드 처리 ────────────────────────────────────────────────────────
boundary_nodes: free  # free | fixed | snap
snap_tolerance: 0.001 # snap 모드 탐색 반경

# ─── 품질 보고 ─────────────────────────────────────────────────────────────
quality_check:  true  # 재메시 후 스케일드 자코비안 보고 활성화
warn_min_jac:   0.15  # 이 값 미만 요소에 경고 출력

# ─── 패치 폴리싱 (실험적) ─────────────────────────────────────────────────
polish:          false # 나쁜 요소 클러스터 로컬 재메시 (기본: off)
polish_jac:      0.10  # polish 대상 임계값 (J < polish_jac)
polish_max_iter: 2     # 최대 반복 횟수
```

**표 40-1. meshfix YAML 옵션 요약**

| 파라미터 | 기본값 | 설명 |
|---|---|---|
| `lc_target` | 1.0 | 목표 평균 요소 크기 |
| `lc_min` | auto | 최솟값 (-1=자동, dt 기반 또는 edge_min×0.8) |
| `lc_max` | auto | 최댓값 (-1=lc_target×2) |
| `adaptive` | true | MathEval bbox 코너 거리 필드 |
| `algorithm` | hxt | Gmsh 3D 메셔 선택 |
| `optimize_passes` | 3 | 추가 Netgen 최적화 횟수 |
| `refine_surface` | auto | STL feature-edge 세분화 레벨 |
| `warn_min_jac` | 0.15 | 품질 경고 임계값 |
| `polish` | false | 나쁜 요소 로컬 재메시 |

### 동작 파이프라인

```
[1] K파일 로드 → TET4 파트 추출
        ↓
[2] 메시 분석
    - lc_min/max 자동 계산
    - geomThin 감지 (min_bbox < avg_bbox × 0.3)
    - autoRefineSurface 레벨 결정
        ↓
[3] 경계 STL 추출
    - 비다양체 에지 필터 (Stage 1: 점수 기반)
    - bbox 표면 노드 필터 (Stage 2: 레이캐스팅)
    - 최대 연결 컴포넌트만 유지
        ↓
[4] STL 전처리
    - refine_surface: conforming feature-edge 세분화 (dihedral > 40°)
    - smooth_surface: feature-edge Laplacian 스무딩
        ↓
[5] Gmsh .geo 스크립트 생성
    - ClassifySurfaces{40°} → CreateGeometry → Volume
    - MathEval 적응형 사이즈 필드 (또는 인플레인 4코너 필드)
        ↓
[6] Gmsh 실행 (HXT 알고리즘)
    Mesh 2 → Mesh 3 → OptimizeMesh "Netgen" × N
        ↓
[7] MSH2 파싱 → 스케일드 자코비안 품질 검사
        ↓
[8] (polish=true) 나쁜 클러스터 로컬 재메시
        ↓
[9] K파일 스플라이스 (기존 노드/요소 교체, 다른 파트 보존)
```

> **그림 40-1. meshfix 처리 파이프라인 — 입력 K파일에서 재메시된 출력 K파일까지의 전체 데이터 흐름을 단계별로 나타낸다. [2]~[4]는 전처리, [5]~[6]은 Gmsh 처리, [7]~[9]는 후처리에 해당한다.**

### 적응형 사이즈 필드

MathEval 거리 필드: Gmsh 임베디드 Point 엔티티 없이 순수 수식으로 구현 (HXT 호환).

```
Field[1] = MathEval;
Field[1].F = "Sqrt(Min(Min(d000,d001),Min(...,d111)))";
    ← bbox 8개 코너까지의 최솟값 거리 (수식)

Field[2] = Threshold;
    SizeMin = lc_min     ← 코너 근처: 세밀
    SizeMax = lc_max     ← 내부: 큰 요소

Field[3] = MathEval; F = "lc_target";
    ← 균일 배경 필드

Background Field = Min(Field[2], Field[3]);
```

> **그림 40-2. MathEval 적응형 사이즈 필드 구성 — 8개 bbox 코너까지의 최솟값 거리를 기준으로 코너 근처에서는 lc_min, 내부에서는 lc_target의 요소 크기를 유도한다.**

**얇은 형상 (geomThin) 처리:**
min_bbox < avg_bbox × 0.3인 경우 자동 감지하여 다음을 전환한다.

**표 40-2. geomThin 감지 시 동작 변경**

| 항목 | 일반 형상 | 얇은 형상 (geomThin) |
|---|---|---|
| Mesh 2 (표면 메시) | 실행 | **스킵** |
| 사이즈 필드 | 8코너 3D 필드 | **4코너 인플레인 필드** |
| Netgen OptimizePasses | 실행 | **스킵** |
| autoRefineSurface | 0 (Mesh 2가 대신함) | 1 (필요 시) |

### 스케일드 자코비안 (Scaled Jacobian)

TET4 요소의 형상 품질을 나타내는 무차원 지표.

$$J_s = \\frac{6\\sqrt{2} \\cdot V}{L_{max}^3}$$

여기서 V = TET4 체적, $L_{max}$ = 가장 긴 엣지 길이. 정규 TET4에서 $J_s = 1.0$.

**표 40-3. 스케일드 자코비안 판정 기준**

| 범위 | 판정 | 설명 |
|---|---|---|
| $J_s \\geq 0.5$ | **우수** | LS-DYNA 권장 범위 |
| $0.2 \\leq J_s < 0.5$ | 양호 | 실용적으로 허용 |
| $0.15 \\leq J_s < 0.2$ | 주의 | warn_min_jac 경고 기준 |
| $0.0 < J_s < 0.15$ | **불량** | 재메시 또는 개선 필요 |
| $J_s \\leq 0$ | 역전 | 음의 체적 — 해석 불가 |

### 기하학적 품질 한계

90° 직각 코너에 인접한 TET4는 기하 구속으로 인해 이론적 최솟값이 존재한다.

$$J_{s,min}^{corner} \\approx 0.03 \\sim 0.07 \\quad \\text{(기하 구속, 요소 크기와 무관)}$$

이 한계는 기하학 수정(코너 라운딩, 필렛 추가) 없이는 개선할 수 없다.

### 패치 폴리싱 (polish)

`polish: true` 설정 시 재메시 후 추가 국소 개선을 시도한다.

**표 40-4. 패치 폴리싱 이중 품질 게이트**

| 조건 | 식 | 의미 |
|---|---|---|
| 최솟값 Jac 유지 | $J_{min,new} \\geq J_{min,orig} \\times 0.95$ | 최솟값 5% 이상 회귀 시 거부 |
| 불량 수 감소 | $N_{bad,new} < N_{bad,orig}$ | 불량 요소 수가 줄어야 수락 |

두 조건을 모두 만족한 패치만 메시에 머지한다. 하나라도 불만족이면 해당 클러스터는 원본 유지.

### 경계 노드 처리 모드

**표 40-5. boundary_nodes 모드별 동작**

| 모드 | 동작 | 사용 예 |
|---|---|---|
| `free` (기본) | 모든 Gmsh 노드에 새 ID 부여 | 독립 파트 재메시 |
| `fixed` | 경계 노드를 원본 ID에 고정 (좌표 매칭) | 인접 파트와 공유 노드 |
| `snap` | 인접 파트 노드에 snap_tolerance 이내이면 병합 | 파트 간 접합 재메시 |

### 실행 예시

```yaml
# 예시: arc30 평면 TET4 메시 재메시
model:      examples/arc30/arc30_flat_tet.k
output:     output/remeshed.k
pid:        1
lc_target:  5.0
adaptive:   true
warn_min_jac: 0.15
```

실행 출력:
```
PID 1 TET4:              3000
  lc_min=1  lc_max=10
Gmsh TET4:               34695
Min scaled Jacobian:     0.0675
Avg scaled Jacobian:     0.405
[WARN] 847 elements below warn_min_jac=0.15
Total time: 13.6 s
```

> **그림 40-3. meshfix 실행 출력 예 — 입력 3,000개 TET4가 34,695개로 재메시됨. 스케일드 자코비안 통계와 품질 경고가 출력된다. 847개 불량 요소는 박스 90° 코너의 기하 구속에 의한 것으로, 기하학 수정 없이는 개선 불가하다.**

### 주의사항

- **Gmsh 필수**: `dist/gmsh/gmsh.exe` 또는 `dist/gmsh-<ver>/gmsh.exe` 위치에 배치 필요
- **TET4 전용**: 입력 파트는 TET4 (또는 퇴화 HEX8) 형식이어야 함
- **처리 시간**: 10만 요소 이상에서 수 분 소요 가능
- **polish 제한**: `polish: true`는 실험적 기능. 90° 코너 구속 형상에서는 불량 수 감소 불가로 자동 스킵

---
'''

# ─────────────────────────────────────────────────────────────────────────────
# Table caption registry
# (line_hint: first |-row line number, 1-based; caption text)
# ─────────────────────────────────────────────────────────────────────────────

# We'll auto-detect all tables and assign captions in order.
# Tables that already have a "**표 N" line above them will be skipped.
# For named tables (in meshfix section added above), we embed the caption
# directly in the markdown text.

TABLE_CAPTIONS = [
    "표 1-1. KooRemapper 핵심 기능 범주 — 각 범주별 주요 기능과 해당 명령어를 요약한다.",
    "표 7-1. squeeze YAML 파트 설정 예 — 직접 변형률(eps_x/y/z)과 등방 팽창(swelling) 두 가지 방법의 비교.",
    "표 8-1. generate-var 두께 분포 정의 — 영역(zone)별 lc와 두께를 지정하여 변밀도 메시를 생성한다.",
    "표 9-1. unfold 파라미터 — 굽힘 메시 전개 시 호(arc), 너비(width), 두께(thickness) 축 방향 설정.",
    "표 13-1. bend YAML 설정 파라미터 — 굽힘 반경, 각도, 중립면 위치, 굽힘 축 방향 등 핵심 파라미터.",
    "표 14-1. indent YAML 설정 파라미터 — 압입 깊이, 위치, 반경, 방향 등 압입/엠보싱 제어 파라미터.",
    "표 15-1. formstrain 출력 — 이면각 기반 소성 변형률 계산 결과 및 LS-DYNA *INITIAL_STRAIN_SOLID 출력.",
    "표 16-1. convert 지원 변환 유형 — TET4→TET10, HEX8→HEX20, QUAD4→QUAD8, TRIA3→TRIA6 변환 지원.",
    "표 17-1. refine 세분화 비율 — 요소 유형별 1:2, 1:3 세분화 시 생성 요소 수 비교.",
    "표 18-1. ELFORM 코드 목록 — LS-DYNA 요소 공식(ELFORM) 번호와 각 공식의 특성 요약.",
    "표 19-1. disconnect 모드 — full/czm/mefem 세 가지 노드 분리 모드와 생성 키워드.",
    "표 20-1. IGA 생성 파일 — 파트별 IGA NURBS 박스 파일과 메인 파일의 *INCLUDE 구조.",
    "표 21-1. warpage 보정 파라미터 — 워피지 측정 기준면, 보정 방향, 스케일 팩터 설정.",
    "표 22-1. offset 모드 — tied/czm/contact 세 가지 인터페이스 연결 방법과 생성 키워드.",
    "표 22-2. offset local_normals 효과 — 전역 평균 법선 대비 로컬 법선 사용 시 품질 개선.",
    "표 23-1. matswap 번들 파라미터 타입 — ID 접두어(HGID/LCID/SECID/MID/PID)별 자동 인식 규칙.",
    "표 24-1. matdb 재료 매칭 규칙 — 제목(title)/이름(name)/태그(tag) 우선순위 기반 자동 매칭.",
    "표 24-2. matdb 구조 카드 타입 — MAT_ELASTIC, MAT_024, MAT_RIGID 등 지원 카드 목록.",
    "표 25-1. contact 접촉 type 값 목록 — YAML type 키워드와 LS-DYNA *CONTACT_* 키워드 대응.",
    "표 25-2. contact modify 수정 가능 필드 — Card 1/2/A/C 필드명과 대응하는 LS-DYNA 필드.",
    "표 25-3. contact Optional Card 지원 목록 — A~G 카드별 주요 파라미터와 기본값.",
    "표 26-1. load 유형 목록 — 지원하는 하중 종류와 각 하중의 적용 대상(노드/파트/세그먼트).",
    "표 27-1. boundary 조건 유형 — 지원하는 경계 조건 종류와 자유도(DOF) 구성.",
    "표 28-1. rbe 구속 유형 — RBE2/RBE3 구속 조건 생성 방법과 마스터/슬레이브 설정.",
    "표 29-1. implicit 변환 레벨 (1~8) — 공격적→보수적 순으로 정렬된 8단계 변환 레벨과 활성화 키워드.",
    "표 29-2. implicit 오버라이드 파라미터 — dt0, dtmax, nsolvr 등 사용자 정의 시 기본값을 덮어쓰는 파라미터.",
    "표 30-1. modal 해석 파라미터 — 모드 수, 주파수 범위, 고유값 해석 방법(eigmth) 코드 목록.",
    "표 35-1. ALE 프리셋 목록 (14종) — 기체/액체/폭약/진공 프리셋별 적용 재료 모델과 상태방정식.",
    "표 36-1. stabilize 12단계 설정 — 단계별 누적 적용 안정화 옵션과 활성화 조건.",
    "표 37-1. database 프리셋 종류 — crash/drop/nve/all 프리셋별 출력 키워드 목록.",
    "표 39-1. assemble 오퍼레이션 목록 — type 필드로 지정 가능한 전체 오퍼레이션과 주요 파라미터.",
]

CODE_DIAGRAM_CAPTIONS = {
    # section_keyword → caption
    "KooRemapper.exe <command>": "그림 3-1. KooRemapper 전체 명령어 목록 — 기능 범주별로 분류한 지원 명령어 리스트. 각 명령어의 상세 설명은 해당 섹션을 참조한다.",
    "meshfix 처리 파이프라인은 위 그림 참조": None,  # handled inline
}


def read_file(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.readlines()


def write_file(path, lines):
    with open(path, 'w', encoding='utf-8') as f:
        f.writelines(lines)


def is_table_start(line):
    return line.strip().startswith('|') and '|' in line[1:]


# ─────────────────────────────────────────────────────────────────────────────
# LaTeX → Unicode + Word run formatting
# ─────────────────────────────────────────────────────────────────────────────

LATEX_SYMBOLS = {
    # Greek lowercase
    r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
    r'\epsilon': 'ϵ', r'\varepsilon': 'ε', r'\zeta': 'ζ', r'\eta': 'η',
    r'\theta': 'θ', r'\vartheta': 'ϑ', r'\iota': 'ι', r'\kappa': 'κ',
    r'\lambda': 'λ', r'\mu': 'μ', r'\nu': 'ν', r'\xi': 'ξ',
    r'\omicron': 'ο', r'\pi': 'π', r'\varpi': 'ϖ', r'\rho': 'ρ',
    r'\varrho': 'ϱ', r'\sigma': 'σ', r'\varsigma': 'ς', r'\tau': 'τ',
    r'\upsilon': 'υ', r'\phi': 'ϕ', r'\varphi': 'φ', r'\chi': 'χ',
    r'\psi': 'ψ', r'\omega': 'ω',
    # Greek uppercase
    r'\Gamma': 'Γ', r'\Delta': 'Δ', r'\Theta': 'Θ', r'\Lambda': 'Λ',
    r'\Xi': 'Ξ', r'\Pi': 'Π', r'\Sigma': 'Σ', r'\Upsilon': 'Υ',
    r'\Phi': 'Φ', r'\Psi': 'Ψ', r'\Omega': 'Ω',
    # Operators
    r'\cdot': '·', r'\times': '×', r'\div': '÷', r'\pm': '±', r'\mp': '∓',
    r'\leq': '≤', r'\geq': '≥', r'\neq': '≠', r'\approx': '≈',
    r'\equiv': '≡', r'\sim': '∼', r'\propto': '∝',
    r'\in': '∈', r'\notin': '∉', r'\subset': '⊂', r'\supset': '⊃',
    r'\cup': '∪', r'\cap': '∩', r'\emptyset': '∅',
    r'\forall': '∀', r'\exists': '∃', r'\partial': '∂',
    r'\nabla': '∇', r'\infty': '∞', r'\to': '→', r'\rightarrow': '→',
    r'\leftarrow': '←', r'\Rightarrow': '⇒', r'\Leftarrow': '⇐',
    r'\Leftrightarrow': '⇔', r'\mapsto': '↦',
    r'\sum': 'Σ', r'\prod': '∏', r'\int': '∫',
    r'\ldots': '…', r'\cdots': '⋯', r'\vdots': '⋮', r'\ddots': '⋱',
    # NOTE: \hat, \bar, \overline, \tilde, \widehat are handled by stripping below
    # (they take an argument like \hat{x}; the symbol form alone is rare here)
    r'\quad': '   ', r'\qquad': '      ',
    r'\,': ' ', r'\;': ' ', r'\:': ' ', r'\!': '',
    r'\%': '%', r'\$': '$', r'\&': '&', r'\#': '#', r'\_': '_',
    r'\{': '{', r'\}': '}', r'\\': '',
    r'\sqrt': '√',
}

# Unicode subscript/superscript digits (used for short single-char sub/sup)
_SUB_DIGITS = {'0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅','6':'₆','7':'₇','8':'₈','9':'₉',
               '+':'₊','-':'₋','=':'₌','(':'₍',')':'₎','i':'ᵢ','j':'ⱼ','x':'ₓ'}
_SUP_DIGITS = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹',
               '+':'⁺','-':'⁻','=':'⁼','(':'⁽',')':'⁾','i':'ⁱ','n':'ⁿ'}


def _strip_braces(s):
    s = s.strip()
    if s.startswith('{') and s.endswith('}'):
        return s[1:-1]
    return s


def _balanced_arg(text, start):
    """Return content inside {...} starting at start (must point to '{').
    Returns (content, end_index_after_closing_brace)."""
    if start >= len(text) or text[start] != '{':
        return None, start
    depth = 0
    i = start
    while i < len(text):
        if text[i] == '{':
            depth += 1
        elif text[i] == '}':
            depth -= 1
            if depth == 0:
                return text[start+1:i], i+1
        i += 1
    return None, start


def _try_unicode_sub(text):
    """Convert text to Unicode subscript if all chars are mappable."""
    if all(c in _SUB_DIGITS for c in text):
        return ''.join(_SUB_DIGITS[c] for c in text), True
    return text, False


def _try_unicode_sup(text):
    """Convert text to Unicode superscript if all chars are mappable."""
    if all(c in _SUP_DIGITS for c in text):
        return ''.join(_SUP_DIGITS[c] for c in text), True
    return text, False


class MathToken:
    """Token for a math expression: ('text'|'sub'|'sup'|'frac'|'sqrt', content, ...)."""
    __slots__ = ('kind', 'data', 'extra')
    def __init__(self, kind, data, extra=None):
        self.kind = kind
        self.data = data
        self.extra = extra


def parse_latex_math(text):
    """Parse LaTeX math text into a list of MathToken. Greedy left-to-right."""
    # First, replace LaTeX symbols with Unicode (longest-first to avoid prefix issues)
    sorted_syms = sorted(LATEX_SYMBOLS.items(), key=lambda kv: -len(kv[0]))
    # Use a regex to replace whole-command tokens (followed by non-letter)
    def repl(m):
        tok = m.group(0)
        return LATEX_SYMBOLS.get(tok, tok)
    pat = re.compile('|'.join(re.escape(k) for k, _ in sorted_syms) + r'(?![A-Za-z])')
    text = pat.sub(repl, text)

    # Strip \mathbf{...}, \mathrm{...}, \text{...}, \boldsymbol{...}
    for cmd in ['mathbf', 'mathrm', 'mathit', 'text', 'textrm', 'boldsymbol', 'operatorname']:
        pattern = r'\\' + cmd + r'\s*\{([^{}]*)\}'
        text = re.sub(pattern, r'\1', text)

    # Strip \left and \right (delimiter sizing) — keep just the delimiter
    text = re.sub(r'\\left\s*([\(\[\{\|])', r'\1', text)
    text = re.sub(r'\\right\s*([\)\]\}\|])', r'\1', text)
    text = re.sub(r'\\left\.', '', text)
    text = re.sub(r'\\right\.', '', text)
    # Misc spacing/control commands (drop entirely)
    text = re.sub(r'\\(?:displaystyle|textstyle|scriptstyle|scriptscriptstyle)\b', '', text)
    text = re.sub(r'\\(?:limits|nolimits)\b', '', text)
    # \overline{x}, \underline{x}, \widehat{x}, \widetilde{x} — keep content
    for cmd in ['overline', 'underline', 'widehat', 'widetilde', 'overrightarrow',
                'overleftarrow', 'hat', 'tilde', 'bar', 'vec', 'dot', 'ddot']:
        pattern = r'\\' + cmd + r'\s*\{([^{}]*)\}'
        text = re.sub(pattern, r'\1', text)

    # Math function names (ln, log, sin, cos, ...) — render as plain text
    for fn in ['ln', 'log', 'sin', 'cos', 'tan', 'cot', 'sec', 'csc',
               'arcsin', 'arccos', 'arctan', 'sinh', 'cosh', 'tanh',
               'exp', 'lim', 'min', 'max', 'sup', 'inf', 'det', 'arg']:
        text = re.sub(r'\\' + fn + r'(?![A-Za-z])', fn, text)

    # Now parse: text + sub/sup + frac + sqrt
    tokens = []
    i = 0
    buf = ''

    def flush_buf():
        nonlocal buf
        if buf:
            tokens.append(MathToken('text', buf))
            buf = ''

    while i < len(text):
        c = text[i]

        # \frac{a}{b}
        if text[i:i+5] == r'\frac':
            flush_buf()
            j = i + 5
            # skip whitespace
            while j < len(text) and text[j] == ' ':
                j += 1
            num, j = _balanced_arg(text, j)
            while j < len(text) and text[j] == ' ':
                j += 1
            den, j = _balanced_arg(text, j)
            if num is not None and den is not None:
                tokens.append(MathToken('frac', (num, den)))
                i = j
                continue
            else:
                buf += '\\frac'
                i += 5
                continue

        # √{x} or √x — \sqrt was already converted to '√' above
        if c == '√' and i + 1 < len(text) and text[i+1] == '{':
            flush_buf()
            arg, j = _balanced_arg(text, i+1)
            tokens.append(MathToken('sqrt', arg if arg is not None else ''))
            i = j
            continue

        # Subscript: _{xxx} or _x
        if c == '_':
            flush_buf()
            if i + 1 < len(text) and text[i+1] == '{':
                arg, j = _balanced_arg(text, i+1)
                tokens.append(MathToken('sub', arg if arg is not None else ''))
                i = j
            elif i + 1 < len(text):
                tokens.append(MathToken('sub', text[i+1]))
                i += 2
            else:
                i += 1
            continue

        # Superscript: ^{xxx} or ^x
        if c == '^':
            flush_buf()
            if i + 1 < len(text) and text[i+1] == '{':
                arg, j = _balanced_arg(text, i+1)
                tokens.append(MathToken('sup', arg if arg is not None else ''))
                i = j
            elif i + 1 < len(text):
                tokens.append(MathToken('sup', text[i+1]))
                i += 2
            else:
                i += 1
            continue

        buf += c
        i += 1

    flush_buf()
    return tokens


def _convert_matrix_envs(text):
    """Convert pmatrix/bmatrix/matrix environments to bracketed row form:
    \\begin{pmatrix} a & b \\\\ c & d \\end{pmatrix}  →  [ a, b ; c, d ]
    """
    pattern = re.compile(r'\\begin\{(p|b|v|V|B)?matrix\}(.*?)\\end\{\1?matrix\}', re.DOTALL)

    def repl(m):
        kind = m.group(1) or ''
        body = m.group(2)
        # Split rows on \\
        rows = re.split(r'\\\\\s*', body)
        rendered_rows = []
        for row in rows:
            row = row.strip()
            if not row:
                continue
            cells = [c.strip() for c in row.split('&')]
            rendered_rows.append(', '.join(cells))
        joined = ' ; '.join(rendered_rows)
        # Choose brackets by env type
        open_b, close_b = '[', ']'
        if kind == 'p':
            open_b, close_b = '(', ')'
        elif kind == 'b':
            open_b, close_b = '[', ']'
        elif kind in ('v', 'V'):
            open_b, close_b = '|', '|'
        elif kind == 'B':
            open_b, close_b = '{', '}'
        return f' {open_b} {joined} {close_b} '
    return pattern.sub(repl, text)


def render_math_to_paragraph(p, latex_text, base_font='Cambria Math', base_size=11):
    """Render LaTeX math text into a python-docx paragraph p, using Unicode +
    sub/super run formatting + simple a/b fractions."""
    # Convert matrix environments first (before generic begin/end stripping)
    latex_text = _convert_matrix_envs(latex_text)
    # Drop any remaining environments
    latex_text = re.sub(r'\\begin\{[^}]+\}', '', latex_text)
    latex_text = re.sub(r'\\end\{[^}]+\}', '', latex_text)
    # Normalize whitespace
    latex_text = latex_text.replace('\n', ' ').strip()

    tokens = parse_latex_math(latex_text)

    for tok in tokens:
        if tok.kind == 'text':
            r = p.add_run(tok.data)
            r.font.name = base_font
            r.font.size = Pt(base_size)
        elif tok.kind == 'sub':
            inner = tok.data
            uni, ok = _try_unicode_sub(inner)
            if ok:
                r = p.add_run(uni)
                r.font.name = base_font
                r.font.size = Pt(base_size)
            else:
                r = p.add_run(inner)
                r.font.name = base_font
                r.font.size = Pt(base_size)
                r.font.subscript = True
        elif tok.kind == 'sup':
            inner = tok.data
            uni, ok = _try_unicode_sup(inner)
            if ok:
                r = p.add_run(uni)
                r.font.name = base_font
                r.font.size = Pt(base_size)
            else:
                r = p.add_run(inner)
                r.font.name = base_font
                r.font.size = Pt(base_size)
                r.font.superscript = True
        elif tok.kind == 'frac':
            num, den = tok.data
            # Render as "(num) / (den)" — wrap multi-char in parens
            num_str = num.strip()
            den_str = den.strip()
            need_paren_n = len(num_str) > 1 and not (num_str.startswith('(') and num_str.endswith(')'))
            need_paren_d = len(den_str) > 1 and not (den_str.startswith('(') and den_str.endswith(')'))
            # Recursively expand any LaTeX inside
            n_tokens = parse_latex_math(num_str)
            d_tokens = parse_latex_math(den_str)
            if need_paren_n:
                r = p.add_run('(')
                r.font.name = base_font; r.font.size = Pt(base_size)
            for nt in n_tokens:
                _emit_token(p, nt, base_font, base_size)
            if need_paren_n:
                r = p.add_run(')')
                r.font.name = base_font; r.font.size = Pt(base_size)
            r = p.add_run(' / ')
            r.font.name = base_font; r.font.size = Pt(base_size)
            if need_paren_d:
                r = p.add_run('(')
                r.font.name = base_font; r.font.size = Pt(base_size)
            for dt in d_tokens:
                _emit_token(p, dt, base_font, base_size)
            if need_paren_d:
                r = p.add_run(')')
                r.font.name = base_font; r.font.size = Pt(base_size)
        elif tok.kind == 'sqrt':
            r = p.add_run('√')
            r.font.name = base_font; r.font.size = Pt(base_size)
            arg_str = tok.data
            need_paren = len(arg_str.strip()) > 1
            if need_paren:
                r = p.add_run('(')
                r.font.name = base_font; r.font.size = Pt(base_size)
            for at in parse_latex_math(arg_str):
                _emit_token(p, at, base_font, base_size)
            if need_paren:
                r = p.add_run(')')
                r.font.name = base_font; r.font.size = Pt(base_size)


def _emit_token(p, tok, base_font, base_size):
    """Emit a single MathToken into paragraph p."""
    if tok.kind == 'text':
        r = p.add_run(tok.data)
        r.font.name = base_font; r.font.size = Pt(base_size)
    elif tok.kind == 'sub':
        uni, ok = _try_unicode_sub(tok.data)
        if ok:
            r = p.add_run(uni)
        else:
            r = p.add_run(tok.data)
            r.font.subscript = True
        r.font.name = base_font; r.font.size = Pt(base_size)
    elif tok.kind == 'sup':
        uni, ok = _try_unicode_sup(tok.data)
        if ok:
            r = p.add_run(uni)
        else:
            r = p.add_run(tok.data)
            r.font.superscript = True
        r.font.name = base_font; r.font.size = Pt(base_size)
    else:
        # Recurse for frac/sqrt within frac/sqrt
        render_math_to_paragraph(p, _reconstruct(tok), base_font, base_size)


def _reconstruct(tok):
    """Reconstruct LaTeX-like text from a token (for recursion)."""
    if tok.kind == 'frac':
        return r'\frac{' + tok.data[0] + '}{' + tok.data[1] + '}'
    if tok.kind == 'sqrt':
        return r'\sqrt{' + tok.data + '}'
    return tok.data


def has_caption_above(lines, idx):
    """Check if there's already a caption line within 3 lines above idx."""
    for i in range(max(0, idx-3), idx):
        if re.search(r'\*\*표\s+\d', lines[i]):
            return True
    return False


def add_captions_to_md(lines):
    """Add caption lines above each table group that doesn't have one.
    A table 'starts' only when the previous non-empty, non-separator line is NOT a |-row.
    """
    result = []
    table_num = [0]  # mutable counter
    in_code = False
    in_table = False  # track contiguous |-rows

    for i, line in enumerate(lines):
        stripped = line.strip()

        # Track code blocks — don't count |-rows inside code as tables
        if stripped.startswith('```'):
            in_code = not in_code
            in_table = False
            result.append(line)
            continue

        if in_code:
            result.append(line)
            continue

        # Detect table row
        is_row = stripped.startswith('|') and len(stripped) > 1

        if is_row:
            if not in_table:
                # NEW table start
                in_table = True
                already_has = has_caption_above(result, len(result))
                if not already_has:
                    table_num[0] += 1
                    if table_num[0] - 1 < len(TABLE_CAPTIONS):
                        cap = TABLE_CAPTIONS[table_num[0] - 1]
                    else:
                        cap = f"표 {table_num[0]}. (표 설명 — 해당 명령어/기능의 파라미터 또는 옵션 목록)"
                    # Each newline-separated chunk MUST be a separate list item
                    # so the DocxBuilder's line index matches actual file lines.
                    result.append('\n')
                    result.append(f"**{cap}**\n")
                    result.append('\n')
            result.append(line)
        else:
            in_table = False
            result.append(line)

    return result


def insert_meshfix_section(lines):
    """Insert the meshfix section before '## 40. 수학 이론'.
    Split MESHFIX_SECTION into individual lines so the builder's line index
    stays in sync with the actual file lines.
    """
    # Split MESHFIX_SECTION into per-line list items (each ending in '\n')
    meshfix_lines = [ln + '\n' for ln in MESHFIX_SECTION.split('\n')]

    result = []
    inserted = False
    for line in lines:
        if not inserted and line.strip().startswith('## 40. 수학 이론'):
            result.extend(meshfix_lines)
            inserted = True
        result.append(line)

    if not inserted:
        result.extend(meshfix_lines)

    return result


def renumber_sections_40plus(lines):
    """After inserting section 40 (meshfix), renumber old 40→41, 41→42."""
    result = []
    for line in lines:
        # Renumber ## 40. 수학 이론 → ## 41. 수학 이론
        line = re.sub(r'^## 40\. 수학 이론', '## 41. 수학 이론', line)
        # Renumber ## 41. 출력 → ## 42. 출력
        line = re.sub(r'^## 41\. 출력 파일 형식', '## 42. 출력 파일 형식', line)
        # Renumber cross-references
        line = re.sub(r'##41-수학이론', '#41-수학이론', line)
        result.append(line)
    return result


def update_version(lines):
    """Update version string in header and footer."""
    result = []
    for line in lines:
        line = re.sub(r'버전 1\.3\.0', '버전 1.8.0', line)
        line = re.sub(r'v1\.3\.0', 'v1.8.0', line)
        line = re.sub(r'KooRemapper v1\.3\.0', 'KooRemapper v1.8.0', line)
        result.append(line)
    return result


def add_meshfix_to_command_list(lines):
    """Add meshfix to the command list section."""
    result = []
    for i, line in enumerate(lines):
        result.append(line)
        # After the assemble line in command list
        if '  assemble       다중 오퍼레이션 통합 어셈블리' in line:
            # Check if meshfix already there
            if i + 1 < len(lines) and 'meshfix' not in lines[i+1]:
                result.append('  tetremesh      TET4 로컬 재메시 (패치 기반)\n')
                result.append('  meshfix        TET4 파트 전체 재메시 (Gmsh 기반)\n')
    return result


def add_meshfix_to_feature_table(lines):
    """Add TET4 재메시 row to the features table."""
    result = []
    for line in lines:
        result.append(line)
        if '| **출력 관리** | DATABASE 출력 제어 키워드 삽입' in line:
            result.append('| **TET4 재메시** | Gmsh 기반 완전 재메시, 스케일드 자코비안 품질 개선(meshfix) |\n')
    return result


def update_toc(lines):
    """Add meshfix entry to TOC."""
    result = []
    for line in lines:
        result.append(line)
        if '    - 39.16 [wrap](#3916-wrap--와인딩-인장-프리스트레스)' in line:
            result.append('40. [meshfix — TET4 재메시 (Gmsh 기반)](#40-meshfix--tet4-재메시-gmsh-기반)\n')
    return result


# ─────────────────────────────────────────────────────────────────────────────
# DOCX generation
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_bookmark(para, name):
    """Add a Word bookmark to a paragraph."""
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(hash(name) % 100000))
    bm_start.set(qn('w:name'), name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(hash(name) % 100000))
    para._p.insert(0, bm_start)
    para._p.append(bm_end)


class DocxBuilder:
    def __init__(self, md_lines):
        self.doc = Document()
        self.lines = md_lines
        self.fig_num = 0
        self.tbl_num_global = 0
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles."""
        doc = self.doc

        # Normal style: Malgun Gothic 10pt
        style = doc.styles['Normal']
        style.font.name = '맑은 고딕'
        style.font.size = Pt(10)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        # Heading 1
        h1 = doc.styles['Heading 1']
        h1.font.name = '맑은 고딕'
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        h1._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        # Heading 2
        h2 = doc.styles['Heading 2']
        h2.font.name = '맑은 고딕'
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        h2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        # Heading 3
        h3 = doc.styles['Heading 3']
        h3.font.name = '맑은 고딕'
        h3.font.size = Pt(11)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(0x37, 0x86, 0x9B)
        h3._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        # Caption style
        if 'Caption' not in [s.name for s in doc.styles]:
            cap_style = doc.styles.add_style('Caption', 2)  # 2 = paragraph
        else:
            cap_style = doc.styles['Caption']
        cap_style.font.name = '맑은 고딕'
        cap_style.font.size = Pt(9)
        cap_style.font.italic = True
        cap_style.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        cap_style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        # Code style (paragraph style, type=1)
        from docx.enum.style import WD_STYLE_TYPE
        if 'KooCode' not in [s.name for s in doc.styles]:
            code_style = doc.styles.add_style('KooCode', WD_STYLE_TYPE.PARAGRAPH)
        else:
            code_style = doc.styles['KooCode']
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(8.5)
        code_style.font.color.rgb = RGBColor(0x10, 0x10, 0x10)
        pf = code_style.paragraph_format
        pf.left_indent = Cm(0.5)
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)

    def _add_heading(self, text, level):
        # MD heading levels → Word heading styles
        # # (1) = Document title
        # ## (2) = Heading 1 (main sections like "1. 개요")
        # ### (3) = Heading 2
        # #### (4) = Heading 3
        if level == 1:
            p = self.doc.add_paragraph(text, style='Title')
        elif level == 2:
            p = self.doc.add_paragraph(text, style='Heading 1')
        elif level == 3:
            p = self.doc.add_paragraph(text, style='Heading 2')
        else:
            p = self.doc.add_paragraph(text, style='Heading 3')
        return p

    def _add_paragraph(self, text):
        """Add paragraph with inline markdown formatting."""
        if not text.strip():
            return
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['Normal']
        self._apply_inline(p, text)

    def _apply_inline(self, para, text):
        """Parse inline markdown (**bold**, *italic*, `code`) and add runs."""
        # Split by inline patterns
        pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\$[^$]+\$)')
        parts = pattern.split(text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
                run.font.name = '맑은 고딕'
                run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = para.add_run(part[1:-1])
                run.italic = True
                run.font.name = '맑은 고딕'
                run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            elif part.startswith('`') and part.endswith('`'):
                run = para.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            elif part.startswith('$') and part.endswith('$'):
                # Inline LaTeX: render via Unicode + sub/sup formatting
                inner = part[1:-1]
                render_math_to_paragraph(para, inner, base_size=10)
            else:
                # Strip trailing HTML comments
                clean = re.sub(r'<!--.*?-->', '', part)
                run = para.add_run(clean)
                run.font.name = '맑은 고딕'
                run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    def _add_caption(self, text, is_table=True):
        """Add caption paragraph with Caption style."""
        clean = re.sub(r'\*\*|\*', '', text).strip()
        p = self.doc.add_paragraph(clean, style='Caption')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        return p

    def _add_code_block(self, lines_block):
        """Add code block as formatted paragraphs."""
        text = '\n'.join(lines_block)
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['KooCode']
        # Light gray background
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'F5F5F5')
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        return p

    def _add_table(self, table_lines, caption_text=None):
        """Add Word table from markdown table lines."""
        # Parse header and rows
        rows = []
        is_header_sep = re.compile(r'^\|[-| :]+\|$')

        for tl in table_lines:
            if not tl.strip():
                continue
            if is_header_sep.match(tl.strip()):
                continue
            cells = [c.strip() for c in tl.strip().strip('|').split('|')]
            rows.append(cells)

        if not rows:
            return

        # Caption ABOVE the table
        if caption_text:
            self._add_caption(caption_text, is_table=True)

        max_cols = max(len(r) for r in rows)

        tbl = self.doc.add_table(rows=len(rows), cols=max_cols)
        tbl.style = 'Table Grid'

        for ri, row in enumerate(rows):
            for ci, cell_text in enumerate(row):
                if ci >= max_cols:
                    break
                cell = tbl.rows[ri].cells[ci]
                # Strip markdown bold
                clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                clean = re.sub(r'\*([^*]+)\*', r'\1', clean)
                clean = re.sub(r'`([^`]+)`', r'\1', clean)
                p = cell.paragraphs[0]
                p.clear()
                p.style = self.doc.styles['Normal']
                if ri == 0:
                    run = p.add_run(clean)
                    run.bold = True
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(9)
                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    set_cell_bg(cell, 'D6E4F0')
                else:
                    run = p.add_run(clean)
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(9)
                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        return tbl

    def build(self):
        """Parse MD lines and build DOCX."""
        lines = self.lines
        i = 0
        total = len(lines)

        # State
        in_code = False
        code_lang = ''
        code_lines = []
        table_lines = []
        in_table = False
        pending_caption = None  # caption found ABOVE the table

        # Caption counter for auto-detection
        tbl_auto_num = [0]

        while i < total:
            raw = lines[i]
            line = raw.rstrip('\n')

            # ── Code block ──────────────────────────────────────────────────
            if line.startswith('```'):
                if not in_code:
                    in_code = True
                    code_lang = line[3:].strip()
                    code_lines = []
                else:
                    in_code = False
                    self._add_code_block(code_lines)
                    code_lines = []
                i += 1
                continue

            if in_code:
                code_lines.append(line)
                i += 1
                continue

            # ── Table detection ──────────────────────────────────────────────
            if line.strip().startswith('|') and '|' in line[1:]:
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(line)
                i += 1
                continue
            else:
                if in_table:
                    cap = pending_caption
                    pending_caption = None
                    if not cap:
                        tbl_auto_num[0] += 1
                    self._add_table(table_lines, caption_text=cap)
                    self.tbl_num_global += 1
                    in_table = False
                    table_lines = []

            # ── Skip horizontal rules ────────────────────────────────────────
            if re.match(r'^-{3,}$', line.strip()):
                i += 1
                continue

            # ── Headings ─────────────────────────────────────────────────────
            m = re.match(r'^(#{1,4})\s+(.+)', line)
            if m:
                level = len(m.group(1))
                text = m.group(2)
                # Remove markdown links from heading [text](#anchor)
                text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
                self._add_heading(text, level)
                i += 1
                continue

            # ── Caption line (above table) — **표 N. ...** ───────────────────
            cap_m = re.match(r'^\*\*(표\s+[\d\-]+\..+)\*\*\s*$', line.strip())
            if cap_m:
                pending_caption = cap_m.group(1)
                i += 1
                continue

            # ── Figure caption line — > **그림 N. ...** ──────────────────────
            fig_m = re.match(r'^>\s*\*\*(그림\s+[\d\-]+\..+)\*\*', line.strip())
            if fig_m:
                self._add_caption(fig_m.group(1), is_table=False)
                i += 1
                continue

            # ── Blockquote (non-caption) ─────────────────────────────────────
            bq_m = re.match(r'^>\s*(.*)', line)
            if bq_m:
                p = self.doc.add_paragraph(style='Normal')
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                self._apply_inline(p, bq_m.group(1))
                i += 1
                continue

            # ── Math ($$) — single-line OR multi-line block ──────────────────
            stripped = line.strip()
            if stripped.startswith('$$'):
                # Single-line math: $$ ... $$ on one line
                if stripped.endswith('$$') and len(stripped) > 4:
                    math_text = stripped[2:-2].strip()
                    p = self.doc.add_paragraph()
                    p.style = self.doc.styles['Normal']
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    render_math_to_paragraph(p, math_text, base_size=11)
                    i += 1
                    continue
                else:
                    # Multi-line math block: collect until closing $$
                    math_lines = []
                    i += 1
                    while i < total and not lines[i].strip().startswith('$$'):
                        math_lines.append(lines[i].rstrip('\n'))
                        i += 1
                    i += 1  # skip closing $$
                    math_text = ' '.join(math_lines).strip()
                    p = self.doc.add_paragraph()
                    p.style = self.doc.styles['Normal']
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    render_math_to_paragraph(p, math_text, base_size=11)
                    continue

            # ── List items ────────────────────────────────────────────────────
            li_m = re.match(r'^(\s*)[*\-]\s+(.+)', line)
            if li_m:
                indent = len(li_m.group(1))
                text = li_m.group(2)
                p = self.doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Cm(0.5 + indent * 0.25)
                self._apply_inline(p, text)
                i += 1
                continue

            # ── Numbered list ─────────────────────────────────────────────────
            nl_m = re.match(r'^(\s*)\d+\.\s+(.+)', line)
            if nl_m:
                indent = len(nl_m.group(1))
                text = nl_m.group(2)
                p = self.doc.add_paragraph(style='List Number')
                p.paragraph_format.left_indent = Cm(0.5 + indent * 0.25)
                self._apply_inline(p, text)
                i += 1
                continue

            # ── Empty line ────────────────────────────────────────────────────
            if not line.strip():
                i += 1
                continue

            # ── Regular paragraph ─────────────────────────────────────────────
            self._add_paragraph(line)
            i += 1

        # Flush table if file ends in table
        if in_table and table_lines:
            self._add_table(table_lines, caption_text=pending_caption)

    def save(self, path):
        self.doc.save(path)
        print(f"  → DOCX saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    base = os.path.dirname(os.path.abspath(__file__))
    md_path   = os.path.join(base, 'KooRemapper_Manual.md')
    docx_path = os.path.join(base, 'KooRemapper_Manual.docx')

    print("Reading MD...")
    lines = read_file(md_path)
    print(f"  {len(lines)} lines")

    print("Updating version...")
    lines = update_version(lines)

    print("Adding meshfix to command list...")
    lines = add_meshfix_to_command_list(lines)

    print("Adding meshfix to feature table...")
    lines = add_meshfix_to_feature_table(lines)

    print("Inserting meshfix section...")
    lines = insert_meshfix_section(lines)

    print("Renumbering sections 40/41...")
    lines = renumber_sections_40plus(lines)

    print("Adding table captions...")
    lines = add_captions_to_md(lines)

    print("Writing updated MD...")
    write_file(md_path, lines)
    print(f"  → MD saved: {md_path} ({len(lines)} lines)")

    print("Building DOCX...")
    builder = DocxBuilder(lines)
    builder.build()
    builder.save(docx_path)
    print("Done.")


if __name__ == '__main__':
    main()
