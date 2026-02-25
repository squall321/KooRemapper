"""
KooRemapper 기능 설명서 DOCX 생성기
python-docx 1.x 사용

실행: python docs/generate_manual.py
출력: docs/KooRemapper_Manual.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

# ─────────────────────────────────────────────
# 유틸리티
# ─────────────────────────────────────────────

def set_run_font(run, name='맑은 고딕', size=None, bold=None, italic=None, color=None):
    run.font.name = name
    if size:   run.font.size = Pt(size)
    if bold is not None:   run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color:  run.font.color.rgb = RGBColor(*color)
    # 동아시아 폰트 설정
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = '맑은 고딕'
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        rPr.insert(0, rFonts)
    return p

def add_para(doc, text='', style='Normal', bold=False, italic=False,
             size=10.5, color=None, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if not text:
        return p
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_formula(doc, text):
    """수식을 이탤릭체 단락으로 삽입 (유니코드 수학 기호 사용)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_font(run, name='Cambria Math', size=11, italic=True, color=(0x00, 0x00, 0x80))
    return p

def add_code_block(doc, code_text):
    """고정폭 코드 블록"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else ' ')
        set_run_font(run, name='Consolas', size=9.5, color=(0x1e, 0x3a, 0x5f))
        # 배경색 (연한 회색)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0')
        run._r.get_or_add_rPr().append(shd)

def add_table(doc, headers, rows, col_widths=None):
    """표 생성"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # 헤더
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_run_font(cell.paragraphs[0].runs[0], size=10, bold=True)
        # 배경 회색
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)
    # 데이터
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            if cell.paragraphs[0].runs:
                set_run_font(cell.paragraphs[0].runs[0], size=9.5)
    # 열 너비
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()  # 표 뒤 여백

def add_note(doc, text):
    """주의/참고 박스"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('📌 ' + text)
    set_run_font(run, size=10, color=(0x7b, 0x36, 0x06))

def add_sep(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ─────────────────────────────────────────────
# 문서 생성 시작
# ─────────────────────────────────────────────

doc = Document()

# 여백 설정
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# 기본 Normal 스타일
normal = doc.styles['Normal']
normal.font.name = '맑은 고딕'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)

# ─────────────────────────────────────────────
# 표지
# ─────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('KooRemapper')
set_run_font(r, name='맑은 고딕', size=32, bold=True, color=(0x1f, 0x49, 0x7d))

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('기능 설명서')
set_run_font(r2, size=20, color=(0x2e, 0x75, 0xb6))

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('LS-DYNA 메시 전처리 도구  |  Version 1.1.0')
set_run_font(r3, size=12, color=(0x60, 0x60, 0x60))

doc.add_page_break()

# ─────────────────────────────────────────────
# 1. 개요
# ─────────────────────────────────────────────
add_heading(doc, '1. 개요', 1)
add_para(doc,
    'KooRemapper는 LS-DYNA FEA 해석을 위한 메시 전처리 도구입니다. '
    '개략 메시(coarse mesh)로 구성된 전체 모델에 상세 메시(detail mesh)를 매핑하고, '
    '조립 공정에 수반되는 초기 응력 상태(prestress)를 재현하는 것이 주목적입니다.')

add_heading(doc, '1.1 핵심 기능 범위', 2)
add_table(doc,
    ['범주', '기능'],
    [
        ['메시 매핑', 'HEX8 등매개변수 매핑 (map), QUAD4 셸 기반 매핑 (shellmap)'],
        ['초기 응력', '기준-변형 형상 간 변형률/응력 계산, dynain 출력 (prestress)'],
        ['간섭 조립', '부품 압축(squeeze) + 역방향 prestress'],
        ['형상 변형', '굽힘(bend), 압입/엠보싱(indent) + 굽힘 초기 응력'],
        ['적층 구조', '레이어별 두께·재료 재정의 (restack)'],
        ['성형 변형', '이면각 기반 소성 변형률 (formstrain)'],
        ['메시 변환', '2차 요소 변환(tet10/hex20/quad8/tria6), 세분화(refine), ELFORM 변경(elform)'],
        ['토폴로지', '노드 분리 (disconnect): full / CZM / MEFEM 모드'],
        ['등기하해석', 'FE solid → IGA NURBS box 래핑 (iga)'],
        ['메시 생성', '변밀도 메시 생성 (generate-var)'],
    ],
    col_widths=[4, 12]
)

add_heading(doc, '1.2 시스템 요구사항 및 빌드', 2)
add_table(doc,
    ['항목', '요구사항'],
    [
        ['운영체제', 'Windows 10/11 x64'],
        ['컴파일러', 'MSVC 2019/2022 (Visual Studio)'],
        ['CMake', '3.16 이상'],
        ['C++ 표준', 'C++17'],
    ],
    col_widths=[4, 12]
)
add_code_block(doc, '''
cmake -B build -G "Visual Studio 17 2022" -A x64
cmake --build build --config Release

# 실행 파일: build\\bin\\Release\\KooRemapper.exe
''')

# ─────────────────────────────────────────────
# 2. 명령어 목록
# ─────────────────────────────────────────────
add_heading(doc, '2. 명령어 목록', 1)
add_code_block(doc, '''
KooRemapper.exe <command> [options] ...

  map            HEX8 구조화 메시를 굽힘 참조 형상에 매핑
  shellmap       QUAD4 셸 참조 기반 상세 메시 매핑
  unfold         굽힘 구조화 메시로부터 평면 메시 생성
  generate       테스트용 예제 메시 생성
  generate-var   YAML 설정 기반 변밀도 메시 생성
  strain         두 메시 간 변형률 계산 (CSV 출력)
  prestress      변형 형상 기반 초기 응력 계산 + dynain 출력
  squeeze        간섭 끼워맞춤 초기 변형 계산
  assemble       다중 오퍼레이션 통합 어셈블리
  info           메시 파일 정보 출력
  help           도움말
  version        버전 정보
''')

# ─────────────────────────────────────────────
# 3. map
# ─────────────────────────────────────────────
add_heading(doc, '3. map — HEX8 구조화 메시 매핑', 1)
add_para(doc, '평면(flat) 상세 메시(HEX8)를 굽힘(bent) 참조 구조화 메시에 등매개변수 방법으로 매핑합니다.')
add_code_block(doc, 'KooRemapper.exe map [--single] <bent_mesh.k> <flat_mesh.k> <output.k>\n\n  --single, -s    단일 스레드 모드 (기본: 병렬)')

add_heading(doc, '3.1 동작 원리', 2)
add_para(doc, '각 상세 메시 노드 p에 대해 다음 과정을 수행합니다:')
add_para(doc, '① 참조 메시에서 포함하는 HEX8 요소 검색')
add_para(doc, '② 자연 좌표 (ξ, η, ζ) 역계산 (Newton-Raphson)')
add_para(doc, '③ 굽힘 참조 형상의 동일 자연 좌표로 위치 변환')

add_heading(doc, '3.2 수식', 2)
add_para(doc, 'HEX8 형상 함수 (자연 좌표 ξ, η, ζ ∈ [-1, 1]³):')
add_formula(doc, 'x(ξ,η,ζ) = Σᵢ Nᵢ(ξ,η,ζ) · xᵢ    (i = 1…8)')
add_formula(doc, 'Nᵢ = ⅛ (1 + ξᵢξ)(1 + ηᵢη)(1 + ζᵢζ)')
add_para(doc, '역매핑 (Newton-Raphson 반복):')
add_formula(doc, '(Δξ, Δη, Δζ)ᵀ = J⁻¹ · (x_target − x(ξ,η,ζ))')
add_para(doc, '야코비안:  Jᵢⱼ = ∂xᵢ/∂ξⱼ = Σₖ (∂Nₖ/∂ξⱼ) xₖᵢ')

add_heading(doc, '3.3 주의사항', 2)
add_note(doc, '참조 메시는 반드시 규칙적인 HEX8 구조화 메시이어야 합니다.')
add_note(doc, '상세 메시 노드가 참조 요소 외부에 있으면 경고가 출력됩니다.')

# ─────────────────────────────────────────────
# 4. shellmap
# ─────────────────────────────────────────────
add_heading(doc, '4. shellmap — QUAD4 셸 기반 매핑', 1)
add_para(doc, 'QUAD4 셸 참조 형상을 기반으로 평면 상세 고체 메시(solid detail mesh)를 굽힘 형상으로 매핑합니다.')
add_code_block(doc, 'KooRemapper.exe shellmap [--thickness <t>] <bent_shell.k> <flat_detail.k> <output.k>\n\n  --thickness <t>    두께 명시적 지정 (기본: Z-범위 자동 감지)')

add_heading(doc, '4.1 수식', 2)
add_para(doc, '셸 법선 n̂(u,v)와 두께 t에 대해:')
add_formula(doc, "x' = x_shell(u,v) + (z / (t/2)) · (t/2) · n̂(u,v)")
add_para(doc, '여기서 z는 평면 메시에서의 두께 방향 상대 위치.')

add_note(doc, '가전개(developable) 면에 최적화되어 있습니다. 비가전개 면에서는 왜곡 경고가 출력됩니다.')
add_note(doc, 'QUAD4 전용 (TRIA3 미지원).')

# ─────────────────────────────────────────────
# 5. prestress
# ─────────────────────────────────────────────
add_heading(doc, '5. prestress — 초기 응력/변형률 계산', 1)
add_para(doc, '기준 형상(reference)과 변형 형상(deformed) 메시 쌍으로부터 각 요소의 초기 응력을 계산하여 LS-DYNA *INITIAL_STRESS_SOLID 형식으로 출력합니다.')
add_code_block(doc, '''KooRemapper.exe prestress [options] <ref.k> <def.k> <output_prefix>

  --E <value>                  영률 (K-파일 재료 카드 대체)
  --nu <value>                 푸아송 비
  --strain engineering|green|log   변형률 계산 방식 (기본: engineering)
  --csv                        CSV 형식 추가 출력''')

add_heading(doc, '5.1 변형률 계산', 2)
add_para(doc, '공학 변형률 (Engineering Strain):', bold=True)
add_formula(doc, 'εᵢⱼ = ½ (∂uᵢ/∂xⱼ + ∂uⱼ/∂xᵢ)')
add_para(doc, 'Green-Lagrange 변형률:', bold=True)
add_formula(doc, 'Eᵢⱼ = ½ (∂uᵢ/∂Xⱼ + ∂uⱼ/∂Xᵢ + ∂uₖ/∂Xᵢ · ∂uₖ/∂Xⱼ)')
add_para(doc, '로그 변형률 (Logarithmic / True Strain):', bold=True)
add_formula(doc, 'ε_log = ln(L / L₀)')

add_heading(doc, '5.2 응력 계산 (선형 탄성)', 2)
add_para(doc, '라메 상수:')
add_formula(doc, 'λ = Eν / [(1+ν)(1−2ν)],    μ = E / [2(1+ν)]')
add_para(doc, 'Cauchy 응력 (3D 이방성 Hooke 법칙):')
add_formula(doc, 'σᵢⱼ = λ εₖₖ δᵢⱼ + 2μ εᵢⱼ')
add_para(doc, '재료 우선순위: 명령행 --E/--nu  >  K-파일 *MAT_ELASTIC (파트별)')

# ─────────────────────────────────────────────
# 6. squeeze (standalone)
# ─────────────────────────────────────────────
add_heading(doc, '6. squeeze — 간섭 끼워맞춤', 1)
add_para(doc, '간섭(interference fit) 조립 시뮬레이션을 위해 대상 파트를 지정 변형률로 압축하고, 그 역방향 응력을 dynain으로 출력합니다.')
add_code_block(doc, 'KooRemapper.exe squeeze <mesh.k> <config.yaml> <output_prefix>')

add_heading(doc, '6.1 YAML 설정', 2)
add_code_block(doc, '''
parts:
  - pid: 3
    eps_x: -0.02    # x방향 2% 압축 (음수=압축)
    eps_y: -0.02
    eps_z:  0.0

material:           # 전역 재료 (K-파일 재료 없을 때)
  E: 210000
  nu: 0.3
''')

add_heading(doc, '6.2 수식', 2)
add_para(doc, '파트 바운딩 박스 중심 c에 대해 각 노드 위치:')
add_formula(doc, "x' = c + diag(1+εₓ, 1+εᵧ, 1+ε_z) · (x − c)")
add_para(doc, '초기 응력 (압축에 대한 역방향):')
add_formula(doc, 'σₓₓ = −[(λ+2μ)εₓ + λ(εᵧ+ε_z)]')

# ─────────────────────────────────────────────
# 7. generate / generate-var
# ─────────────────────────────────────────────
add_heading(doc, '7. generate / generate-var — 메시 생성', 1)
add_para(doc, 'generate: 테스트용 기하학 메시 생성', bold=True)
add_code_block(doc, 'KooRemapper.exe generate <type> [options] <output.k>\n\nTypes: teardrop, arc, scurve, helix, torus, twist, wave, bulge, taper, waterdrop')

add_para(doc, 'generate-var: YAML 기반 변밀도 메시 생성', bold=True)
add_code_block(doc, 'KooRemapper.exe generate-var [--ref <ref.k>] [--no-scale] <config.yaml> <output.k>')
add_heading(doc, '7.1 YAML 설정 예 (평면 타입)', 2)
add_code_block(doc, '''
type: flat
zones:
  - id: 1
    nx: 10
    ny: 8
    nz: 2
    x_min: 0.0
    x_max: 50.0
    y_min: 0.0
    y_max: 40.0
''')

# ─────────────────────────────────────────────
# 8. assemble — 개요
# ─────────────────────────────────────────────
add_heading(doc, '8. assemble — 통합 어셈블리 명령', 1)
add_para(doc,
    '여러 오퍼레이션을 순차적으로 적용하는 통합 명령입니다. '
    '기본 모델을 로드하고 각 오퍼레이션을 순서대로 실행하며, '
    '누적된 초기 응력을 단일 dynain 파일로 출력합니다.')
add_code_block(doc, 'KooRemapper.exe assemble <config.yaml>')

add_heading(doc, '8.1 공통 YAML 구조', 2)
add_code_block(doc, '''
base_model: model.k             # 기본 모델 (필수)
output: result                  # 출력 접두어 (필수)

material:                       # 전역 재료 (선택)
  E: 210000
  nu: 0.3

dynamic_relaxation: true        # *CONTROL_DYNAMIC_RELAXATION 자동 삽입
dynain_embed: false             # true: dynain을 메인 파일에 인라인 삽입

operations:
  - type: <op_type>
    ...
''')

add_heading(doc, '8.2 공통 특성', 2)
add_table(doc,
    ['특성', '설명'],
    [
        ['원본 키워드 보존', '*CONTACT, *BOUNDARY, *LOAD 등 미파싱 키워드 그대로 유지'],
        ['응력 누적', '동일 요소에 여러 오퍼레이션 적용 시 응력 합산'],
        ['ID 자동 관리', '파트/섹션/노드/요소 ID 자동 발급 (충돌 방지)'],
        ['동적 완화', 'dynamic_relaxation: true → *CONTROL_DYNAMIC_RELAXATION 자동 삽입'],
    ],
    col_widths=[5, 11]
)

# ─────────────────────────────────────────────
# 8.3 replace
# ─────────────────────────────────────────────
add_heading(doc, '8.3 replace — 상세 메시 교체', 2)
add_para(doc, '모델 내 특정 파트(coarse)를 상세 메시(detail)로 교체합니다. 선택적으로 굽힘 초기 응력(prestress)을 자동 계산합니다.')
add_code_block(doc, '''
- type: replace
  target_pid: 3           # 교체 대상 파트 ID
  detail_flat: detail.k   # 평면 상세 메시
  shell_bent: bent.k      # 굽힘 QUAD4 셸 참조 (필수)
  prestress: true         # 굽힘 초기 응력 계산 (기본: false)
''')
add_para(doc, '동작 순서:', bold=True)
add_para(doc, '① target_pid 파트 제거  → ② shellmap으로 상세 메시 매핑  → ③ 새 ID 발급  → ④ prestress: true 시 dynain 축적')
add_para(doc, 'ID 재번호화:')
add_formula(doc, 'new_node_id = old_node_id + max_node_id')
add_formula(doc, 'new_elem_id = old_elem_id + max_elem_id')

# ─────────────────────────────────────────────
# 8.4 squeeze (assemble 내)
# ─────────────────────────────────────────────
add_heading(doc, '8.4 squeeze (assemble 내)', 2)
add_code_block(doc, '''
- type: squeeze
  target_pid: 5
  eps_x: -0.015
  eps_y: -0.015
  eps_z:  0.0
''')
add_para(doc, '독립형 squeeze 명령과 동일 알고리즘. assemble 파이프라인 내에서 다른 오퍼레이션과 순차 결합 가능합니다.')

# ─────────────────────────────────────────────
# 8.5 restack
# ─────────────────────────────────────────────
add_heading(doc, '8.5 restack — 레이어 재적층', 2)
add_para(doc, '기존 파트를 두께 방향으로 제거하고, 각기 다른 두께와 재료를 가진 레이어 스택으로 재생성합니다.')
add_code_block(doc, '''
- type: restack
  target_pid: 2
  direction: z          # auto | x | y | z
  element_type: solid   # solid | tshell | shell
  layers:
    - thickness: 0.3
      material_card: |
        *MAT_ELASTIC
        $#     mid        ro         e        pr
          MID001  7.85E-09    210000       0.3
    - thickness: 0.5
      material_card: |
        *MAT_ELASTIC
        $#     mid        ro         e        pr
          MID001  2.50E-09     70000       0.33
''')
add_note(doc, 'MID001, MID002 등의 플레이스홀더가 자동으로 실제 MID로 치환됩니다.')

# ─────────────────────────────────────────────
# 8.6 bend
# ─────────────────────────────────────────────
add_heading(doc, '8.6 bend — 굽힘 변형 + 초기 응력', 2)
add_para(doc, '처짐 함수 w(x₁, x₂)로 기술되는 굽힘을 파트에 적용합니다. deform(노드 이동) 또는 stress(응력만) 모드 선택 가능.')
add_code_block(doc, '''
- type: bend
  target_pid: 1
  plane: xy           # xy | yz | zx
  mode: deform        # deform | stress
  source: formula     # formula | dat | dat_pair

  # source: formula
  expression: "0.5 * sin(pi * x1 / L1) * sin(pi * x2 / L2)"

  # source: dat (2D 격자 파일)
  # dat_file: deflection.dat

  # source: dat_pair (상/하면 별도 파일)
  # dat_top: top.dat
  # dat_bottom: bottom.dat
''')

add_heading(doc, '8.6.1 수식 변수', 3)
add_table(doc,
    ['변수', '의미'],
    [
        ['x1', '면내 좌표 1 (바운딩 박스 최솟값 기준 상대값)'],
        ['x2', '면내 좌표 2'],
        ['L1', 'x1 방향 바운딩 박스 길이'],
        ['L2', 'x2 방향 바운딩 박스 길이'],
        ['pi', '원주율 π = 3.14159…'],
    ],
    col_widths=[3, 13]
)
add_para(doc, '지원 함수: sin, cos, tan, sqrt, exp, log, abs, pow')

add_heading(doc, '8.6.2 굽힘 이론 (Kirchhoff 판이론)', 3)
add_para(doc, '처짐 함수 w(x₁, x₂)로부터 곡률:')
add_formula(doc, 'κ₁₁ = −∂²w/∂x₁²,    κ₂₂ = −∂²w/∂x₂²,    κ₁₂ = −∂²w/∂x₁∂x₂')
add_para(doc, '수치 미분 (dat 소스, 유한 차분):')
add_formula(doc, 'κ ≈ −[w(x+h) − 2w(x) + w(x−h)] / h²')
add_para(doc, '중립면에서 거리 d인 지점의 굽힘 변형률:')
add_formula(doc, 'ε₁₁ = d·κ₁₁,    ε₂₂ = d·κ₂₂,    ε₁₂ = d·κ₁₂')
add_para(doc, '평면응력 가정 하 응력:')
add_formula(doc, 'σ₁₁ = E/(1−ν²) · (ε₁₁ + ν·ε₂₂)')
add_formula(doc, 'σ₂₂ = E/(1−ν²) · (ε₂₂ + ν·ε₁₁)')
add_formula(doc, 'σ₁₂ = E/(1+ν) · ε₁₂')
add_note(doc, '응력은 노드 변위 적용 전에 계산합니다 (중립면 위치 보존).')

# ─────────────────────────────────────────────
# 8.7 indent
# ─────────────────────────────────────────────
add_heading(doc, '8.7 indent — 압입 / 엠보싱', 2)
add_para(doc, '폐곡선 경계(다각형 또는 스플라인) 안쪽 영역에 quarter-arc 필렛 프로파일로 압입(depth>0) 또는 엠보싱(depth<0)을 적용합니다.')
add_code_block(doc, '''
- type: indent
  target_pid: 2
  plane: xy
  direction: -z
  depth: 2.0          # 양수=압입, 음수=엠보싱
  r1: 1.5             # 펀치 측 필렛 반경
  r2: 1.0             # 다이 측 필렛 반경
  bottom_ratio: 0.5   # 두께 방향 관통 비율 (0~1)
  stress: true        # 굽힘 응력 계산 여부
  shape:
    type: polygon     # polygon | spline
    points:
      - [0.0, 0.0]
      - [10.0, 0.0]
      - [10.0, 8.0]
      - [0.0, 8.0]
''')

add_heading(doc, '8.7.1 부호 있는 거리 함수 (SDF)', 3)
add_para(doc, '경계 다각형에 대한 SDF는 Winding Number 방법으로 계산합니다:')
add_formula(doc, 'w(p) = (1/2π) ∮_C dθ')
add_para(doc, '내부이면 w ≠ 0 (양수), 외부이면 w = 0 (음수). SDF = ±최근접 경계까지의 거리.')

add_heading(doc, '8.7.2 압입 프로파일 h(d)', 3)
add_para(doc, 'k = depth / (r₁ + r₂)')
add_para(doc, 'r₁ 구역  (0 ≤ d ≤ r₁, 펀치 측 필렛):', bold=True)
add_formula(doc, 'h(d) = −depth + k·r₁·[1 − √(1 − (d/r₁)²)]')
add_formula(doc, "h''(d) = (k/r₁) / [1 − (d/r₁)²]^(3/2)")
add_para(doc, '평탄 구역  (r₁ < d ≤ D − r₂):', bold=True)
add_formula(doc, "h(d) = −depth,    h''(d) = 0")
add_para(doc, 'r₂ 구역  (D − r₂ < d ≤ D, 다이 측 필렛):', bold=True)
add_formula(doc, 'h(d) = depth · k·r₂ · [1 − √(1 − ((D−d)/r₂)²)] − depth')
add_note(doc, '엠보싱(depth < 0): 위 수식의 정확한 수학적 거울상 (부호 반전).')
add_note(doc, "h''(d)의 특이점 (d = r₁ 부근 발산) → strainLimit/(thickness/2)으로 상한 제한.")

add_heading(doc, '8.7.3 굽힘 응력 분해 (stress: true)', 3)
add_para(doc, '부호 있는 거리 d에서 경계까지의 구배 방향 (gₓ, gᵧ):')
add_formula(doc, "κₓ = −h''(d)·gₓ²,    κᵧ = −h''(d)·gᵧ²,    κₓᵧ = −h''(d)·gₓ·gᵧ")

# ─────────────────────────────────────────────
# 8.8 formstrain
# ─────────────────────────────────────────────
add_heading(doc, '8.8 formstrain — 성형 소성 변형률', 2)
add_para(doc, '셸 메시의 이면각(dihedral angle)으로부터 굽힘 곡률을 계산하여 등가 소성 변형률(EPS)을 *INITIAL_STRESS_SHELL로 출력합니다.')
add_code_block(doc, '''
- type: formstrain
  target_pid: 0          # 0 = 전체 셸 파트 자동 감지
  shell_thickness: 0.0   # 0 = *SECTION_SHELL에서 자동
  min_curvature: 0.001   # 잡음 필터 임계값
''')
add_heading(doc, '8.8.1 이론', 3)
add_para(doc, '인접 셸 요소 쌍의 이면각 θ, 중심 간 거리 L:')
add_formula(doc, 'κ = θ / L')
add_para(doc, '두께 t인 셸의 표면 굽힘 변형률:')
add_formula(doc, 'ε = (t/2) · κ')
add_para(doc, '등가 소성 변형률 (Von Mises 기준):')
add_formula(doc, 'EPS = (2/√3) · ε = t·θ / (√3·L)')
add_note(doc, '동일 요소에 복수 이웃 곡률이 있을 경우 최대값(max) 적용 (합산 아님).')

# ─────────────────────────────────────────────
# 8.9 tet10/hex20/quad8/tria6
# ─────────────────────────────────────────────
add_heading(doc, '8.9 tet10 / hex20 / quad8 / tria6 — 2차 요소 변환', 2)
add_para(doc, '1차 요소(TET4, HEX8, QUAD4, TRIA3)를 2차 요소로 변환합니다. 각 요소 엣지의 중간점 노드를 자동 생성합니다.')
add_code_block(doc, '''
- type: tet10          # tet10 | hex20 | quad8 | tria6
  target_pid: 0        # 0 = 전체 파트
  elform: 17           # ELFORM 지정 (0=자동)
''')
add_table(doc,
    ['convertType', '원본 요소', '변환 요소', '기본 ELFORM'],
    [
        ['tet10', 'TET4', 'TET10', '17'],
        ['hex20', 'HEX8', 'HEX20', '23'],
        ['quad8', 'QUAD4', 'QUAD8', '23'],
        ['tria6', 'TRIA3', 'TRIA6', '24'],
    ],
    col_widths=[4, 4, 4, 4]
)
add_para(doc, '중간점 노드:')
add_formula(doc, 'x_mid = (x_n₁ + x_n₂) / 2')
add_note(doc, '공유 엣지의 중간점은 단일 노드로 중복 제거됩니다 (edgeMidNodeMap_ 캐싱).')

# ─────────────────────────────────────────────
# 8.10 refine
# ─────────────────────────────────────────────
add_heading(doc, '8.10 refine — 메시 세분화', 2)
add_para(doc, '요소를 엣지 방향으로 1:2 또는 1:3 비율로 균일 세분화합니다.')
add_code_block(doc, '''
- type: refine
  target_pid: 0    # 0 = 전체
  ratio: 2         # 2 또는 3
''')
add_table(doc,
    ['요소 유형', 'ratio=2 결과', 'ratio=3 결과'],
    [
        ['QUAD4', '4개 서브 쿼드', '9개 서브 쿼드'],
        ['TRIA3', '4개 서브 삼각형', '9개 서브 삼각형'],
        ['HEX8', '8개 서브 헥스', '27개 서브 헥스'],
        ['TET4', '8개 서브 테트 (HEX 기반)', '미지원'],
    ],
    col_widths=[4, 7, 7]
)

# ─────────────────────────────────────────────
# 8.11 elform
# ─────────────────────────────────────────────
add_heading(doc, '8.11 elform — 요소 공식 변경', 2)
add_para(doc, '기존 요소의 ELFORM 번호를 변경합니다. 업그레이드/다운그레이드/동일 차수 모두 지원.')
add_code_block(doc, '''
- type: elform
  target_pid: 0
  target_elform: "17"        # 숫자 또는 별칭
''')
add_table(doc,
    ['별칭 (고체)', 'ELFORM', '설명'],
    [
        ['constant_stress', '1', '상수 응력 (UR)'],
        ['fully_integrated', '2', '완전 적분'],
        ['tet4', '13', '4절점 사면체'],
        ['tet10', '17', '10절점 사면체'],
        ['hex20', '23', '20절점 헥사'],
    ],
    col_widths=[5, 3, 8]
)
add_table(doc,
    ['별칭 (셸)', 'ELFORM'],
    [
        ['belytschko_tsay', '2'],
        ['hughes_liu', '1'],
        ['fully_integrated_shell', '16'],
        ['quad8', '23'],
        ['tria6', '24'],
    ],
    col_widths=[7, 4]
)

# ─────────────────────────────────────────────
# 8.12 disconnect
# ─────────────────────────────────────────────
add_heading(doc, '8.12 disconnect — 노드 분리', 2)
add_para(doc, '지정 파트의 경계면 노드를 분리하여 비연속 인터페이스를 생성합니다. CZM 또는 MEFEM 모드를 지원합니다.')
add_code_block(doc, '''
- type: disconnect
  target_pid: 3          # 0 = 전체 파트
  mode: full             # full | czm | mefem

  # mode: czm 추가 설정
  cohesive_part_id: 0    # 0 = 자동 발급

  # mode: mefem 추가 설정
  failure_strain: 0.05   # *MAT_ADD_EROSION EPPF 값
''')
add_table(doc,
    ['모드', '동작', 'LS-DYNA 출력'],
    [
        ['full', '경계 노드 단순 분리', '노드 복사만'],
        ['czm', '분리 면에 응집 요소 삽입', '*ELEMENT_SOLID (cohesive) + *MAT_COHESIVE_*'],
        ['mefem', '미세균열 확장 파라미터 설정', '*MAT_ADD_EROSION (EPPF 값)'],
    ],
    col_widths=[3, 6, 7]
)

# ─────────────────────────────────────────────
# 8.13 iga
# ─────────────────────────────────────────────
add_heading(doc, '8.13 iga — 등기하해석 NURBS 박스 생성', 2)
add_para(doc,
    'FE solid 파트를 3D NURBS B-Spline 박스(trivariate)로 래핑하여 LS-DYNA IGA(Isogeometric Analysis) '
    '해석이 가능하도록 변환합니다. 원본 FE 메시는 TETMSH=-1 임베딩으로 그대로 유지됩니다.')

add_code_block(doc, '''
- type: iga
  targets:
    - target_pid: 1
      element_size: 4.0       # NURBS 복셀 크기 (rr=rs=rt, 필수)
      element_size_r: 2.0     # r방향 개별 지정 (0=element_size)
      element_size_s: 2.0
      element_size_t: 4.0
      offset: -1.0            # bbox 고정 확장량 (-1=auto=element_size)
      bbox_scale: 1.5         # 균일 배율 (IGA박스=파트bbox×scale)
      bbox_scale_r: 2.0       # r방향 배율 개별
      bbox_scale_s: 1.3
      bbox_scale_t: 1.0
      ir: 0                   # 0=reduced Gauss, 1=full Gauss
      styp: 4                 # LCP stabilization type
      tollg: 1.0e-3
      pr: 1  ps: 1  pt: 1     # polynomial order
      nisr: 1  niss: 1  nist: 1  # 적분점 수
''')

add_heading(doc, '8.13.1 파라미터 상세', 3)
add_table(doc,
    ['파라미터', '기본값', '설명'],
    [
        ['target_pid (필수)', '—', '대상 FE solid 파트 ID'],
        ['element_size (필수)', '1.0', 'NURBS 복셀 크기 (rr=rs=rt 공통)'],
        ['element_size_r/s/t', '0 (auto)', '축별 복셀 크기 개별 지정 (0=element_size)'],
        ['offset', '-1 (auto)', 'bbox 고정 확장량 (-1=element_size)'],
        ['bbox_scale', '0 (off)', '균일 배율: IGA박스=파트bbox×scale'],
        ['bbox_scale_r/s/t', '0 (auto)', '축별 배율 개별 지정 (0=bbox_scale)'],
        ['ir', '0', '적분 규칙 (0=reduced Gauss, 1=full Gauss)'],
        ['styp', '4', 'LCP stabilization type'],
        ['tollg', '1.0e-3', 'Light Control Point 임계값'],
        ['pr / ps / pt', '1', 'B-Spline polynomial order (각 방향)'],
        ['nisr / niss / nist', '1', '각 방향 적분점 수'],
    ],
    col_widths=[4.5, 3, 8.5]
)

add_heading(doc, '8.13.2 bbox offset 우선순위', 3)
add_para(doc, '① bbox_scale_r/s/t (축별 배율)  →  ② bbox_scale (균일)  →  ③ offset ≥ 0 (고정)  →  ④ element_size (auto)')
add_para(doc, 'bbox_scale 계산 공식:')
add_formula(doc, 'off_r = (scale_r − 1) / 2 × L_r')
add_para(doc, '예) bbox_scale=1.5, L_r=20  →  off_r = 0.25 × 20 = 5.0')

add_heading(doc, '8.13.3 생성 파일', 3)
add_table(doc,
    ['파일', '내용'],
    [
        ['<output>.k', '원본 FE 유지 + *INCLUDE <output>_iga_p{pid}.k'],
        ['<output>_iga_p{pid}.k', '파트별 IGA 파일 (PARAMETER + MAT + IGA 키워드)'],
    ],
    col_widths=[7, 9]
)

add_heading(doc, '8.13.4 MID 격리 규칙', 3)
add_note(doc, 'IGA 파트와 일반 FE 파트는 반드시 다른 MID를 사용해야 합니다. 동일 MID 공유 시 LS-DYNA 오류 발생.')
add_para(doc, '→ 각 IGA 타겟마다 새 MID 자동 발급 후 재료 카드 복사. 원본 MID는 메인 파일에서 절대 변경하지 않습니다.')

add_heading(doc, '8.13.5 NURBS 이론 기초', 3)
add_para(doc, 'B-Spline 기저 함수 (차수 p, 매듭 벡터 Ξ = {ξ₀, …, ξₘ}):')
add_formula(doc, 'N_{i,0}(ξ) = 1  if ξᵢ ≤ ξ < ξᵢ₊₁,  else 0')
add_formula(doc, 'N_{i,p}(ξ) = [(ξ−ξᵢ)/(ξᵢ₊ₚ−ξᵢ)]·N_{i,p-1}(ξ) + [(ξᵢ₊ₚ₊₁−ξ)/(ξᵢ₊ₚ₊₁−ξᵢ₊₁)]·N_{i+1,p-1}(ξ)')
add_para(doc, 'NURBS 볼륨 (가중치 wᵢⱼₖ):')
add_formula(doc, 'S(ξ,η,ζ) = Σᵢⱼₖ Nᵢ,ₚ(ξ)·Nⱼ,q(η)·Nₖ,ᵣ(ζ)·wᵢⱼₖ·Pᵢⱼₖ / Σᵢⱼₖ Nᵢ,ₚ·Nⱼ,q·Nₖ,ᵣ·wᵢⱼₖ')
add_para(doc, 'KooRemapper 생성: 2×2×2 knot + rr/rs/rt 균일 h-refinement (기하 변화 없음)')

# ─────────────────────────────────────────────
# 9. 수학 이론 종합
# ─────────────────────────────────────────────
add_heading(doc, '9. 수학 이론 종합', 1)

add_heading(doc, '9.1 선형 탄성 재료 모델 (Hooke 법칙)', 2)
add_para(doc, '라메 상수:')
add_formula(doc, 'λ = Eν / [(1+ν)(1−2ν)]')
add_formula(doc, 'μ = G = E / [2(1+ν)]')
add_para(doc, '구성 방정식 (Voigt 표기):')
add_formula(doc, '[σₓₓ σᵧᵧ σ_zz σₓᵧ σᵧ_z σₓ_z]ᵀ = C · [εₓₓ εᵧᵧ ε_zz 2εₓᵧ 2εᵧ_z 2εₓ_z]ᵀ')
add_para(doc, 'C₁₁ = C₂₂ = C₃₃ = λ+2μ,  C₁₂ = C₁₃ = C₂₃ = λ,  C₄₄ = C₅₅ = C₆₆ = μ')

add_heading(doc, '9.2 HEX8 등매개변수 요소', 2)
add_para(doc, '형상 함수 (i = 1…8):')
add_formula(doc, 'Nᵢ(ξ,η,ζ) = ⅛ (1+ξᵢξ)(1+ηᵢη)(1+ζᵢζ)')
add_para(doc, '야코비안:  Jᵢⱼ = ∂xᵢ/∂ξⱼ = Σₖ (∂Nₖ/∂ξⱼ) xₖᵢ')
add_para(doc, '수치 역매핑 (Newton-Raphson): 수렴 기준 ‖Δξ‖ < 10⁻⁸')

add_heading(doc, '9.3 Kirchhoff 판 이론', 2)
add_para(doc, '중립면에서 거리 z인 지점:')
add_formula(doc, 'εᵢⱼ(z) = −z·κᵢⱼ')
add_para(doc, '굽힘 강성:')
add_formula(doc, 'D = E·t³ / [12(1−ν²)]')
add_para(doc, '모멘트-곡률:')
add_formula(doc, 'M₁₁ = D(κ₁₁ + ν·κ₂₂),    M₂₂ = D(κ₂₂ + ν·κ₁₁),    M₁₂ = D(1−ν)·κ₁₂')

add_heading(doc, '9.4 이면각 기반 곡률 (formstrain)', 2)
add_formula(doc, 'κ = θ / L    (θ=이면각, L=인접 요소 중심 간 거리)')
add_formula(doc, 'ε_surface = ±(t/2)·κ')
add_formula(doc, 'EPS = (2/√3)·|ε_surface| = t·θ / (√3·L)')

# ─────────────────────────────────────────────
# 10. 출력 파일 형식
# ─────────────────────────────────────────────
add_heading(doc, '10. 출력 파일 형식', 1)

add_heading(doc, '10.1 dynain 파일 (*INITIAL_STRESS_SOLID)', 2)
add_code_block(doc, '''
*INITIAL_STRESS_SOLID
$#     eid    numint
   12345         1
$# hisv1~7 (미사용)
         0         0         0         0         0         0         0
$#  sig-xx    sig-yy    sig-zz    sig-xy    sig-yz    sig-xz
  5654.00  2423.00  2423.00     0.00     0.00     0.00
''')

add_heading(doc, '10.2 셸 초기 응력 (*INITIAL_STRESS_SHELL)', 2)
add_code_block(doc, '''
*INITIAL_STRESS_SHELL
$#     eid    nplane    nthick     nhisv
      100         1         2         0
$# Thickness integration point T=-1 (bottom)
$#    sig-xx    sig-yy    sig-zz    sig-xy ...
   -100.0    -80.0     0.0     0.0  ...
$# Thickness integration point T=+1 (top)
    100.0     80.0     0.0     0.0  ...
''')

add_heading(doc, '10.3 IGA 포함 메인 파일 구조', 2)
add_code_block(doc, '''
*KEYWORD
... 원본 FE 노드/요소/파트 키워드 ...
*INCLUDE
 result_iga_p1.k
*INCLUDE
 result_iga_p2.k
*END
''')

add_heading(doc, '10.4 IGA 파일 *PARAMETER_LOCAL 구조', 2)
add_code_block(doc, '''
*PARAMETER_LOCAL
$    PRMR1      VAL1
Iid            3       ← 새 파트/섹션/솔리드/리파인 ID (동일값)
Imid           2       ← 새 재료 ID (원본과 다름)
Ifepid         1       ← 원본 FE 파트 ID
Rxmin          0       ← 파트 bbox 경계
Rxmax         20
Rymin          0
Rymax         10
Rzmin          0
Rzmax          5
Rrr            4       ← r방향 복셀 크기 (h-refinement용)
Rrs            4
Rrt            4
Rofr           4       ← r방향 실제 적용 offset
Rofs           4       ← s방향 실제 적용 offset
Roft           4       ← t방향 실제 적용 offset
Iir            0
Istyp          4
Rtollg     0.001
*PARAMETER_EXPRESSION_LOCAL
rxminn, &xmin-&ofr     ← 실제 IGA 박스 범위 계산
rxmaxx, &xmax+&ofr
...
''')

# ─────────────────────────────────────────────
# 저장
# ─────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), 'KooRemapper_Manual.docx')
doc.save(out_path)
print(f'[OK] DOCX 저장 완료: {out_path}')
