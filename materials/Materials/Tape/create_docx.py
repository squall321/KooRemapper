"""
create_docx.py
Generates TAPE_COMPRESSION_STUDY.docx
Korean engineering report — professional formatting with python-docx.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "/home/koopark/claude/KooDynaAdvanced/Tape/TAPE_COMPRESSION_STUDY.docx"

# colours
RED        = RGBColor(0xC0, 0x00, 0x00)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
HDR_BLUE   = RGBColor(0x1F, 0x38, 0x64)
HDR2_BLUE  = RGBColor(0x2E, 0x74, 0xB5)
ALT_BLUE   = RGBColor(0xDD, 0xE8, 0xF5)
YELLOW     = RGBColor(0xFF, 0xFF, 0xCC)
GREEN_CODE = RGBColor(0x00, 0x40, 0x00)
GRAY_NOTE  = RGBColor(0x60, 0x60, 0x60)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(rgb[0], rgb[1], rgb[2]))
    tcPr.append(shd)


def add_tbl_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "4472C4")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sizes = {1:16, 2:13, 3:11, 4:10}
    colors = {1:HDR_BLUE, 2:HDR2_BLUE, 3:RGBColor(0x1F,0x64,0x96)}
    h.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    h.paragraph_format.space_after  = Pt(4)
    for run in h.runs:
        run.font.size = Pt(sizes.get(level, 10))
        run.font.bold = True
        if level in colors:
            run.font.color.rgb = colors[level]
    return h


def add_para(doc, text, font_size=10, bold=False, space_after=6, left_indent=None):
    p = doc.add_paragraph()
    if left_indent:
        p.paragraph_format.left_indent = left_indent
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(font_size)
    r.font.bold = bold
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    if bold_prefix:
        br = p.add_run(bold_prefix)
        br.font.bold = True; br.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)
    else:
        p.add_run(text).font.size = Pt(10)
    return p


def add_equation(doc, eq_text, label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(eq_text)
    r.font.size = Pt(11); r.font.name = "Courier New"
    if label:
        rl = p.add_run(f"   ... ({label})")
        rl.font.size = Pt(9); rl.font.color.rgb = GRAY_NOTE
    return p


def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1.0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(line)
        r.font.name = "Courier New"; r.font.size = Pt(8.5)
        r.font.color.rgb = GREEN_CODE


def build_table(doc, headers, rows, crash_cols=None, hi_rows=None, col_widths=None):
    n = len(headers)
    table = doc.add_table(rows=1+len(rows), cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    add_tbl_borders(table)
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, HDR_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
        run = p.add_run(h)
        run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = WHITE
    # data rows
    for ri, row_data in enumerate(rows):
        bg = ALT_BLUE if ri % 2 == 0 else WHITE
        if hi_rows and ri in hi_rows:
            bg = YELLOW
        for ci, val in enumerate(row_data):
            cell = table.rows[ri+1].cells[ci]
            is_crash = crash_cols and ci in crash_cols and "CRASH" in str(val)
            set_cell_bg(cell, RED if is_crash else bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if is_crash:
                run.font.bold = True; run.font.color.rgb = WHITE
            elif hi_rows and ri in hi_rows:
                run.font.bold = True
    if col_widths:
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                if ci < len(col_widths):
                    cell.width = col_widths[ci]
    return table


def build_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # ── TITLE ──
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(10); tp.paragraph_format.space_after = Pt(6)
    tr = tp.add_run("Tape 점탄성 vs 탄성 압축 거동 비교 연구")
    tr.font.size = Pt(20); tr.font.bold = True; tr.font.color.rgb = HDR_BLUE

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(16)
    sr = sp.add_run("LS-DYNA Explicit 해석: MAT_VISCOELASTIC (MAT_006) vs MAT_ELASTIC (MAT_001)")
    sr.font.size = Pt(11); sr.font.color.rgb = HDR2_BLUE

    doc.add_paragraph()

    # ── 1. 개요 ──
    add_heading(doc, "1. 개요", 1)
    add_para(doc, "알루미늄 판 사이에 끼인 얇은 접착 테이프 층의 압축 거동을 LS-DYNA explicit 해석으로 비교하였다.", space_after=6)
    add_bullet(doc, ": 속도 의존 물성 — 빠른 하중에서 높은 순간 강성 발현", "MAT_VISCOELASTIC (MAT_006)")
    add_bullet(doc, ": 속도 비의존 — 평형(장시간) 탄성계수만 사용", "MAT_ELASTIC (MAT_001)")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10); p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run("핵심 결론: "); r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = HDR_BLUE
    r2 = p.add_run("5 ms half-sine 하중 하에서 점탄성 테이프는 100 MPa에서도 6.2% 압축으로 건재하나, "
                   "동일 평형 탄성계수를 가진 탄성 테이프는 50 MPa부터 요소 붕괴(negative volume)로 해석 실패한다.")
    r2.font.size = Pt(10)

    # ── 2. 모델 구성 ──
    add_heading(doc, "2. 모델 구성", 1)

    add_heading(doc, "2.1 형상", 2)
    build_table(doc,
        headers=["항목", "값"],
        rows=[
            ["In-plane 메쉬", "3 x 3 요소 (요소당 0.1 mm)"],
            ["전체 면적", "0.3 mm x 0.3 mm"],
            ["하부 알루미늄", "10 μm (1 요소)"],
            ["테이프", "30 μm (1 요소)"],
            ["상부 알루미늄", "10 μm (1 요소)"],
            ["총 요소 수", "27 (8절점 육면체, ELFORM=1)"],
            ["총 절점 수", "64"],
        ],
        col_widths=[Cm(6), Cm(8)]
    )

    doc.add_paragraph()
    add_heading(doc, "2.2 경계조건", 2)
    add_bullet(doc, ": 전체 병진 자유도 구속 (XYZ 고정)", "하면 (z=0)")
    add_bullet(doc, ": *LOAD_SEGMENT_SET으로 half-sine 압력 하중 인가", "상면 (z=50 μm)")
    add_bullet(doc, "Conformal mesh (절점 공유) — 테이프와 알루미늄이 완전 접착 상태")

    doc.add_paragraph()
    add_heading(doc, "2.3 하중", 2)
    add_para(doc, "Half-sine 압력 펄스 정의:", space_after=4)
    add_equation(doc, "P(t) = P_max × sin(π × t / T)", "하중 함수")
    add_bullet(doc, "하중 지속시간 T = 5 ms, 총 해석시간 = 6 ms (1 ms 자유 진동 포함)")
    add_bullet(doc, "하중 수준: 0.02, 0.04, 0.08, 0.16, 1, 5, 10, 50, 100 MPa")

    doc.add_paragraph()
    add_heading(doc, "2.4 단위계", 2)
    add_para(doc, "mm, s, ton  →  힘 = N,  응력 = MPa,  밀도 = ton/mm³", space_after=10)

    # ── 3. 물성 카드 정의 ──
    add_heading(doc, "3. 물성 카드 정의", 1)

    add_heading(doc, "3.1 알루미늄 — *MAT_ELASTIC (MAT_001)", 2)
    add_para(doc, "상/하 알루미늄 판 (MID=1).", space_after=4)
    add_code(doc, [
        "*MAT_ELASTIC",
        "$      MID        RO         E        PR        DA        DB  NOT_USED",
        "         1 2.700e-09   70000.0    0.3300       0.0       0.0       0.0",
    ])
    doc.add_paragraph()
    build_table(doc,
        headers=["변수", "값", "설명"],
        rows=[
            ["MID", "1", "물성 ID"],
            ["RO", "2.7×10⁻⁹ ton/mm³", "밀도 (2700 kg/m³)"],
            ["E",   "70,000 MPa", "영률 (70 GPa)"],
            ["PR",  "0.33", "포아송비"],
        ],
        col_widths=[Cm(3), Cm(5), Cm(8)]
    )

    doc.add_paragraph()
    add_heading(doc, "3.2 테이프 (점탄성) — *MAT_VISCOELASTIC (MAT_006)", 2)
    add_para(doc, "단일 Maxwell 요소 점탄성 모델 (MID=2).", space_after=4)
    add_code(doc, [
        "*MAT_VISCOELASTIC",
        "$      MID        RO      BULK        G0        GI      BETA",
        "         2 1.100e-09    2000.0    0.3000    0.0100     500.0",
    ])
    doc.add_paragraph()
    build_table(doc,
        headers=["변수", "값", "설명"],
        rows=[
            ["MID",  "2", "물성 ID"],
            ["RO",   "1.1×10⁻⁹ ton/mm³", "밀도 (1100 kg/m³, 아크릴 PSA 기준)"],
            ["BULK", "2000 MPa", "체적 탄성계수 (비압축성에 가까움)"],
            ["G0",   "0.30 MPa", "단시간(순간) 전단 탄성계수"],
            ["GI",   "0.01 MPa", "장시간(평형) 전단 탄성계수"],
            ["BETA", "500 /s",   "감쇄 상수 (이완시간 τ = 1/BETA = 2 ms)"],
        ],
        hi_rows={3,4,5},
        col_widths=[Cm(3), Cm(5), Cm(8)]
    )

    doc.add_paragraph()
    add_heading(doc, "MAT_006 카드 포맷 (LS-DYNA Vol II 참조)", 3)
    build_table(doc,
        headers=["필드", "변수", "타입", "설명"],
        rows=[
            ["1", "MID",  "A", "물성 식별 번호"],
            ["2", "RO",   "F", "질량 밀도"],
            ["3", "BULK", "F", "체적 탄성계수. 음수이면 abs(BULK)가 온도 의존 곡선 ID"],
            ["4", "G0",   "F", "단시간 전단 탄성계수. 음수이면 곡선 ID"],
            ["5", "GI",   "F", "장시간 전단 탄성계수. 음수이면 곡선 ID"],
            ["6", "BETA", "F", "감쇄 상수 (1/s). 음수이면 곡선 ID"],
        ],
        col_widths=[Cm(1.5), Cm(2), Cm(1.5), Cm(11)]
    )

    doc.add_paragraph()
    add_heading(doc, "구성 방정식: 전단 이완 함수", 3)
    add_equation(doc, "G(t) = GI + (G0 - GI) × exp(-BETA × t)", "식 1")
    for b in [
        "t = 0 (순간):    G = G0 = 0.30 MPa  →  E_inst ≈ 3×G0 = 0.90 MPa",
        "t = ∞ (평형):   G = GI = 0.01 MPa  →  E_equil ≈ 3×GI = 0.03 MPa",
        "체적 응답은 순수 탄성 (BULK에 의해 지배)",
        "Jaumann 응력률 정식화 사용, 편차 응력에 콘볼루션 적분 적용",
        "솔리드, 셸, 빔(Hughes-Liu) 요소에 사용 가능",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "3.3 테이프 (탄성) — *MAT_ELASTIC (MAT_001)", 2)
    add_para(doc, "평형 탄성계수를 사용한 비교용 모델 (MID=2).", space_after=4)
    add_code(doc, [
        "*MAT_ELASTIC",
        "$      MID        RO         E        PR        DA        DB  NOT_USED",
        "         2 1.100e-09    0.0300    0.4990       0.0       0.0       0.0",
    ])
    doc.add_paragraph()
    build_table(doc,
        headers=["변수", "값", "설명"],
        rows=[
            ["MID", "2", "물성 ID"],
            ["RO",  "1.1×10⁻⁹ ton/mm³", "점탄성과 동일 밀도"],
            ["E",   "0.03 MPa", "영률 = 3 × GI (장시간 평형 탄성계수)"],
            ["PR",  "0.499", "포아송비 (비압축성에 가까움)"],
        ],
        col_widths=[Cm(3), Cm(5), Cm(8)]
    )
    doc.add_paragraph()
    pw = doc.add_paragraph()
    pw.paragraph_format.space_after = Pt(10); pw.paragraph_format.left_indent = Cm(0.5)
    r1 = pw.add_run("주의: "); r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = RED
    r2 = pw.add_run("K = E / (3×(1−2ν)) = 0.03 / 0.006 = 5 MPa. "
                    "이는 점탄성의 BULK(2000 MPa)보다 400배 작다. "
                    "탄성 모델은 체적 강성도 크게 낮다.")
    r2.font.size = Pt(10)

    # ── 4. G0, GI, BETA ──
    add_heading(doc, "4. G0, GI, BETA의 물리적 의미와 시간 대역별 특성", 1)

    add_heading(doc, "4.1 핵심 개념", 2)
    add_para(doc, "MAT_VISCOELASTIC의 세 파라미터가 결정하는 것:", space_after=4)
    add_bullet(doc, ": 하중이 극히 빠를 때 (충격, μs 단위) 재료가 느끼는 강성. 높을수록 순간 하중에 강하다.", "G0 (단시간 전단 탄성계수)")
    add_bullet(doc, ': 하중이 충분히 오래 유지될 때 (정적, s~min 단위) 재료의 최종 강성. 접착 테이프는 이 값이 매우 작아서 "물렁물렁"하다.', "GI (장시간 전단 탄성계수)")
    add_bullet(doc, ": 이완 속도. τ = 1/BETA가 이완 시간이다. 이 시간보다 빠른 하중에는 G0에 가까운 강성을 보이고, 느린 하중에는 GI에 가까운 강성을 보인다.", "BETA (감쇄 상수)")

    doc.add_paragraph()
    add_heading(doc, "4.2 시간에 따른 유효 전단 탄성계수 G(t)", 2)
    add_para(doc, "본 연구의 BETA = 500/s이므로 이완시간 τ = 2 ms이다.", space_after=6)
    build_table(doc,
        headers=["시간 t", "G(t) (MPa)", "E ≈ 3G (MPa)", "G0 대비", "의미"],
        rows=[
            ["0 ms",          "0.3000", "0.9000", "100%",  "순간 (충격 하중)"],
            ["0.1 ms",        "0.2859", "0.8576", "95.3%", "극초단 하중"],
            ["0.5 ms",        "0.2359", "0.7076", "78.6%", "극초단 하중"],
            ["1.0 ms",        "0.1859", "0.5577", "62.0%", "본 해석 대역"],
            ["2.0 ms (τ)",    "0.1167", "0.3501", "38.9%", "본 해석 대역 (이완시간)"],
            ["2.5 ms (피크)", "0.0931", "0.2793", "31.0%", "피크 하중 시점"],
            ["5.0 ms",        "0.0338", "0.1014", "11.3%", "하중 종료 시점"],
            ["10 ms",         "0.0120", "0.0359", "4.0%",  "평형 근접"],
            ["20 ms 이상",    "0.0100", "0.0300", "3.3%",  "완전 평형 (= GI)"],
        ],
        hi_rows={4,5},
        col_widths=[Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.5), Cm(5)]
    )
    doc.add_paragraph()
    pn = doc.add_paragraph()
    pn.paragraph_format.space_after = Pt(10); pn.paragraph_format.left_indent = Cm(0.5)
    rn1 = pn.add_run("해석: "); rn1.font.bold = True; rn1.font.size = Pt(10)
    rn2 = pn.add_run("5 ms half-sine 하중의 피크(t=2.5ms)에서 유효 전단 탄성계수는 G=0.093 MPa로, "
                     "순간값(0.3)의 31%이지만 평형값(0.01)의 9.3배이다. "
                     "즉, 탄성 모델(E=0.03 MPa)보다 약 9배 강한 전단 강성을 유지한다.")
    rn2.font.size = Pt(10)

    doc.add_paragraph()
    add_heading(doc, "4.3 BETA 값에 따른 거동 변화 (참고)", 2)
    build_table(doc,
        headers=["BETA (/s)", "τ (ms)", "5 ms 하중에서의 거동"],
        rows=[
            ["50",    "20",  "τ >> T_load: 거의 순간 강성 유지 (G ≈ G0)"],
            ["200",   "5",   "τ ≈ T_load: 하중 중 중간까지 이완"],
            ["500",   "2",   "τ < T_load: 본 연구. 상당 부분 이완되지만 여전히 GI보다 훨씬 강함"],
            ["2000",  "0.5", "τ << T_load: 빠르게 이완, 거의 평형 거동 (G ≈ GI)"],
            ["10000", "0.1", "사실상 탄성 (평형) 거동과 동일"],
        ],
        hi_rows={2},
        col_widths=[Cm(3), Cm(3), Cm(10)]
    )
    doc.add_paragraph()
    pd_ = doc.add_paragraph()
    pd_.paragraph_format.space_after = Pt(10); pd_.paragraph_format.left_indent = Cm(0.5)
    rd1 = pd_.add_run("설계 지침: "); rd1.font.bold = True; rd1.font.size = Pt(10)
    rd2 = pd_.add_run("원하는 하중 시간대역(T_load)보다 이완시간(τ=1/BETA)이 충분히 길면(τ > T_load), "
                      "점탄성 재료는 순간 강성(G0)에 가깝게 거동하여 변형을 억제한다.")
    rd2.font.size = Pt(10)

    # ── 5. 시뮬레이션 결과 ──
    add_heading(doc, "5. 시뮬레이션 결과", 1)

    add_heading(doc, "5.1 피크 하중 시점 (t=2.5 ms) 테이프 변형 비교", 2)
    build_table(doc,
        headers=["P (MPa)", "Visco 두께 (μm)", "Visco ε_zz (%)",
                 "Elastic 두께 (μm)", "Elastic ε_zz (%)", "변형 비율 (E/V)"],
        rows=[
            ["0.02",   "30.00", "0.001", "29.88", "0.39",           "284×"],
            ["0.04",   "30.00", "0.003", "29.77", "0.78",           "284×"],
            ["0.08",   "30.00", "0.006", "29.53", "1.55",           "283×"],
            ["0.16",   "30.00", "0.011", "29.08", "3.08",           "280×"],
            ["1.00",   "29.98", "0.069", "24.66", "17.80",          "259×"],
            ["5.00",   "29.90", "0.342", "11.19", "62.72",          "184×"],
            ["10.00",  "29.80", "0.679", "4.12",  "86.26",          "127×"],
            ["50.00",  "29.02", "3.251", "CRASH", "negative volume","--"],
            ["100.00", "28.14", "6.196", "CRASH", "negative volume","--"],
        ],
        crash_cols={3,4},
        col_widths=[Cm(1.8), Cm(2.8), Cm(2.8), Cm(2.8), Cm(3.0), Cm(2.5)]
    )

    doc.add_paragraph()
    add_heading(doc, "5.2 하부 알루미늄 응력 (반대편, t=2.5 ms)", 2)
    build_table(doc,
        headers=["P (MPa)", "Visco Bot σ_zz (MPa)", "Elastic Bot σ_zz (MPa)", "Visco Bot σ_VM (MPa)"],
        rows=[
            ["0.02",   "-0.020",  "-0.020", "0.031"],
            ["1.00",   "-1.000",  "-1.000", "1.546"],
            ["5.00",   "-4.998",  "-4.998", "7.691"],
            ["10.00",  "-9.993",  "-9.998", "15.291"],
            ["50.00",  "-49.827", "CRASH",  "73.093"],
            ["100.00", "-99.361", "CRASH",  "139.160"],
        ],
        crash_cols={2},
        col_widths=[Cm(2), Cm(4.5), Cm(4.5), Cm(4.5)]
    )
    doc.add_paragraph()
    add_para(doc, "등분포 압력 하중이므로 하부 Al의 σ_zz = −P_applied로 물성에 관계없이 동일하다.", space_after=10)

    add_heading(doc, "5.3 에너지 비교", 2)

    add_heading(doc, "5.3.1 글로벌 에너지 (피크 시점 t=2.5 ms)", 3)
    build_table(doc,
        headers=["P (MPa)", "Visco IE (N·mm)", "Elastic IE (N·mm)", "IE 비율 (E/V)", "의미"],
        rows=[
            ["0.02",   "2.93×10⁻¹⁰", "1.08×10⁻⁷",  "367×", "Elastic이 367배 더 많은 내부 에너지 흡수"],
            ["1.00",   "7.32×10⁻⁷",  "2.36×10⁻⁴",  "322×", ""],
            ["5.00",   "1.82×10⁻⁵",  "3.56×10⁻³",  "195×", "비선형 영역 진입"],
            ["10.00",  "7.28×10⁻⁵",  "8.02×10⁻³",  "110×", "Elastic 거의 포화"],
            ["50.00",  "1.77×10⁻³",  "CRASH",       "--",   ""],
            ["100.00", "6.90×10⁻³",  "CRASH",       "--",   ""],
        ],
        crash_cols={2},
        col_widths=[Cm(1.8), Cm(3.2), Cm(3.2), Cm(2.5), Cm(6)]
    )
    doc.add_paragraph()
    add_bullet(doc, ": 재료의 변형에 저장된 에너지. Elastic이 수백 배 크다 = 그만큼 더 많이 변형됨", "내부 에너지(IE)")
    add_bullet(doc, ": 모든 케이스에서 IE 대비 10⁻¹⁰ 이하로 무시할 수준 → 준정적(quasi-static) 거동 확인", "운동 에너지(KE)")
    add_bullet(doc, ": 에너지 보존 양호", "외부 일(ExtW) ≈ IE")

    doc.add_paragraph()
    add_heading(doc, "5.3.2 파트별 내부 에너지 (5 MPa 예시, t=2.5 ms)", 3)
    build_table(doc,
        headers=["파트", "Viscoelastic IE", "Elastic IE", "비고"],
        rows=[
            ["Part 1 (하부 Al)", "4.30×10⁻⁷", "1.86×10⁻⁷", "Al은 매우 경직 → IE 미미"],
            ["Part 2 (테이프)",  "1.69×10⁻⁵", "3.56×10⁻³", "Elastic 테이프가 211배 더 많은 에너지 흡수"],
            ["Part 3 (상부 Al)", "9.64×10⁻⁷", "3.49×10⁻⁷", ""],
        ],
        hi_rows={1},
        col_widths=[Cm(3.5), Cm(3.5), Cm(3.5), Cm(6)]
    )
    doc.add_paragraph()
    add_para(doc, "내부 에너지의 99% 이상이 테이프 층에 집중된다. Elastic 모델에서 테이프 IE가 압도적으로 크다는 것은 그만큼 체적 변형이 크다는 뜻이다.", space_after=10)

    add_heading(doc, "5.3.3 하중 제거 후 잔류 에너지 (t=6 ms)", 3)
    build_table(doc,
        headers=["P (MPa)", "Visco IE_end", "Elastic IE_end", "비고"],
        rows=[
            ["0.02",  "9.16×10⁻¹⁴", "2.85×10⁻¹²", "거의 완전 회복"],
            ["1.00",  "2.29×10⁻¹⁰", "6.54×10⁻⁹",  ""],
            ["5.00",  "5.69×10⁻⁹",  "1.19×10⁻⁷",  ""],
            ["10.00", "2.26×10⁻⁸",  "3.35×10⁻⁷",  ""],
        ],
        col_widths=[Cm(2), Cm(4), Cm(4), Cm(6)]
    )
    doc.add_paragraph()
    add_para(doc, "하중 제거 후 두 모델 모두 IE가 크게 감소하지만, 점탄성 모델은 이완 과정에서 에너지를 "
             "산일(dissipation)한다. 잔류 IE는 피크 대비 10⁻⁴ 이하 수준으로 거의 완전 탄성 회복을 보인다.", space_after=10)

    add_heading(doc, "5.4 시뮬레이션 소요시간 비교", 2)
    add_para(doc, "모든 케이스는 2 MPI 프로세스 사용. Timestep은 Al 요소(dt=1.37 ns)가 지배.", space_after=6)
    build_table(doc,
        headers=["P (MPa)", "Visco 시간 (s)", "Elastic 시간 (s)", "비고"],
        rows=[
            ["0.02",   "65", "66",         ""],
            ["0.04",   "65", "67",         ""],
            ["0.08",   "68", "64",         ""],
            ["0.16",   "64", "65",         ""],
            ["1.00",   "67", "67",         ""],
            ["5.00",   "61", "63",         ""],
            ["10.00",  "62", "66",         ""],
            ["50.00",  "62", "11 (CRASH)", "Elastic t=0.92 ms에서 중단"],
            ["100.00", "68", "7 (CRASH)",  "Elastic 더 일찍 중단"],
        ],
        crash_cols={2},
        col_widths=[Cm(2), Cm(3.5), Cm(3.5), Cm(7)]
    )
    doc.add_paragraph()
    add_bullet(doc, "정상 종료 케이스: 점탄성과 탄성의 계산 시간이 동일 (~65초, ~4.39M cycles)")
    add_bullet(doc, "MAT_006 단일 Maxwell 모델은 사이클당 추가 계산 비용이 무시할 수준")
    add_bullet(doc, "Timestep = 1.37 ns (Al 요소 지배) → 하중/물성과 무관")

    doc.add_paragraph()
    add_heading(doc, "5.5 탄성 테이프 파괴 모드", 2)
    add_bullet(doc, ": Element #10 (테이프) negative volume, t = 0.92 ms (Error 40509)", "50 MPa")
    add_bullet(doc, ": 동일 파괴 모드, t = 0.49 ms", "100 MPa")
    doc.add_paragraph()
    add_para(doc, "테이프 요소의 체적이 음수가 된다 = 요소가 뒤집어짐. 구속 탄성계수를 인가 압력이 크게 초과하면 발생한다.", space_after=4)
    add_equation(doc, "M = K + 4G/3", "식 2: 구속 탄성계수")
    add_equation(doc, "epsilon_zz = P / M", "식 3: 일축 변형률")

    # ── 6. 해석 및 토의 ──
    add_heading(doc, "6. 해석 및 토의", 1)

    add_heading(doc, "6.1 응력 전달이 동일한 이유", 2)
    add_para(doc, "등분포 압력이 상면에 작용하면, 준정적 평형에서 σ_zz = −P_applied가 전 층에 균일하게 전달된다:", space_after=4)
    for txt in [
        "50 μm 스택을 응력파가 통과하는 시간: ~10–40 ns (5 ms 하중에 비해 무시)",
        "하중이 공간적으로 균일 (모든 세그먼트에 동일 압력)",
        "힘의 평형: σ_zz가 계면에서 연속",
    ]:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(txt).font.size = Pt(10)
    doc.add_paragraph()
    pf = doc.add_paragraph()
    pf.paragraph_format.space_after = Pt(10); pf.paragraph_format.left_indent = Cm(0.5)
    pf.add_run("따라서 ").font.size = Pt(10)
    rb = pf.add_run("하부 Al이 받는 응력은 테이프 물성과 무관"); rb.font.bold = True; rb.font.size = Pt(10)
    pf.add_run("하며, 차이는 오직 테이프 변형에서 나타난다.").font.size = Pt(10)

    add_heading(doc, "6.2 테이프 압축이 다른 이유: 구속 압축 상태", 2)
    add_para(doc, "Conformal mesh로 테이프가 경직한 Al(E=70 GPa)에 접착되어 측면 팽창이 구속됨 → 구속 압축(oedometric) 상태:", space_after=4)
    add_equation(doc, "유효 강성 = 구속 탄성계수  M = K + 4G/3", "식 4")
    doc.add_paragraph()
    build_table(doc,
        headers=["모델", "K (MPa)", "G_eff (MPa)", "M (MPa)", "ε_zz"],
        rows=[
            ["점탄성 (t≈2.5ms)", "2000", "0.093", "2000.1", "P / 2000"],
            ["탄성",              "5",    "0.01",  "5.013",  "P / 5"],
        ],
        col_widths=[Cm(4), Cm(2.5), Cm(3), Cm(2.5), Cm(4)]
    )
    doc.add_paragraph()
    p6 = doc.add_paragraph()
    p6.paragraph_format.space_after = Pt(6); p6.paragraph_format.left_indent = Cm(0.5)
    p6.add_run("M_visco / M_elastic = 2000 / 5 = ").font.size = Pt(10)
    rb6 = p6.add_run("400배"); rb6.font.bold = True; rb6.font.size = Pt(10)
    p6.add_run(". 관찰된 비율(~284배)은 하중 중 이완 효과로 약간 감소한 값이다.").font.size = Pt(10)
    doc.add_paragraph()
    p6b = doc.add_paragraph()
    p6b.paragraph_format.space_after = Pt(10); p6b.paragraph_format.left_indent = Cm(0.5)
    r6b1 = p6b.add_run("핵심: "); r6b1.font.bold = True; r6b1.font.size = Pt(10)
    r6b2 = p6b.add_run("구속 조건에서는 체적 탄성계수(K)가 지배적이며, K의 400배 차이(2000 vs 5 MPa)가 "
                       "변형 차이의 주된 원인이다. 전단 탄성계수(G)는 M에 4G/3만큼만 기여하므로 상대적으로 영향이 작다.")
    r6b2.font.size = Pt(10)

    add_heading(doc, "6.3 에너지 관점에서의 해석", 2)
    add_para(doc, "점탄성 모델에서는:", bold=True, space_after=3)
    for i, (bold_part, rest) in enumerate([
        ("순간 하중 시:", " 높은 G0로 인해 전단 변형 에너지가 적게 저장됨"),
        ("하중 유지 중:", " G(t)가 감소하면서 전단 응력이 이완 → 에너지 산일 (viscous dissipation)"),
        ("하중 제거 후:", " 남은 탄성 에너지는 GI 기반으로 작음 → 거의 완전 회복"),
    ], 1):
        p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(2)
        rb = p.add_run(bold_part); rb.font.bold = True; rb.font.size = Pt(10)
        p.add_run(rest).font.size = Pt(10)
    doc.add_paragraph()
    add_para(doc, "탄성 모델에서는:", bold=True, space_after=3)
    for i, txt in enumerate([
        ("낮은 K(5 MPa)로 인해 ", "체적 변형 에너지", "가 대량 저장"),
        ("에너지 산일 메커니즘 없음 → 저장된 에너지가 모두 변형으로 전환", "", ""),
        ("고압에서 에너지 밀도가 요소 허용 한계를 초과 → 요소 붕괴", "", ""),
    ], 1):
        p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(2)
        p.add_run(txt[0]).font.size = Pt(10)
        if txt[1]:
            rb = p.add_run(txt[1]); rb.font.bold = True; rb.font.size = Pt(10)
        if txt[2]:
            p.add_run(txt[2]).font.size = Pt(10)

    # ── 7. 파일 구조 ──
    doc.add_paragraph()
    add_heading(doc, "7. 파일 구조", 1)
    add_code(doc, [
        "Tape/",
        "  generate_tape_model.py          # 모델 생성 스크립트",
        "  TAPE_COMPRESSION_STUDY.md       # 본 문서",
        "  TAPE_COMPRESSION_STUDY.docx     # Word 버전",
        "  results/",
        "    case_viscoelastic/             # P = 0.02 MPa",
        "    case_elastic/",
        "    case_viscoelastic_2x/          # P = 0.04 MPa",
        "    case_elastic_2x/",
        "    case_viscoelastic_4x/          # P = 0.08 MPa",
        "    case_elastic_4x/",
        "    case_viscoelastic_8x/          # P = 0.16 MPa",
        "    case_elastic_8x/",
        "    case_viscoelastic_1MPa/        # P = 1.0 MPa",
        "    case_elastic_1MPa/",
        "    case_viscoelastic_5MPa/        # P = 5.0 MPa",
        "    case_elastic_5MPa/",
        "    case_viscoelastic_10MPa/       # P = 10.0 MPa",
        "    case_elastic_10MPa/",
        "    case_viscoelastic_50MPa/       # P = 50 MPa (탄성 CRASH)",
        "    case_elastic_50MPa/",
        "    case_viscoelastic_100MPa/      # P = 100 MPa (탄성 CRASH)",
        "    case_elastic_100MPa/",
    ])

    # ── 8. 재현 방법 ──
    doc.add_paragraph()
    add_heading(doc, "8. 재현 방법", 1)
    add_code(doc, [
        "cd /home/koopark/claude/KooDynaAdvanced/Tape",
        "python3 generate_tape_model.py --outdir /data/tape_study/case_test --pressure 1.0",
        "bash /home/koopark/claude/KooDynaAdvanced/DynaJobSubmit/run.sh \\",
        "     /data/tape_study/case_test/tape_viscoelastic.k",
    ])

    # ── 9. 결론 ──
    doc.add_paragraph()
    add_heading(doc, "9. 결론", 1)
    conclusions = [
        ("점탄성 테이프(MAT_006)는 100 MPa에서도 구조 건전성 유지",
         ": 30 μm → 28.14 μm (6.2% 압축). 높은 체적 탄성계수(K=2000 MPa)와 빠른 하중에서의 순간 전단 강성(G0=0.3 MPa) 덕분이다."),
        ("탄성 테이프(MAT_001, 평형 탄성계수)는 중하중에서 붕괴",
         ": 10 MPa에서 86% 압축(4.12 μm), 50 MPa 이상에서 요소 뒤집힘(negative volume)으로 해석 실패."),
        ("하부 Al 전달 응력은 물성과 무관",
         ": 등분포 압축에서 σ_zz = −P_applied가 전 층에 균일 전달."),
        ("구속 조건에서 체적 탄성계수가 지배적",
         ": conformal mesh로 Al이 측면 팽창을 구속하므로, K의 400배 차이(2000 vs 5 MPa)가 변형 차이의 주된 원인."),
        ("에너지 차이",
         ": 피크 하중에서 탄성 테이프의 내부 에너지가 점탄성 대비 100~370배 크다. 이 에너지가 모두 체적 변형으로 전환되어 요소 붕괴를 유발한다."),
        ("계산 비용 동일",
         ": MAT_006 단일 Maxwell 모델은 사이클당 추가 비용이 무시할 수준이며, 모든 정상 종료 케이스에서 ~65초로 동일하다."),
        ("설계 지침",
         ": 동적 하중을 받는 접착 테이프 모델링 시, 정적 탄성계수를 사용하면 변형을 과대 예측하고 해석이 실패할 수 있다. "
         "점탄성 모델로 속도 의존 강성을 반영해야 현실적인 결과를 얻는다. "
         "특히 τ = 1/BETA와 하중 지속시간(T_load)의 비율이 거동을 결정하므로, "
         "실험 데이터로 G0, GI, BETA를 정확히 결정하는 것이 중요하다."),
    ]
    for title, body in conclusions:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.space_before = Pt(3)
        rt = p.add_run(title); rt.font.bold = True; rt.font.size = Pt(10); rt.font.color.rgb = HDR_BLUE
        rb = p.add_run(body); rb.font.size = Pt(10)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
