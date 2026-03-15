"""
Convert KooRemapper_Manual.md to KooRemapper_Manual.docx
Parses markdown headings, paragraphs, code blocks, tables, blockquotes, lists, math
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def add_run_with_style(paragraph, text, bold=False, italic=False, code=False, size=None):
    """Add a run with formatting."""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    if size:
        run.font.size = Pt(size)
    return run

def parse_inline(paragraph, text):
    """Parse inline markdown: **bold**, *italic*, `code`, $$math$$"""
    # Pattern for inline elements
    pattern = re.compile(r'(\$\$.*?\$\$|\$[^$]+\$|`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)')
    pos = 0
    for m in pattern.finditer(text):
        # Add text before match
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group()
        if token.startswith('$$'):
            # Display math - just italicize
            add_run_with_style(paragraph, token[2:-2], italic=True, code=True)
        elif token.startswith('$') and not token.startswith('$$'):
            add_run_with_style(paragraph, token[1:-1], italic=True, code=True)
        elif token.startswith('`'):
            add_run_with_style(paragraph, token[1:-1], code=True)
        elif token.startswith('**'):
            add_run_with_style(paragraph, token[2:-2], bold=True)
        elif token.startswith('*'):
            add_run_with_style(paragraph, token[1:-1], italic=True)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def md_to_docx(md_path, docx_path):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Configure heading styles
    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Malgun Gothic'
        if i == 1:
            hs.font.size = Pt(22)
            hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        elif i == 2:
            hs.font.size = Pt(16)
            hs.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        elif i == 3:
            hs.font.size = Pt(13)
            hs.font.color.rgb = RGBColor(0x34, 0x49, 0x5e)
        elif i == 4:
            hs.font.size = Pt(11)
            hs.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # Heading
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped[level:].strip()
            if level <= 4:
                doc.add_heading(text, level=level)
            else:
                p = doc.add_paragraph()
                add_run_with_style(p, text, bold=True, size=10)
            i += 1
            continue

        # Code block
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # skip closing ```

            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)

            # Add light gray shading to code block paragraph
            pPr = p._p.get_or_add_pPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F5F5F5')
            shading.set(qn('w:val'), 'clear')
            pPr.append(shading)

            continue

        # Table
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_rows = []
            while i < n and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                # Skip separator row
                if re.match(r'\|[-:\s|]+\|', row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1

            if table_rows:
                num_cols = max(len(r) for r in table_rows)
                # Pad rows
                for r in table_rows:
                    while len(r) < num_cols:
                        r.append('')

                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.LEFT

                for ri, row_data in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_data):
                        cell = table.rows[ri].cells[ci]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(1)
                        p.paragraph_format.space_after = Pt(1)
                        parse_inline(p, cell_text)
                        for run in p.runs:
                            run.font.size = Pt(9)

                        # Header row styling
                        if ri == 0:
                            set_cell_shading(cell, '2C3E50')
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.bold = True

                doc.add_paragraph()  # spacing after table
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_text = stripped[1:].strip()
            # Continue collecting blockquote lines
            i += 1
            while i < n and lines[i].strip().startswith('>'):
                quote_text += ' ' + lines[i].strip()[1:].strip()
                i += 1

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

            # Add left border via shading
            pPr = p._p.get_or_add_pPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'EBF5FB')
            shading.set(qn('w:val'), 'clear')
            pPr.append(shading)

            parse_inline(p, quote_text)
            for run in p.runs:
                run.font.size = Pt(9)
                run.italic = True
            continue

        # Display math (standalone $$...$$)
        if stripped.startswith('$$'):
            inner = stripped[2:]
            # Check if closing $$ is on the same line
            if '$$' in inner:
                math_text = inner[:inner.index('$$')]
                i += 1
            else:
                math_text = inner
                i += 1
                while i < n and '$$' not in lines[i]:
                    math_text += ' ' + lines[i].strip()
                    i += 1
                if i < n:
                    closing = lines[i].strip()
                    math_text += ' ' + closing[:closing.index('$$')]
                    i += 1

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(math_text.strip())
            run.font.name = 'Cambria Math'
            run.font.size = Pt(10)
            run.italic = True
            continue

        # List items
        if re.match(r'^(\d+)\.\s', stripped) or stripped.startswith('- '):
            if stripped.startswith('- '):
                text = stripped[2:]
                p = doc.add_paragraph(style='List Bullet')
            else:
                m = re.match(r'^(\d+)\.\s(.*)', stripped)
                text = m.group(2)
                p = doc.add_paragraph(style='List Number')

            p.text = ''
            parse_inline(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f'Saved: {docx_path}')

if __name__ == '__main__':
    md_to_docx(
        r'd:\KooRemapper\docs\KooRemapper_Manual.md',
        r'd:\KooRemapper\docs\KooRemapper_Manual_new.docx'
    )
